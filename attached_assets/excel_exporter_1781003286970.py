"""Excel workbook exporter for Survey Insight Engine single cuts."""

from __future__ import annotations

from collections import defaultdict
from dataclasses import dataclass, field
from datetime import datetime, timezone
import gc
from io import BytesIO
import logging
import math
import os
from pathlib import Path
import re
import sys
import tempfile
from typing import Any
import xml.etree.ElementTree as ET
from xml.sax.saxutils import escape
from zipfile import ZIP_DEFLATED, ZipFile

try:
    import xlsxwriter
except ModuleNotFoundError:
    xlsxwriter = None

from src.calculation_log import CalculationLog
from src.filter_options import NPS_BUCKETS, _is_computed_multi_select_column, filter_question_options
from src.models import (
    AuditRecord,
    AnalysisType,
    CrossCutResult,
    DataQualityReport,
    FilteredSingleCutResult,
    FilterSpec,
    GridBinaryPivotResult,
    GridRatedResult,
    GridSingleSelectResult,
    HypothesisResult,
    MultiSelectResult,
    NPSResult,
    NumericResult,
    QuestionType,
    RankOrderResult,
    SegmentDefinition,
    SingleCutResult,
    SingleSelectResult,
    SkipRecord,
    SurveySchema,
)
from src.single_cut import compute_multi_select, compute_single_select


_CHART_COLORS = [
    "#CC0000",
    "#0A0A0A",
    "#666666",
    "#990000",
    "#999999",
    "#330000",
    "#444444",
    "#FF6666",
]

PASS_WORKBOOK_FILTERS_DATA_NAME = "passes_workbook_filters_data"
PASS_WORKBOOK_CUSTOM_FILTERS_DATA_NAME = "passes_workbook_custom_filters_data"
_EXCEL_SHORT_LABEL_CACHE: dict[str, Any] = {}
_RAW_DATA_STREAMING_ROW_THRESHOLD = 2000
DEFAULT_WORKBOOK_CUSTOM_FILTER_COUNT = 2
DEFAULT_PER_QUESTION_FILTER_COUNT = 1
INPUT_RAW_SHEET_NAME = "Raw Data (Input)"
INPUT_DATAMAP_SHEET_NAME = "Data Map (Input)"
RANK_CROSS_TAB_METRICS = (
    "Weighted Average",
    "Sum of ranks",
    "Rank position count",
)

try:
    from config import CROSS_TAB_MAX_GROUPS, RAW_DATA_SHEET_ROW_LIMIT
except ModuleNotFoundError:
    RAW_DATA_SHEET_ROW_LIMIT = 50000
    CROSS_TAB_MAX_GROUPS = 12

GRID_RATED = "GRID_RATED"
GRID_CATEGORICAL = "GRID_CATEGORICAL"
GRID_BINARY_SELECT = "GRID_BINARY_SELECT"
_FORMULA_DEFINED_NAME_PATTERN = re.compile(
    r"\b([A-Za-z_][A-Za-z0-9_]*(?:_data|_resolved|_wrapped))\b"
)
_FORMULA_STRING_LITERAL_PATTERN = re.compile(r'"(?:[^"]|"")*"')


def _normalise_slot_count(value: int | None, default: int, maximum: int) -> int:
    try:
        count = int(value if value is not None else default)
    except (TypeError, ValueError):
        count = default
    return max(0, min(maximum, count))


def export_single_cuts(
    results: list[SingleCutResult],
    skips: list[SkipRecord],
    schema: SurveySchema,
    quality_report: DataQualityReport,
    log: CalculationLog,
    output_path: str,
    cross_cut_results: list[CrossCutResult] | None = None,
    cross_cut_skips: list[SkipRecord] | None = None,
    themes: dict | None = None,
    decoded_df: Any | None = None,
    demo_priority: dict | None = None,
    short_labels: dict[str, str] | None = None,
    workbook_custom_filter_count: int = DEFAULT_WORKBOOK_CUSTOM_FILTER_COUNT,
    per_question_filter_count: int = DEFAULT_PER_QUESTION_FILTER_COUNT,
    rank_cross_tab_settings: dict[str, Any] | None = None,
    hypothesis_results: list[HypothesisResult] | None = None,
    embed_input_files: bool = False,
    input_file_sources: dict[str, Any] | None = None,
    strict_formula_name_validation: bool = False,
) -> None:
    """Write a complete live-filterable single-cut workbook."""

    from openpyxl import Workbook
    from src.ai_insights import generate_short_labels
    from src import memory_profiler as export_memory

    if os.environ.get("SURVEY_PROFILE_MEMORY") == "1":
        export_memory.reset_log()
        export_memory.enable_profiling()
    elif export_memory.is_profiling_enabled():
        export_memory.reset_log()

    with export_memory.memory_step("load_or_receive_decoded_df"):
        cross_cut_results = cross_cut_results or []
        cross_cut_skips = cross_cut_skips or []

    workbook = Workbook()
    default_sheet = workbook.active
    workbook.remove(default_sheet)

    with export_memory.memory_step("generate_short_labels"):
        questions_for_labels = [
            {"question_id": q.canonical_id, "question_text": q.question_text}
            for q in schema.questions
        ]
        short_labels_map = dict(
            generate_short_labels(
                questions_for_labels,
                cache=_EXCEL_SHORT_LABEL_CACHE,
            )
        )
        explicit_short_labels: dict[str, str] = {}
        if short_labels:
            short_labels_map.update(short_labels)
            explicit_short_labels = dict(short_labels)

    live_context = _LiveWorkbookContext(
        schema=schema,
        results=results,
        log=log,
        short_labels_map=short_labels_map,
        explicit_short_labels=explicit_short_labels,
        workbook_custom_filter_count=_normalise_slot_count(
            workbook_custom_filter_count,
            DEFAULT_WORKBOOK_CUSTOM_FILTER_COUNT,
            5,
        ),
        per_question_filter_count=_normalise_slot_count(
            per_question_filter_count,
            DEFAULT_PER_QUESTION_FILTER_COUNT,
            3,
        ),
        rank_cross_tab_settings=dict(rank_cross_tab_settings or {}),
    )
    used_sheet_names = _reserved_sheet_names().union({"_RawData", "_Options", "Filters"})
    theme_groups = _theme_groups_for_results(schema, results, themes)
    theme_sheet_names = {
        theme_name: _unique_sheet_name(_safe_sheet_name(theme_name), used_sheet_names)
        for theme_name, _theme_results in theme_groups
    }
    result_theme_names = {
        result.question_id: theme_name
        for theme_name, theme_results in theme_groups
        for result in theme_results
    }
    nps_sheet_names = {
        result.question_id: _unique_sheet_name(f"SC_{result.question_id}", used_sheet_names)
        for result in results
        if isinstance(result, NPSResult)
    }
    result_sheet_names = {
        result.question_id: theme_sheet_names[result_theme_names[result.question_id]]
        for result in results
        if result.question_id in result_theme_names
    }
    result_sheet_names.update(nps_sheet_names)

    raw_row_count = _raw_data_source_row_count(decoded_df, schema, results)
    stream_raw_data = _should_stream_raw_data(raw_row_count)
    with export_memory.memory_step("build_raw_data_sheet"):
        if stream_raw_data:
            _prepare_live_columns(schema, live_context)
            demo_questions = _ordered_demographic_questions(schema, demo_priority)
            sheet_filters = _planned_sheet_filters(demo_questions, live_context, demo_priority)
            _build_raw_data_sheet_streaming(
                workbook,
                decoded_df,
                schema,
                results,
                live_context,
                sheet_filters,
                theme_groups,
                theme_sheet_names,
            )
        else:
            _build_raw_data_sheet(workbook, decoded_df, schema, results, live_context)
    with export_memory.memory_step("build_options_sheet"):
        _build_options_sheet(workbook, schema, live_context)
    demo_questions = _ordered_demographic_questions(schema, demo_priority)
    with export_memory.memory_step("build_filters_sheet"):
        sheet_filters = _build_filters_sheet(
            workbook,
            schema,
            results,
            demo_questions,
            live_context,
            demo_priority,
            planned_filters=sheet_filters if stream_raw_data else None,
        )
    with export_memory.memory_step("build_helper_columns"):
        _build_raw_filter_helper_columns(workbook, sheet_filters, live_context)

    with export_memory.memory_step("build_run_summary_sheet"):
        _live_write_run_summary(
            workbook,
            schema,
            quality_report,
            results,
            skips,
            log,
            cross_cut_results,
            cross_cut_skips,
        )
    with export_memory.memory_step("build_question_metadata_sheet"):
        _live_write_question_metadata(workbook, schema)

    with export_memory.memory_step("build_single_cut_index_sheet"):
        _live_write_single_cut_index(
            workbook,
            results,
            result_sheet_names,
            result_theme_names,
        )

    with export_memory.memory_step("build_theme_sheets"):
        for theme_name, theme_results in theme_groups:
            _live_write_theme_sheet(
                workbook,
                theme_name,
                theme_sheet_names[theme_name],
                theme_results,
                schema,
                live_context,
                sheet_filters,
                live_context.short_labels_map,
            )
    with export_memory.memory_step("build_nps_sheets"):
        for result in results:
            if not isinstance(result, NPSResult):
                continue
            question = schema.get_question(result.question_id)
            sheet_name = nps_sheet_names.get(result.question_id)
            if question is None or sheet_name is None:
                continue
            _live_write_nps_sheet(
                workbook,
                sheet_name,
                result,
                question,
                live_context,
                sheet_filters,
            )

    with export_memory.memory_step("build_calculation_log_sheet"):
        _live_write_calculation_log(workbook, log)
    with export_memory.memory_step("build_hypothesis_check_sheet"):
        _live_write_hypothesis_check_sheet(
            workbook,
            schema,
            hypothesis_results or [],
            log,
        )
    with export_memory.memory_step("build_filter_log_sheet"):
        _live_write_filter_log(workbook, cross_cut_results)
    with export_memory.memory_step("build_warnings_sheet"):
        _live_write_warnings(
            workbook,
            schema,
            quality_report,
            results,
            skips,
            cross_cut_results,
            cross_cut_skips,
        )
    with export_memory.memory_step("embed_input_files"):
        _embed_input_file_sheets_openpyxl(
            workbook,
            schema,
            embed_input_files=embed_input_files,
            input_file_sources=input_file_sources,
        )
    with export_memory.memory_step("save_workbook"):
        _validate_formula_names_for_export(
            workbook,
            strict=strict_formula_name_validation,
        )
        workbook.save(output_path)
        _replace_streamed_raw_data_sheet(output_path, workbook)
    _write_formula_caches(output_path, workbook)
    _write_memory_report_if_enabled(output_path, export_memory)


def write_workbook(
    output_path: str,
    schema: SurveySchema,
    results: list[SingleCutResult],
    skips: list[SkipRecord],
    audit_log: CalculationLog,
    filter_log: Any | None = None,
    decoded_df: Any | None = None,
    themes: dict | None = None,
    cross_cut_results: list[CrossCutResult] | None = None,
    segmentation_result: Any | None = None,
    quality_report: DataQualityReport | None = None,
    cross_cut_skips: list[SkipRecord] | None = None,
    demo_priority: dict | None = None,
    short_labels: dict[str, str] | None = None,
    workbook_custom_filter_count: int = DEFAULT_WORKBOOK_CUSTOM_FILTER_COUNT,
    per_question_filter_count: int = DEFAULT_PER_QUESTION_FILTER_COUNT,
    hypothesis_results: list[HypothesisResult] | None = None,
    embed_input_files: bool = False,
    input_file_sources: dict[str, Any] | None = None,
    strict_formula_name_validation: bool = False,
) -> None:
    """Compatibility wrapper for callers using the proposed workbook API."""

    del filter_log, segmentation_result
    if quality_report is None:
        row_count = int(getattr(decoded_df, "shape", [schema.total_respondents, 0])[0])
        col_count = int(getattr(decoded_df, "shape", [0, 0])[1])
        quality_report = DataQualityReport(
            total_rows=row_count,
            total_columns=col_count,
            columns_in_datamap=len(schema.questions),
            columns_not_in_datamap=(),
            per_column_missing_pct={},
            per_column_out_of_range_pct={},
            coercion_log=(),
            warnings=(),
        )
    export_single_cuts(
        results=results,
        skips=skips,
        schema=schema,
        quality_report=quality_report,
        log=audit_log,
        output_path=output_path,
        cross_cut_results=cross_cut_results,
        cross_cut_skips=cross_cut_skips,
        themes=themes,
        decoded_df=decoded_df,
        demo_priority=demo_priority,
        short_labels=short_labels,
        workbook_custom_filter_count=workbook_custom_filter_count,
        per_question_filter_count=per_question_filter_count,
        hypothesis_results=hypothesis_results,
        embed_input_files=embed_input_files,
        input_file_sources=input_file_sources,
        strict_formula_name_validation=strict_formula_name_validation,
    )


def export_winners_vs_laggards_workbook(
    output_path: Path | str,
    decoded_df: Any,
    schema: SurveySchema,
    single_cut_results: dict[str, SingleCutResult] | list[SingleCutResult],
    segment_definition: SegmentDefinition,
    laggard_outcome_question_id: str | None = None,
    laggard_segment_definition: SegmentDefinition | None = None,
    workbook_filter_state: Any | None = None,
    themes: dict | None = None,
    short_labels: dict[str, str] | None = None,
    workbook_custom_filter_count: int = DEFAULT_WORKBOOK_CUSTOM_FILTER_COUNT,
    per_question_filter_count: int = DEFAULT_PER_QUESTION_FILTER_COUNT,
    calculation_log: CalculationLog | None = None,
    embed_input_files: bool = False,
    input_file_sources: dict[str, Any] | None = None,
    strict_formula_name_validation: bool = False,
) -> None:
    """Write a live-filterable winners-vs-laggards single-cut workbook."""

    del workbook_filter_state, short_labels
    from dataclasses import replace
    from openpyxl import Workbook
    from src.outcome_segmentation import _build_segment_masks

    results = (
        list(single_cut_results.values())
        if isinstance(single_cut_results, dict)
        else list(single_cut_results)
    )
    log = calculation_log if calculation_log is not None else CalculationLog()

    workbook = Workbook()
    workbook.remove(workbook.active)
    context = _LiveWorkbookContext(
        schema=schema,
        results=results,
        log=log,
        workbook_custom_filter_count=_normalise_slot_count(
            workbook_custom_filter_count,
            DEFAULT_WORKBOOK_CUSTOM_FILTER_COUNT,
            5,
        ),
        per_question_filter_count=_normalise_slot_count(
            per_question_filter_count,
            DEFAULT_PER_QUESTION_FILTER_COUNT,
            3,
        ),
    )
    _prepare_live_columns(schema, context)

    if segment_definition.segment_mode == "manual_uuid":
        winner_spec = None
        laggard_spec = None
        laggard_definition = laggard_segment_definition or segment_definition
        winner_mask, laggard_mask, _valid_mask, _warnings = _build_segment_masks(
            decoded_df,
            None,
            segment_definition,
            respondent_id_column=schema.respondent_id_column,
        )
        _record_wvl_manual_cohort_audit(log, "winner_mask", schema, segment_definition, int(winner_mask.sum()))
        _record_wvl_manual_cohort_audit(log, "laggard_mask", schema, segment_definition, int(laggard_mask.sum()))
    else:
        winner_spec = schema.get_question(segment_definition.outcome_question_id)
        if winner_spec is None:
            raise ValueError(
                f"winner outcome question {segment_definition.outcome_question_id!r} not in schema"
            )
        winner_mask, _unused_laggard_mask, _valid_mask, _warnings = _build_segment_masks(
            decoded_df,
            winner_spec,
            segment_definition,
            respondent_id_column=schema.respondent_id_column,
        )

        laggard_definition = laggard_segment_definition or segment_definition
        laggard_qid = (
            laggard_outcome_question_id
            or laggard_definition.laggard_outcome_question_id
            or segment_definition.laggard_outcome_question_id
            or segment_definition.outcome_question_id
        )
        laggard_spec = schema.get_question(laggard_qid)
        if laggard_spec is None:
            raise ValueError(f"laggard outcome question {laggard_qid!r} not in schema")
        laggard_sub_question_id = (
            laggard_definition.laggard_outcome_sub_question_id
            or segment_definition.laggard_outcome_sub_question_id
        )
        if laggard_sub_question_id:
            laggard_spec = replace(laggard_spec, raw_columns=(laggard_sub_question_id,))
        laggard_definition = replace(laggard_definition, outcome_question_id=laggard_qid)
        _lag_winner_mask, laggard_mask, _lag_valid_mask, _lag_warnings = _build_segment_masks(
            decoded_df,
            laggard_spec,
            laggard_definition,
            respondent_id_column=schema.respondent_id_column,
        )

        _record_wvl_cohort_audit(
            log,
            "winner_mask",
            winner_spec,
            segment_definition,
            int(winner_mask.sum()),
        )
        _record_wvl_cohort_audit(
            log,
            "laggard_mask",
            laggard_spec,
            laggard_definition,
            int(laggard_mask.sum()),
        )

    winner_mask = winner_mask.reindex(decoded_df.index, fill_value=False)
    laggard_mask = laggard_mask.reindex(decoded_df.index, fill_value=False)
    others_mask = ~(winner_mask | laggard_mask)

    _wvl_build_raw_data_sheet(workbook, decoded_df, schema, results, context, winner_mask, laggard_mask)
    _build_options_sheet(workbook, schema, context)
    demo_questions = _ordered_demographic_questions(schema, None)
    sheet_filters = _build_filters_sheet(
        workbook,
        schema,
        results,
        demo_questions,
        context,
        None,
    )
    _build_raw_filter_helper_columns(workbook, sheet_filters, context)
    _wvl_write_run_summary(
        workbook,
        schema,
        segment_definition,
        laggard_definition,
        int(winner_mask.sum()),
        int(laggard_mask.sum()),
    )
    _wvl_write_cohort_definition_sheet(
        workbook,
        segment_definition,
        int(winner_mask.sum()),
        int(laggard_mask.sum()),
    )
    _live_write_question_metadata(workbook, schema)

    used_sheet_names = _reserved_sheet_names().union({"_RawData", "_Options", "Filters"})
    theme_groups = _theme_groups_for_results(schema, results, themes)
    theme_sheet_names = {
        theme_name: _unique_sheet_name(_safe_sheet_name(theme_name), used_sheet_names)
        for theme_name, _theme_results in theme_groups
    }
    result_theme_names = {
        result.question_id: theme_name
        for theme_name, theme_results in theme_groups
        for result in theme_results
    }
    result_sheet_names = {
        result.question_id: theme_sheet_names[result_theme_names[result.question_id]]
        for result in results
        if result.question_id in result_theme_names
    }
    _live_write_single_cut_index(workbook, results, result_sheet_names, result_theme_names)

    for theme_name, theme_results in theme_groups:
        sheet_name = theme_sheet_names[theme_name]
        worksheet = workbook.create_sheet(sheet_name)
        theme_prefix = _theme_defined_prefix(sheet_name)
        row_index, local_filters = _build_theme_local_filters_block(
            workbook,
            worksheet,
            1,
            theme_prefix,
            sheet_filters,
            context,
        )
        _build_theme_local_filter_helper_column(
            workbook,
            theme_prefix,
            local_filters,
            context,
        )
        row_index = _wvl_write_category_sheet_header(worksheet, theme_name, row_index)
        for result in theme_results:
            question = schema.get_question(result.question_id)
            if question is None:
                continue
            row_index = _wvl_write_question_block(
                workbook,
                worksheet,
                row_index,
                result,
                question,
                context,
                log,
                segment_definition.winner_label,
                laggard_definition.laggard_label or laggard_definition.loser_label,
                schema,
                local_filters,
                theme_prefix,
                decoded_df,
                winner_mask,
                laggard_mask,
                others_mask,
            )
            row_index += 2
        _wvl_apply_arial_to_used_cells(worksheet)
        _live_autofit(worksheet)

    _live_write_calculation_log(workbook, log)
    _live_write_filter_log(workbook, [])
    _wvl_apply_arial_to_used_cells(workbook["Run_Summary"])
    _wvl_apply_arial_to_used_cells(workbook["Question_Metadata"])
    _wvl_apply_arial_to_used_cells(workbook["Single_Cut_Index"])
    _wvl_apply_arial_to_used_cells(workbook["Calculation_Log"])
    _wvl_apply_arial_to_used_cells(workbook["Filter_Log"])
    _embed_input_file_sheets_openpyxl(
        workbook,
        schema,
        embed_input_files=embed_input_files,
        input_file_sources=input_file_sources,
    )
    _validate_formula_names_for_export(
        workbook,
        strict=strict_formula_name_validation,
    )
    workbook.save(str(output_path))
    _write_formula_caches(str(output_path), workbook)


@dataclass
class _LiveColumnSpec:
    key: str
    header: str
    data_name: str
    question: Any | None
    source_column: str | None
    kind: str


def _column_data_name(header: str) -> str:
    return f"{header.rstrip('_')}_data"


def _nps_filter_column_key(question_id: str) -> str:
    return f"{question_id}__nps_filter"


@dataclass
class _LiveWorkbookContext:
    schema: SurveySchema
    results: list[SingleCutResult]
    log: CalculationLog | None = None
    short_labels_map: dict[str, str] = field(default_factory=dict)
    explicit_short_labels: dict[str, str] = field(default_factory=dict)
    columns: list[_LiveColumnSpec] = field(default_factory=list)
    column_by_key: dict[str, _LiveColumnSpec] = field(default_factory=dict)
    option_values: dict[str, list[str]] = field(default_factory=dict)
    warnings: list[str] = field(default_factory=list)
    raw_data_streamed: bool = False
    raw_data_row_count: int = 0
    workbook_custom_filter_count: int = DEFAULT_WORKBOOK_CUSTOM_FILTER_COUNT
    per_question_filter_count: int = DEFAULT_PER_QUESTION_FILTER_COUNT
    rank_cross_tab_settings: dict[str, Any] = field(default_factory=dict)


def _wvl_build_raw_data_sheet(
    workbook: Any,
    decoded_df: Any,
    schema: SurveySchema,
    results: list[SingleCutResult],
    context: _LiveWorkbookContext,
    winner_mask: Any,
    laggard_mask: Any,
) -> None:
    worksheet = workbook.create_sheet("_RawData")
    worksheet.sheet_state = "hidden"
    columns = context.columns
    helper_headers = [
        "winners_mask_data",
        "laggards_mask_data",
        "others_mask_data",
    ]
    worksheet.cell(row=1, column=1, value="respondent_id")
    for index, column in enumerate(columns, start=2):
        worksheet.cell(row=1, column=index, value=column.header)
    helper_start = len(columns) + 2
    for offset, header in enumerate(helper_headers):
        worksheet.cell(row=1, column=helper_start + offset, value=header)

    rows = _raw_rows_from_dataframe(decoded_df, schema, columns)
    if not rows:
        rows = _synthetic_raw_rows(schema, results, columns)
    winner_values = _wvl_mask_values(winner_mask, decoded_df, len(rows))
    laggard_values = _wvl_mask_values(laggard_mask, decoded_df, len(rows))
    others_values = [
        not (
            bool(winner_values[index]) if index < len(winner_values) else False
        )
        and not (
            bool(laggard_values[index]) if index < len(laggard_values) else False
        )
        for index in range(len(rows))
    ]

    for row_index, row_payload in enumerate(rows, start=2):
        pos = row_index - 2
        worksheet.cell(row=row_index, column=1, value=row_payload.get("respondent_id", pos + 1))
        for col_index, column in enumerate(columns, start=2):
            worksheet.cell(row=row_index, column=col_index, value=row_payload.get(column.key))
        helper_values = [
            bool(winner_values[pos]) if pos < len(winner_values) else False,
            bool(laggard_values[pos]) if pos < len(laggard_values) else False,
            bool(others_values[pos]) if pos < len(others_values) else False,
        ]
        for offset, value in enumerate(helper_values):
            worksheet.cell(row=row_index, column=helper_start + offset, value=value)

    last_row = max(2, len(rows) + 1)
    _add_named_range(workbook, "respondent_id_data", "_RawData", f"$A$2:$A${last_row}")
    for index, column in enumerate(columns, start=2):
        col_letter = _openpyxl_column_letter(index)
        _add_named_range(workbook, column.data_name, "_RawData", f"${col_letter}$2:${col_letter}${last_row}")
    for offset, header in enumerate(helper_headers):
        col_letter = _openpyxl_column_letter(helper_start + offset)
        _add_named_range(workbook, header, "_RawData", f"${col_letter}$2:${col_letter}${last_row}")
    worksheet.freeze_panes = "A2"


def _wvl_mask_values(mask: Any, decoded_df: Any, expected_len: int) -> list[bool]:
    if hasattr(mask, "reindex") and decoded_df is not None and hasattr(decoded_df, "index"):
        return [bool(value) for value in mask.reindex(decoded_df.index, fill_value=False).tolist()]
    try:
        return [bool(value) for value in list(mask)]
    except TypeError:
        return [False] * expected_len


def _wvl_write_run_summary(
    workbook: Any,
    schema: SurveySchema,
    segment_definition: SegmentDefinition,
    laggard_definition: SegmentDefinition,
    winner_n: int,
    laggard_n: int,
) -> None:
    worksheet = workbook.create_sheet("Run_Summary")
    rows = [
        ("Source datamap:", schema.source_datamap_path),
        ("Source raw data:", schema.source_rawdata_path),
        ("Run timestamp:", schema.parsed_at.isoformat()),
        ("Workbook:", "Outcome Segmented Workbook (Winners vs Laggards)"),
        ("Winner label:", segment_definition.winner_label),
        ("Winner outcome:", segment_definition.outcome_question_id),
        ("Winner definition:", _wvl_segment_description(segment_definition)),
        ("Winner cohort size:", winner_n),
        ("Laggard label:", laggard_definition.laggard_label or laggard_definition.loser_label),
        ("Laggard outcome:", laggard_definition.outcome_question_id),
        ("Laggard definition:", _wvl_segment_description(laggard_definition, laggard=True)),
        ("Laggard cohort size:", laggard_n),
    ]
    for row_index, (label, value) in enumerate(rows, start=1):
        worksheet.cell(row=row_index, column=1, value=label).font = _live_font(bold=True)
        worksheet.cell(row=row_index, column=2, value=value)
    _live_autofit(worksheet)


def _wvl_write_cohort_definition_sheet(
    workbook: Any,
    segment_definition: SegmentDefinition,
    winner_n: int,
    laggard_n: int,
) -> None:
    worksheet = workbook.create_sheet("_CohortDefinition")
    worksheet.sheet_state = "hidden"
    worksheet.cell(row=1, column=1, value="Cohort Source").font = _live_font(bold=True)
    outcome_value = (
        "(N/A for manual)"
        if segment_definition.segment_mode == "manual_uuid"
        else segment_definition.outcome_question_id
    )
    rows = [
        ("Mode", segment_definition.segment_mode),
        ("Outcome question", outcome_value),
        ("Winner count", winner_n),
        ("Laggard count", laggard_n),
    ]
    if segment_definition.segment_mode == "manual_uuid":
        overlap_count = len(
            set(segment_definition.manual_winner_uuids)
            & set(segment_definition.manual_laggard_uuids)
        )
        rows.extend(
            [
                ("Invalid uuid count", 0),
                ("Overlap count", overlap_count),
            ]
        )

    for row_index, (label, value) in enumerate(rows, start=2):
        worksheet.cell(row=row_index, column=1, value=label).font = _live_font(bold=True)
        worksheet.cell(row=row_index, column=2, value=value)

    if segment_definition.segment_mode == "manual_uuid":
        row_index = 10
        worksheet.cell(row=row_index, column=1, value="Winner uuids:").font = _live_font(bold=True)
        for uuid in segment_definition.manual_winner_uuids:
            row_index += 1
            worksheet.cell(row=row_index, column=1, value=uuid)
        row_index += 2
        worksheet.cell(row=row_index, column=1, value="Laggard uuids:").font = _live_font(bold=True)
        for uuid in segment_definition.manual_laggard_uuids:
            row_index += 1
            worksheet.cell(row=row_index, column=1, value=uuid)
    _wvl_apply_arial_to_used_cells(worksheet)
    _live_autofit(worksheet)


def _wvl_segment_description(segment_definition: SegmentDefinition, laggard: bool = False) -> str:
    if segment_definition.segment_mode == "manual_uuid":
        count = (
            len(segment_definition.manual_laggard_uuids)
            if laggard
            else len(segment_definition.manual_winner_uuids)
        )
        label = "laggard" if laggard else "winner"
        return f"manual_uuid {label} list ({count} respondent ids)"
    if segment_definition.segment_mode == "categorical":
        values = (
            segment_definition.laggard_values
            if laggard and segment_definition.laggard_values
            else segment_definition.winner_values
        )
        return f"{segment_definition.outcome_question_id} IN {tuple(values)}"
    if segment_definition.segment_mode == "numeric_threshold":
        threshold = (
            segment_definition.laggard_threshold
            if laggard and segment_definition.laggard_threshold is not None
            else segment_definition.winner_threshold
        )
        direction = (
            segment_definition.laggard_threshold_direction
            if laggard and segment_definition.laggard_threshold is not None
            else segment_definition.threshold_direction
        )
        return f"{segment_definition.outcome_question_id} {direction} {threshold}"
    return f"{segment_definition.outcome_question_id} quartile winner={segment_definition.quartile_winner}"


def _wvl_write_category_sheet_header(
    worksheet: Any,
    theme_name: str,
    start_row: int = 1,
) -> int:
    from openpyxl.styles import Alignment

    banner = worksheet.cell(row=start_row, column=1, value=theme_name)
    banner.font = _live_font(bold=True, size=12, color="FFFFFF")
    banner.fill = _live_fill("CC0000")
    banner.alignment = Alignment(vertical="center")
    worksheet.row_dimensions[start_row].height = 25
    for col_index in range(1, 13):
        cell = worksheet.cell(row=start_row, column=col_index)
        cell.fill = _live_fill("CC0000")
        cell.font = _live_font(bold=True, size=12, color="FFFFFF")
    worksheet.cell(
        row=start_row + 1,
        column=1,
        value="Tip: Use the Filters sheet for workbook-wide values, or override below.",
    ).font = _live_font(italic=True, size=9, color="666666")
    return start_row + 2


def _wvl_write_question_block(
    workbook: Any,
    worksheet: Any,
    start_row: int,
    result: SingleCutResult,
    question: Any,
    context: _LiveWorkbookContext,
    log: CalculationLog,
    winner_label: str,
    laggard_label: str,
    schema: SurveySchema,
    sheet_filters: list[dict[str, str]],
    theme_prefix: str,
    decoded_df: Any,
    winner_mask: Any,
    laggard_mask: Any,
    others_mask: Any,
) -> int:
    del sheet_filters
    heading_fill = _live_fill("FFCC0000")
    heading_font = _live_font(bold=True, size=11, color="FFFFFFFF")
    for col_index in range(1, 13):
        cell = worksheet.cell(row=start_row, column=col_index)
        cell.fill = heading_fill
        cell.font = heading_font
    worksheet.cell(
        row=start_row,
        column=1,
        value=_question_heading_text(question, {}),
    ).font = heading_font

    q_prefix = f"{theme_prefix}_{_safe_defined_name(question.canonical_id)}"
    q_filter_prefix = f"{q_prefix}_F"
    fq_name = ""
    if context.per_question_filter_count > 0:
        for slot in range(1, context.per_question_filter_count + 1):
            slot_fq_name, _slot_fv_name = _live_write_per_question_filter_slot(
                workbook,
                worksheet,
                start_row + slot,
                q_filter_prefix,
                slot,
            )
            if slot == 1:
                fq_name = slot_fq_name
        _build_per_question_filter_helper_column(
            workbook,
            q_filter_prefix,
            context,
            context.per_question_filter_count,
        )

    row_index = start_row + max(1, context.per_question_filter_count) + 1
    ct_name = f"{q_prefix}_CT"
    if not _question_is_grid_target(question):
        worksheet.cell(row=row_index, column=1, value="Cross-tab by").font = _live_font(bold=True, size=9)
        ct_cell = worksheet.cell(row=row_index, column=3, value="(None)")
        _add_named_cell(workbook, ct_name, worksheet, ct_cell.coordinate)
        _add_dropdown_to_cell(worksheet, ct_cell.coordinate, "=Cross_Tab_Questions")
        row_index += 1

    note = _subset_denominator_note(result, question, schema) or (
        f"Note: This question was shown to {int(getattr(result, 'valid_n', 0)):,} respondents."
    )
    note_cell = worksheet.cell(row=row_index, column=1, value=note)
    note_cell.font = _live_font(italic=True, size=9, color="808080")

    table_start_row = row_index + 2
    if isinstance(result, SingleSelectResult):
        return _wvl_write_single_select_layout(worksheet, table_start_row, result, question, context, log, winner_label, laggard_label, fq_name, theme_prefix, decoded_df, winner_mask, laggard_mask, others_mask)
    if isinstance(result, MultiSelectResult):
        return _wvl_write_multi_select_layout(worksheet, table_start_row, result, question, context, log, winner_label, laggard_label, fq_name, theme_prefix, decoded_df, winner_mask, laggard_mask, others_mask)
    if isinstance(result, GridRatedResult):
        return _wvl_write_grid_rated_layout(worksheet, table_start_row, result, question, context, log, winner_label, laggard_label, fq_name, theme_prefix)
    if isinstance(result, NumericResult):
        return _wvl_write_numeric_layout(worksheet, table_start_row, result, question, context, log, winner_label, laggard_label, fq_name, theme_prefix)
    return table_start_row + 1


def _wvl_write_grouped_header(
    worksheet: Any,
    row_index: int,
    last_col: int,
    include_percent_groups: bool = True,
) -> None:
    from openpyxl.styles import Alignment

    for col_index in range(1, last_col + 1):
        cell = worksheet.cell(row=row_index, column=col_index)
        cell.font = _live_font(bold=True, color="FFFFFF")
        cell.alignment = Alignment(horizontal="center", vertical="center")

    if include_percent_groups:
        groups = (
            (3, 4, "Winners", "CC0000"),
            (5, 6, "Laggards", "5C5C5C"),
            (7, 8, "Others", "858585"),
            (9, 10, "Total", "333333"),
        )
        for start_col, end_col, label, color in groups:
            worksheet.merge_cells(
                start_row=row_index,
                start_column=start_col,
                end_row=row_index,
                end_column=end_col,
            )
            cell = worksheet.cell(row=row_index, column=start_col, value=label)
            cell.fill = _live_fill(color)
            cell.font = _live_font(bold=True, color="FFFFFF")
            cell.alignment = Alignment(horizontal="center", vertical="center")
            for col_index in range(start_col, end_col + 1):
                worksheet.cell(row=row_index, column=col_index).fill = _live_fill(color)
        who_cell = worksheet.cell(row=row_index, column=11, value="Who prioritizes more?")
        who_cell.fill = _live_fill("5C5C5C")
        who_cell.font = _live_font(bold=True, color="FFFFFF")
        who_cell.alignment = Alignment(horizontal="center", vertical="center")
    else:
        labels = {
            3: ("Winners", "CC0000"),
            4: ("Laggards", "5C5C5C"),
            5: ("Others", "858585"),
            6: ("Total", "333333"),
            7: ("Delta", "5C5C5C"),
            8: ("Winners", "CC0000"),
            9: ("Laggards", "5C5C5C"),
            10: ("Others", "858585"),
            11: ("Total", "333333"),
            12: ("Who prioritizes more?", "5C5C5C"),
        }
        for col_index, (label, color) in labels.items():
            cell = worksheet.cell(row=row_index, column=col_index, value=label)
            cell.fill = _live_fill(color)
            cell.font = _live_font(bold=True, color="FFFFFF")
            cell.alignment = Alignment(horizontal="center", vertical="center")


def _wvl_priority_formula(
    row_index: int,
    winner_col: str,
    laggard_col: str,
    others_col: str,
    margin: str,
) -> str:
    return (
        f'=IF(AND({winner_col}{row_index}-{others_col}{row_index}>={margin},'
        f'{winner_col}{row_index}-{laggard_col}{row_index}>={margin}),"Winners",'
        f'IF(AND({laggard_col}{row_index}-{others_col}{row_index}>={margin},'
        f'{laggard_col}{row_index}-{winner_col}{row_index}>={margin}),"Laggards",'
        f'IF(AND({others_col}{row_index}-{winner_col}{row_index}>={margin},'
        f'{others_col}{row_index}-{laggard_col}{row_index}>={margin}),"Others",'
        '"Tied")))'
    )


def _wvl_apply_priority_conditional_formatting(
    worksheet: Any,
    start_row: int,
    end_row: int,
    col_index: int,
) -> None:
    if end_row < start_row:
        return
    from openpyxl.formatting.rule import CellIsRule

    column_letter = _openpyxl_column_letter(col_index)
    cell_range = f"{column_letter}{start_row}:{column_letter}{end_row}"
    rules = (
        ("Winners", "83AC9A", "FFFFFF"),
        ("Laggards", "CC0000", "FFFFFF"),
        ("Others", "858585", "FFFFFF"),
        ("Tied", "D6D6D6", "000000"),
    )
    for label, fill_color, font_color in rules:
        worksheet.conditional_formatting.add(
            cell_range,
            CellIsRule(
                operator="equal",
                formula=[f'"{label}"'],
                fill=_live_fill(fill_color),
                font=_live_font(bold=True, color=font_color),
            ),
        )


def _wvl_apply_percent_heatmaps(
    worksheet: Any,
    data_start: int,
    data_end: int,
    columns: tuple[int, ...],
) -> None:
    for col_index in columns:
        _apply_color_scale_range(worksheet, data_start, col_index, data_end, col_index)


def _wvl_apply_mean_heatmaps(
    worksheet: Any,
    data_start: int,
    data_end: int,
    columns: tuple[int, ...],
) -> None:
    if data_end < data_start:
        return
    from openpyxl.formatting.rule import ColorScaleRule

    for col_index in columns:
        cell_range = (
            f"{_openpyxl_column_letter(col_index)}{data_start}:"
            f"{_openpyxl_column_letter(col_index)}{data_end}"
        )
        worksheet.conditional_formatting.add(
            cell_range,
            ColorScaleRule(
                start_type="num",
                start_value=0,
                start_color="F8696B",
                mid_type="num",
                mid_value=5,
                mid_color="FFEB84",
                end_type="num",
                end_value=10,
                end_color="63BE7B",
            ),
        )


def _wvl_compute_single_select_cohort_results(
    question: Any,
    decoded_df: Any,
    winner_mask: Any,
    laggard_mask: Any,
    others_mask: Any,
) -> tuple[SingleSelectResult, SingleSelectResult, SingleSelectResult]:
    return (
        compute_single_select(
            question,
            decoded_df,
            CalculationLog(),
            filter_mask=winner_mask,
            filter_expr="Winners cohort",
        ),
        compute_single_select(
            question,
            decoded_df,
            CalculationLog(),
            filter_mask=laggard_mask,
            filter_expr="Laggards cohort",
        ),
        compute_single_select(
            question,
            decoded_df,
            CalculationLog(),
            filter_mask=others_mask,
            filter_expr="Others cohort",
        ),
    )


def _wvl_compute_multi_select_cohort_results(
    question: Any,
    decoded_df: Any,
    winner_mask: Any,
    laggard_mask: Any,
    others_mask: Any,
) -> tuple[MultiSelectResult, MultiSelectResult, MultiSelectResult]:
    return (
        compute_multi_select(
            question,
            decoded_df,
            CalculationLog(),
            filter_mask=winner_mask,
            filter_expr="Winners cohort",
        ),
        compute_multi_select(
            question,
            decoded_df,
            CalculationLog(),
            filter_mask=laggard_mask,
            filter_expr="Laggards cohort",
        ),
        compute_multi_select(
            question,
            decoded_df,
            CalculationLog(),
            filter_mask=others_mask,
            filter_expr="Others cohort",
        ),
    )


def _wvl_single_select_payload(
    result: SingleSelectResult,
    option_code: Any,
    label: Any,
) -> dict[str, Any]:
    if option_code in result.distribution:
        return result.distribution[option_code]
    option_text = str(option_code)
    label_text = str(label)
    for candidate_code, payload in result.distribution.items():
        if str(candidate_code) == option_text or str(payload.get("label", "")) == label_text:
            return payload
    return {}


def _wvl_single_select_count(
    result: SingleSelectResult,
    option_code: Any,
    label: Any,
) -> int:
    return int(_wvl_single_select_payload(result, option_code, label).get("count", 0) or 0)


def _wvl_multi_select_count(result: MultiSelectResult, option_id: str) -> int:
    payload = result.selections.get(option_id, {})
    return int(payload.get("count", 0) or 0)


def _wvl_pct(count: int, denominator: int) -> float:
    return float(count / denominator) if denominator else 0.0


def _wvl_write_single_select_layout(
    worksheet: Any,
    start_row: int,
    result: SingleSelectResult,
    question: Any,
    context: _LiveWorkbookContext,
    log: CalculationLog,
    winner_label: str,
    laggard_label: str,
    fq_name: str,
    theme_prefix: str,
    decoded_df: Any,
    winner_mask: Any,
    laggard_mask: Any,
    others_mask: Any,
) -> int:
    _wvl_write_grouped_header(worksheet, start_row, 11)
    header_row = start_row + 1
    headers = [
        "Option ID",
        "Option",
        "Count",
        "%",
        "Count",
        "%",
        "Count",
        "%",
        "Count",
        "%",
        "Who prioritizes more?",
    ]
    _live_header_row(worksheet, header_row, headers)
    column = context.column_by_key.get(question.canonical_id)
    data_name = column.data_name if column is not None else f"{question.canonical_id}_data"
    winner_result, laggard_result, others_result = _wvl_compute_single_select_cohort_results(
        question,
        decoded_df,
        winner_mask,
        laggard_mask,
        others_mask,
    )
    winner_denominator = int(winner_result.valid_n)
    laggard_denominator = int(laggard_result.valid_n)
    others_denominator = int(others_result.valid_n)
    total_denominator = int(result.valid_n)
    winner_total_row = header_row + len(result.distribution) + 2
    laggard_total_row = winner_total_row + 1
    others_total_row = laggard_total_row + 1
    total_respondents_row = others_total_row + 1
    row_index = header_row + 1
    for option_code, payload in result.distribution.items():
        label = payload.get("label", str(option_code))
        worksheet.cell(row=row_index, column=1, value=option_code)
        worksheet.cell(row=row_index, column=2, value=label)
        winner_count_formula = _wvl_countifs(data_name, option_code, fq_name, "winners_mask_data", theme_prefix)
        laggard_count_formula = _wvl_countifs(data_name, option_code, fq_name, "laggards_mask_data", theme_prefix)
        others_count_formula = _wvl_countifs(data_name, option_code, fq_name, "others_mask_data", theme_prefix)
        total_count_formula = _wvl_total_countifs(data_name, option_code, fq_name, theme_prefix)
        winner_count = _wvl_single_select_count(winner_result, option_code, label)
        laggard_count = _wvl_single_select_count(laggard_result, option_code, label)
        others_count = _wvl_single_select_count(others_result, option_code, label)
        total_count = _wvl_single_select_count(result, option_code, label)
        winner_count_cell = _live_formula(worksheet, row_index, 3, winner_count_formula, winner_count)
        winner_count_cell.number_format = "#,##0"
        _record_wvl_metric(log, worksheet.title, result.question_id, (data_name,), "winner_count", winner_count_formula, winner_count)
        winner_pct_formula = f'=IFERROR(C{row_index}/$C${winner_total_row},0)'
        winner_pct_cell = _live_formula(worksheet, row_index, 4, winner_pct_formula, _wvl_pct(winner_count, winner_denominator))
        winner_pct_cell.number_format = "0.0%"
        _record_wvl_metric(log, worksheet.title, result.question_id, (data_name,), "winner_pct", winner_pct_formula, _wvl_pct(winner_count, winner_denominator))
        laggard_count_cell = _live_formula(worksheet, row_index, 5, laggard_count_formula, laggard_count)
        laggard_count_cell.number_format = "#,##0"
        _record_wvl_metric(log, worksheet.title, result.question_id, (data_name,), "laggard_count", laggard_count_formula, laggard_count)
        laggard_pct_formula = f'=IFERROR(E{row_index}/$C${laggard_total_row},0)'
        laggard_pct_cell = _live_formula(worksheet, row_index, 6, laggard_pct_formula, _wvl_pct(laggard_count, laggard_denominator))
        laggard_pct_cell.number_format = "0.0%"
        _record_wvl_metric(log, worksheet.title, result.question_id, (data_name,), "laggard_pct", laggard_pct_formula, _wvl_pct(laggard_count, laggard_denominator))
        others_count_cell = _live_formula(worksheet, row_index, 7, others_count_formula, others_count)
        others_count_cell.number_format = "#,##0"
        _record_wvl_metric(log, worksheet.title, result.question_id, (data_name,), "others_count", others_count_formula, others_count)
        others_pct_formula = f'=IFERROR(G{row_index}/$C${others_total_row},0)'
        others_pct_cell = _live_formula(worksheet, row_index, 8, others_pct_formula, _wvl_pct(others_count, others_denominator))
        others_pct_cell.number_format = "0.0%"
        _record_wvl_metric(log, worksheet.title, result.question_id, (data_name,), "others_pct", others_pct_formula, _wvl_pct(others_count, others_denominator))
        total_count_cell = _live_formula(worksheet, row_index, 9, total_count_formula, total_count)
        total_count_cell.number_format = "#,##0"
        _record_wvl_metric(log, worksheet.title, result.question_id, (data_name,), "total_count", total_count_formula, total_count)
        total_pct_formula = f'=IFERROR(I{row_index}/$C${total_respondents_row},0)'
        total_pct_cell = _live_formula(worksheet, row_index, 10, total_pct_formula, _wvl_pct(total_count, total_denominator))
        total_pct_cell.number_format = "0.0%"
        _record_wvl_metric(log, worksheet.title, result.question_id, (data_name,), "total_pct", total_pct_formula, _wvl_pct(total_count, total_denominator))
        priority_formula = _wvl_priority_formula(row_index, "D", "F", "H", "0.02")
        priority_cell = _live_formula(worksheet, row_index, 11, priority_formula, "Tied")
        priority_cell.number_format = "@"
        _record_wvl_metric(log, worksheet.title, result.question_id, (data_name,), "who_prioritizes_more", priority_formula, 0)
        row_index += 1
    data_start = header_row + 1
    data_end = row_index - 1
    if data_end >= data_start:
        _wvl_apply_percent_heatmaps(worksheet, data_start, data_end, (4, 6, 8, 10))
        _wvl_apply_priority_conditional_formatting(worksheet, data_start, data_end, 11)
    _wvl_write_totals(
        worksheet,
        winner_total_row,
        laggard_total_row,
        others_total_row,
        total_respondents_row,
        data_name,
        fq_name,
        winner_denominator,
        laggard_denominator,
        others_denominator,
        total_denominator,
        theme_prefix,
    )
    return total_respondents_row + 2


def _wvl_write_multi_select_layout(
    worksheet: Any,
    start_row: int,
    result: MultiSelectResult,
    question: Any,
    context: _LiveWorkbookContext,
    log: CalculationLog,
    winner_label: str,
    laggard_label: str,
    fq_name: str,
    theme_prefix: str,
    decoded_df: Any,
    winner_mask: Any,
    laggard_mask: Any,
    others_mask: Any,
) -> int:
    _wvl_write_grouped_header(worksheet, start_row, 11)
    header_row = start_row + 1
    headers = [
        "Option ID",
        "Option",
        "Count",
        "%",
        "Count",
        "%",
        "Count",
        "%",
        "Count",
        "%",
        "Who prioritizes more?",
    ]
    _live_header_row(worksheet, header_row, headers)
    winner_total_row = header_row + len(result.selections) + 2
    laggard_total_row = winner_total_row + 1
    others_total_row = laggard_total_row + 1
    total_respondents_row = others_total_row + 1
    row_index = header_row + 1
    anchor_data_name = ""
    denominator_data_names: list[str] = []
    winner_result, laggard_result, others_result = _wvl_compute_multi_select_cohort_results(
        question,
        decoded_df,
        winner_mask,
        laggard_mask,
        others_mask,
    )
    winner_denominator = int(winner_result.respondents_who_answered_any)
    laggard_denominator = int(laggard_result.respondents_who_answered_any)
    others_denominator = int(others_result.respondents_who_answered_any)
    total_denominator = int(result.respondents_who_answered_any)
    for option_id, payload in result.selections.items():
        column = context.column_by_key.get(option_id)
        if column is None:
            continue
        anchor_data_name = anchor_data_name or column.data_name
        denominator_data_names.append(column.data_name)
        worksheet.cell(row=row_index, column=1, value=option_id)
        worksheet.cell(row=row_index, column=2, value=payload.get("label", option_id))
        winner_count = _wvl_multi_select_count(winner_result, option_id)
        laggard_count = _wvl_multi_select_count(laggard_result, option_id)
        others_count = _wvl_multi_select_count(others_result, option_id)
        total_count = int(payload.get("count", 0) or 0)
        winner_formula = _wvl_countifs(column.data_name, "Selected", fq_name, "winners_mask_data", theme_prefix)
        laggard_formula = _wvl_countifs(column.data_name, "Selected", fq_name, "laggards_mask_data", theme_prefix)
        others_formula = _wvl_countifs(column.data_name, "Selected", fq_name, "others_mask_data", theme_prefix)
        total_formula = _wvl_total_countifs(column.data_name, "Selected", fq_name, theme_prefix)
        winner_count_cell = _live_formula(worksheet, row_index, 3, winner_formula, winner_count)
        winner_count_cell.number_format = "#,##0"
        _record_wvl_metric(log, worksheet.title, result.question_id, (column.data_name,), "winner_count", winner_formula, winner_count)
        winner_pct_formula = f'=IFERROR(C{row_index}/$C${winner_total_row},0)'
        winner_pct_cell = _live_formula(worksheet, row_index, 4, winner_pct_formula, _wvl_pct(winner_count, winner_denominator))
        winner_pct_cell.number_format = "0.0%"
        _record_wvl_metric(log, worksheet.title, result.question_id, (column.data_name,), "winner_pct", winner_pct_formula, _wvl_pct(winner_count, winner_denominator))
        laggard_count_cell = _live_formula(worksheet, row_index, 5, laggard_formula, laggard_count)
        laggard_count_cell.number_format = "#,##0"
        _record_wvl_metric(log, worksheet.title, result.question_id, (column.data_name,), "laggard_count", laggard_formula, laggard_count)
        laggard_pct_formula = f'=IFERROR(E{row_index}/$C${laggard_total_row},0)'
        laggard_pct_cell = _live_formula(worksheet, row_index, 6, laggard_pct_formula, _wvl_pct(laggard_count, laggard_denominator))
        laggard_pct_cell.number_format = "0.0%"
        _record_wvl_metric(log, worksheet.title, result.question_id, (column.data_name,), "laggard_pct", laggard_pct_formula, _wvl_pct(laggard_count, laggard_denominator))
        others_count_cell = _live_formula(worksheet, row_index, 7, others_formula, others_count)
        others_count_cell.number_format = "#,##0"
        _record_wvl_metric(log, worksheet.title, result.question_id, (column.data_name,), "others_count", others_formula, others_count)
        others_pct_formula = f'=IFERROR(G{row_index}/$C${others_total_row},0)'
        others_pct_cell = _live_formula(worksheet, row_index, 8, others_pct_formula, _wvl_pct(others_count, others_denominator))
        others_pct_cell.number_format = "0.0%"
        _record_wvl_metric(log, worksheet.title, result.question_id, (column.data_name,), "others_pct", others_pct_formula, _wvl_pct(others_count, others_denominator))
        total_count_cell = _live_formula(worksheet, row_index, 9, total_formula, total_count)
        total_count_cell.number_format = "#,##0"
        _record_wvl_metric(log, worksheet.title, result.question_id, (column.data_name,), "total_count", total_formula, total_count)
        total_pct_formula = f'=IFERROR(I{row_index}/$C${total_respondents_row},0)'
        total_pct_cell = _live_formula(worksheet, row_index, 10, total_pct_formula, _wvl_pct(total_count, total_denominator))
        total_pct_cell.number_format = "0.0%"
        _record_wvl_metric(log, worksheet.title, result.question_id, (column.data_name,), "total_pct", total_pct_formula, _wvl_pct(total_count, total_denominator))
        priority_cell = _live_formula(worksheet, row_index, 11, _wvl_priority_formula(row_index, "D", "F", "H", "0.02"), "Tied")
        priority_cell.number_format = "@"
        row_index += 1
    if anchor_data_name:
        data_start = header_row + 1
        data_end = row_index - 1
        if data_end >= data_start:
            _wvl_apply_percent_heatmaps(worksheet, data_start, data_end, (4, 6, 8, 10))
            _wvl_apply_priority_conditional_formatting(worksheet, data_start, data_end, 11)
        _wvl_write_totals(
            worksheet,
            winner_total_row,
            laggard_total_row,
            others_total_row,
            total_respondents_row,
            anchor_data_name,
            fq_name,
            winner_denominator,
            laggard_denominator,
            others_denominator,
            total_denominator,
            theme_prefix,
            denominator_data_names=denominator_data_names,
        )
    return total_respondents_row + 2


def _wvl_write_grid_rated_layout(
    worksheet: Any,
    start_row: int,
    result: GridRatedResult,
    question: Any,
    context: _LiveWorkbookContext,
    log: CalculationLog,
    winner_label: str,
    laggard_label: str,
    fq_name: str,
    theme_prefix: str,
) -> int:
    _wvl_write_grouped_header(worksheet, start_row, 12, include_percent_groups=False)
    header_row = start_row + 1
    headers = [
        "Sub-question ID",
        "Sub-question",
        f"{winner_label} mean",
        f"{laggard_label} mean",
        "Others mean",
        "Total mean",
        "Delta (W-L)",
        f"{winner_label} n",
        f"{laggard_label} n",
        "Others n",
        "Total n",
        "Who prioritizes more?",
    ]
    _live_header_row(worksheet, header_row, headers)
    row_index = header_row + 1
    for result_row in result.rows:
        source_column = _wvl_grid_rated_source_column(question, result_row.row_id, 0)
        column = context.column_by_key.get(source_column)
        if column is None:
            raise ValueError(f"missing named range for grid-rated source column {source_column!r}")
        worksheet.cell(row=row_index, column=1, value=result_row.row_id)
        worksheet.cell(row=row_index, column=2, value=result_row.row_label)
        winner_mean_formula = _wvl_averageifs(column.data_name, fq_name, "winners_mask_data", theme_prefix)
        laggard_mean_formula = _wvl_averageifs(column.data_name, fq_name, "laggards_mask_data", theme_prefix)
        others_mean_formula = _wvl_averageifs(column.data_name, fq_name, "others_mask_data", theme_prefix)
        total_mean_formula = _wvl_total_averageifs(column.data_name, fq_name, theme_prefix)
        winner_mean_cell = _live_formula(worksheet, row_index, 3, winner_mean_formula, result_row.means_per_column[0] if result_row.means_per_column else 0)
        winner_mean_cell.number_format = "0.00"
        _record_wvl_metric(log, worksheet.title, result.question_id, (column.data_name,), "winner_mean", winner_mean_formula, result_row.means_per_column[0] if result_row.means_per_column else 0)
        laggard_mean_cell = _live_formula(worksheet, row_index, 4, laggard_mean_formula, result_row.means_per_column[1] if len(result_row.means_per_column) > 1 else 0)
        laggard_mean_cell.number_format = "0.00"
        _record_wvl_metric(log, worksheet.title, result.question_id, (column.data_name,), "laggard_mean", laggard_mean_formula, result_row.means_per_column[1] if len(result_row.means_per_column) > 1 else 0)
        others_mean_cell = _live_formula(worksheet, row_index, 5, others_mean_formula, 0)
        others_mean_cell.number_format = "0.00"
        _record_wvl_metric(log, worksheet.title, result.question_id, (column.data_name,), "others_mean", others_mean_formula, 0)
        total_mean_cell = _live_formula(worksheet, row_index, 6, total_mean_formula, 0)
        total_mean_cell.number_format = "0.00"
        _record_wvl_metric(log, worksheet.title, result.question_id, (column.data_name,), "total_mean", total_mean_formula, 0)
        delta_formula = f"=C{row_index}-D{row_index}"
        delta_cell = _live_formula(worksheet, row_index, 7, delta_formula, 0)
        delta_cell.number_format = "+0.00;-0.00"
        _record_wvl_metric(log, worksheet.title, result.question_id, (column.data_name,), "delta_mean", delta_formula, 0)
        winner_n_cell = _live_formula(worksheet, row_index, 8, _wvl_countifs(column.data_name, "<>", fq_name, "winners_mask_data", theme_prefix), result_row.valid_n_per_column[0] if result_row.valid_n_per_column else 0)
        winner_n_cell.number_format = "#,##0"
        laggard_n_cell = _live_formula(worksheet, row_index, 9, _wvl_countifs(column.data_name, "<>", fq_name, "laggards_mask_data", theme_prefix), result_row.valid_n_per_column[1] if len(result_row.valid_n_per_column) > 1 else 0)
        laggard_n_cell.number_format = "#,##0"
        others_n_cell = _live_formula(worksheet, row_index, 10, _wvl_countifs(column.data_name, "<>", fq_name, "others_mask_data", theme_prefix), 0)
        others_n_cell.number_format = "#,##0"
        total_n_cell = _live_formula(worksheet, row_index, 11, _wvl_total_countifs(column.data_name, "<>", fq_name, theme_prefix), 0)
        total_n_cell.number_format = "#,##0"
        priority_cell = _live_formula(worksheet, row_index, 12, _wvl_priority_formula(row_index, "C", "D", "E", "0.3"), "Tied")
        priority_cell.number_format = "@"
        row_index += 1
    data_start = header_row + 1
    data_end = row_index - 1
    if data_end >= data_start:
        _wvl_apply_mean_heatmaps(worksheet, data_start, data_end, (3, 4, 5, 6))
        _wvl_apply_priority_conditional_formatting(worksheet, data_start, data_end, 12)
    return row_index


def _wvl_write_numeric_layout(
    worksheet: Any,
    start_row: int,
    result: NumericResult,
    question: Any,
    context: _LiveWorkbookContext,
    log: CalculationLog,
    winner_label: str,
    laggard_label: str,
    fq_name: str,
    theme_prefix: str,
) -> int:
    del log
    headers = ["Stat", winner_label, laggard_label, "Delta"]
    _live_header_row(worksheet, start_row, headers)
    column = context.column_by_key.get(question.canonical_id)
    if column is None and question.raw_columns:
        column = context.column_by_key.get(question.raw_columns[0])
    if column is None:
        return start_row + 1
    metrics = ["Mean", "Median", "Min", "Max"]
    for offset, metric in enumerate(metrics, start=1):
        row_index = start_row + offset
        worksheet.cell(row=row_index, column=1, value=metric)
        winner_formula = _wvl_averageifs(column.data_name, fq_name, "winners_mask_data", theme_prefix) if metric == "Mean" else _wvl_filtered_numeric_formula(metric, column.data_name, fq_name, "winners_mask_data", theme_prefix)
        laggard_formula = _wvl_averageifs(column.data_name, fq_name, "laggards_mask_data", theme_prefix) if metric == "Mean" else _wvl_filtered_numeric_formula(metric, column.data_name, fq_name, "laggards_mask_data", theme_prefix)
        value_format = "#,##0" if metric.lower().endswith("count") else "0.00"
        delta_format = "+#,##0;-#,##0" if value_format == "#,##0" else "+0.00;-0.00"
        winner_cell = _live_formula(worksheet, row_index, 2, winner_formula, _numeric_result_cache_value(metric, result))
        winner_cell.number_format = value_format
        laggard_cell = _live_formula(worksheet, row_index, 3, laggard_formula, _numeric_result_cache_value(metric, result))
        laggard_cell.number_format = value_format
        delta_cell = _live_formula(worksheet, row_index, 4, f"=B{row_index}-C{row_index}", 0)
        delta_cell.number_format = delta_format
    return start_row + len(metrics) + 1


def _wvl_write_totals(
    worksheet: Any,
    winner_total_row: int,
    laggard_total_row: int,
    others_total_row: int,
    total_respondents_row: int,
    data_name: str,
    fq_name: str,
    winner_cache_value: int,
    laggard_cache_value: int,
    others_cache_value: int,
    total_cache_value: int,
    theme_prefix: str,
    denominator_data_names: list[str] | None = None,
) -> None:
    if denominator_data_names:
        winner_formula = "=" + _build_multi_select_respondent_count_formula(
            denominator_data_names,
            [],
            fq_name,
            "",
            theme_prefix,
            extra_pairs=[("winners_mask_data", "TRUE")],
        )
        laggard_formula = "=" + _build_multi_select_respondent_count_formula(
            denominator_data_names,
            [],
            fq_name,
            "",
            theme_prefix,
            extra_pairs=[("laggards_mask_data", "TRUE")],
        )
        others_formula = "=" + _build_multi_select_respondent_count_formula(
            denominator_data_names,
            [],
            fq_name,
            "",
            theme_prefix,
            extra_pairs=[("others_mask_data", "TRUE")],
        )
        total_formula = "=" + _build_multi_select_respondent_count_formula(
            denominator_data_names,
            [],
            fq_name,
            "",
            theme_prefix,
        )
    else:
        winner_formula = _wvl_countifs(
            data_name,
            "<>",
            fq_name,
            "winners_mask_data",
            theme_prefix,
        )
        laggard_formula = _wvl_countifs(
            data_name,
            "<>",
            fq_name,
            "laggards_mask_data",
            theme_prefix,
        )
        others_formula = _wvl_countifs(
            data_name,
            "<>",
            fq_name,
            "others_mask_data",
            theme_prefix,
        )
        total_formula = _wvl_total_countifs(data_name, "<>", fq_name, theme_prefix)

    worksheet.cell(row=winner_total_row, column=1, value="Winners total")
    winner_total_cell = _live_formula(worksheet, winner_total_row, 3, winner_formula, winner_cache_value)
    winner_total_cell.number_format = "#,##0"
    worksheet.cell(row=laggard_total_row, column=1, value="Laggards total")
    laggard_total_cell = _live_formula(worksheet, laggard_total_row, 3, laggard_formula, laggard_cache_value)
    laggard_total_cell.number_format = "#,##0"
    worksheet.cell(row=others_total_row, column=1, value="Others total")
    others_total_cell = _live_formula(worksheet, others_total_row, 3, others_formula, others_cache_value)
    others_total_cell.number_format = "#,##0"
    worksheet.cell(row=total_respondents_row, column=1, value="Total respondents")
    total_cell = _live_formula(worksheet, total_respondents_row, 3, total_formula, total_cache_value)
    total_cell.number_format = "#,##0"


def _wvl_apply_arial_to_used_cells(worksheet: Any) -> None:
    for row in worksheet.iter_rows():
        for cell in row:
            if cell.value is not None:
                current = cell.font or _live_font()
                cell.font = _live_font(
                    bold=bool(current.bold),
                    italic=bool(current.italic),
                    size=int(current.sz or 10),
                    color=current.color.rgb if getattr(current.color, "type", None) == "rgb" else None,
                )


@dataclass(frozen=True)
class _InputSheetSource:
    path: Path | None
    sheet_name: str | None = None
    kind: str = "unknown"


def _embed_input_file_sheets_openpyxl(
    workbook: Any,
    schema: SurveySchema,
    *,
    embed_input_files: bool = False,
    input_file_sources: dict[str, Any] | None = None,
) -> None:
    """Append visible verbatim input sheets to an openpyxl workbook."""

    if not embed_input_files:
        return
    raw_source, datamap_source = _resolve_input_sheet_sources(schema, input_file_sources)
    _write_input_source_to_openpyxl(
        workbook,
        INPUT_RAW_SHEET_NAME,
        raw_source,
        missing_note="Raw data input file was not available for embedding.",
    )
    _write_input_source_to_openpyxl(
        workbook,
        INPUT_DATAMAP_SHEET_NAME,
        datamap_source,
        missing_note="Data map input file was not available for embedding.",
    )
    _move_input_and_helper_sheets_to_end(workbook)


def _embed_input_file_sheets_xlsxwriter(
    workbook: Any,
    schema: SurveySchema,
    formats: dict[str, Any],
    *,
    embed_input_files: bool = False,
    input_file_sources: dict[str, Any] | None = None,
) -> None:
    """Append visible input sheets to an xlsxwriter/fallback workbook."""

    if not embed_input_files:
        return
    raw_source, datamap_source = _resolve_input_sheet_sources(schema, input_file_sources)
    _write_input_source_to_xlsxwriter(
        workbook,
        INPUT_RAW_SHEET_NAME,
        raw_source,
        formats,
        missing_note="Raw data input file was not available for embedding.",
    )
    _write_input_source_to_xlsxwriter(
        workbook,
        INPUT_DATAMAP_SHEET_NAME,
        datamap_source,
        formats,
        missing_note="Data map input file was not available for embedding.",
    )


def _resolve_input_sheet_sources(
    schema: SurveySchema,
    input_file_sources: dict[str, Any] | None,
) -> tuple[_InputSheetSource, _InputSheetSource]:
    sources = input_file_sources or {}
    raw_path = _existing_input_path(
        sources.get("raw_path")
        or sources.get("raw")
        or sources.get("raw_data_path")
        or schema.source_rawdata_path
    )
    datamap_path = _existing_input_path(
        sources.get("datamap_path")
        or sources.get("datamap")
        or sources.get("data_map_path")
        or schema.source_datamap_path
    )
    raw_sheet = sources.get("raw_sheet") or _sheet_name_from_source(schema.source_rawdata_path)
    datamap_sheet = sources.get("datamap_sheet") or _sheet_name_from_source(schema.source_datamap_path)
    return (
        _InputSheetSource(raw_path, str(raw_sheet) if raw_sheet else None, _source_kind(raw_path)),
        _InputSheetSource(datamap_path, str(datamap_sheet) if datamap_sheet else None, _source_kind(datamap_path)),
    )


def _existing_input_path(value: Any) -> Path | None:
    if not value:
        return None
    try:
        path = Path(str(value))
    except TypeError:
        return None
    return path if path.exists() else None


def _sheet_name_from_source(value: Any) -> str | None:
    if isinstance(value, str) and value.startswith("sheet:"):
        return value.removeprefix("sheet:")
    return None


def _source_kind(path: Path | None) -> str:
    if path is None:
        return "missing"
    suffix = path.suffix.lower()
    if suffix in {".xlsx", ".xlsm"}:
        return "excel"
    if suffix == ".csv":
        return "csv"
    if suffix == ".docx":
        return "docx"
    return "unknown"


def _write_input_source_to_openpyxl(
    workbook: Any,
    sheet_name: str,
    source: _InputSheetSource,
    *,
    missing_note: str,
) -> None:
    if sheet_name in workbook.sheetnames:
        del workbook[sheet_name]
    worksheet = workbook.create_sheet(sheet_name)
    if source.path is None:
        worksheet.cell(row=1, column=1, value=missing_note)
    elif source.kind == "excel":
        _copy_excel_input_sheet_openpyxl(source.path, source.sheet_name, workbook, worksheet)
    elif source.kind == "csv":
        _copy_csv_input_sheet_openpyxl(source.path, worksheet)
    elif source.kind == "docx":
        _copy_docx_input_sheet_openpyxl(source.path, worksheet)
    else:
        worksheet.cell(row=1, column=1, value=f"Unsupported input file type: {source.path.name}")
    if source.kind != "excel":
        _live_autofit(worksheet)


def _copy_excel_input_sheet_openpyxl(
    path: Path,
    sheet_name: str | None,
    target_workbook: Any,
    target_worksheet: Any,
) -> None:
    from copy import copy as copy_style
    from openpyxl import load_workbook

    source_workbook = load_workbook(path, data_only=False)
    try:
        source_worksheet = (
            source_workbook[sheet_name]
            if sheet_name and sheet_name in source_workbook.sheetnames
            else source_workbook[source_workbook.sheetnames[0]]
        )
        for row in source_worksheet.iter_rows():
            for source_cell in row:
                target_cell = target_worksheet.cell(
                    row=source_cell.row,
                    column=source_cell.column,
                    value=source_cell.value,
                )
                if source_cell.has_style:
                    target_cell._style = copy_style(source_cell._style)
                    target_cell.number_format = source_cell.number_format
                    target_cell.font = copy_style(source_cell.font)
                    target_cell.fill = copy_style(source_cell.fill)
                    target_cell.border = copy_style(source_cell.border)
                    target_cell.alignment = copy_style(source_cell.alignment)
                    target_cell.protection = copy_style(source_cell.protection)
                if source_cell.hyperlink:
                    target_cell._hyperlink = copy_style(source_cell.hyperlink)
                if source_cell.comment:
                    target_cell.comment = copy_style(source_cell.comment)
        for key, dimension in source_worksheet.column_dimensions.items():
            target_dimension = target_worksheet.column_dimensions[key]
            target_dimension.width = dimension.width
            target_dimension.hidden = dimension.hidden
        for key, dimension in source_worksheet.row_dimensions.items():
            target_dimension = target_worksheet.row_dimensions[key]
            target_dimension.height = dimension.height
            target_dimension.hidden = dimension.hidden
        for merged_range in source_worksheet.merged_cells.ranges:
            target_worksheet.merge_cells(str(merged_range))
        if source_worksheet.freeze_panes:
            target_worksheet.freeze_panes = source_worksheet.freeze_panes
        if source_worksheet.auto_filter.ref:
            target_worksheet.auto_filter.ref = source_worksheet.auto_filter.ref
    finally:
        source_workbook.close()


def _copy_csv_input_sheet_openpyxl(path: Path, worksheet: Any) -> None:
    import csv

    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        for row_index, row in enumerate(csv.reader(handle), start=1):
            for col_index, value in enumerate(row, start=1):
                worksheet.cell(row=row_index, column=col_index, value=value)


def _copy_docx_input_sheet_openpyxl(path: Path, worksheet: Any) -> None:
    for row_index, line in enumerate(_docx_markdown_lines(path), start=1):
        worksheet.cell(row=row_index, column=1, value=line)


def _write_input_source_to_xlsxwriter(
    workbook: Any,
    sheet_name: str,
    source: _InputSheetSource,
    formats: dict[str, Any],
    *,
    missing_note: str,
) -> None:
    worksheet = workbook.add_worksheet(sheet_name)
    if source.path is None:
        _write(worksheet, 0, 0, missing_note, formats.get("italic"))
    elif source.kind == "excel":
        _copy_excel_input_sheet_xlsxwriter(source.path, source.sheet_name, worksheet)
    elif source.kind == "csv":
        _copy_csv_input_sheet_xlsxwriter(source.path, worksheet)
    elif source.kind == "docx":
        for row_index, line in enumerate(_docx_markdown_lines(source.path)):
            _write(worksheet, row_index, 0, line)
    else:
        _write(worksheet, 0, 0, f"Unsupported input file type: {source.path.name}")
    _autofit(worksheet)


def _copy_excel_input_sheet_xlsxwriter(path: Path, sheet_name: str | None, worksheet: Any) -> None:
    from openpyxl import load_workbook

    source_workbook = load_workbook(path, data_only=False)
    try:
        source_worksheet = (
            source_workbook[sheet_name]
            if sheet_name and sheet_name in source_workbook.sheetnames
            else source_workbook[source_workbook.sheetnames[0]]
        )
        for row in source_worksheet.iter_rows():
            for cell in row:
                _write(worksheet, cell.row - 1, cell.column - 1, cell.value)
        for key, dimension in source_worksheet.column_dimensions.items():
            if dimension.width:
                worksheet.set_column(_column_index_from_letter(key), _column_index_from_letter(key), float(dimension.width))
    finally:
        source_workbook.close()


def _copy_csv_input_sheet_xlsxwriter(path: Path, worksheet: Any) -> None:
    import csv

    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        for row_index, row in enumerate(csv.reader(handle)):
            for col_index, value in enumerate(row):
                _write(worksheet, row_index, col_index, value)


def _docx_markdown_lines(path: Path) -> list[str]:
    try:
        from docx import Document
    except Exception:
        return [f"Word input could not be rendered because python-docx is unavailable: {path.name}"]

    document = Document(str(path))
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = getattr(paragraph.style, "name", "") or ""
        if style_name.lower().startswith("heading"):
            level = "".join(ch for ch in style_name if ch.isdigit()) or "1"
            prefix = "#" * max(1, min(6, int(level)))
            lines.append(f"{prefix} {text}")
        else:
            lines.append(text)
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            lines.append("| " + " | ".join(values) + " |")
    return lines or [f"Word input contained no readable paragraphs: {path.name}"]


def _column_index_from_letter(column_letter: str) -> int:
    value = 0
    for char in column_letter.upper():
        if "A" <= char <= "Z":
            value = value * 26 + (ord(char) - ord("A") + 1)
    return max(0, value - 1)


def _move_input_and_helper_sheets_to_end(workbook: Any) -> None:
    end_order = [
        INPUT_RAW_SHEET_NAME,
        INPUT_DATAMAP_SHEET_NAME,
        "_RawData",
        "_Options",
        "Filters",
    ]
    tail = [workbook[name] for name in end_order if name in workbook.sheetnames]
    tail_ids = {id(sheet) for sheet in tail}
    workbook._sheets = [sheet for sheet in workbook._sheets if id(sheet) not in tail_ids] + tail


def _wvl_countifs(
    data_name: str,
    criteria: Any,
    fq_name: str,
    cohort_mask_name: str,
    theme_prefix: str,
) -> str:
    return _build_countifs_formula(
        data_name,
        criteria,
        [],
        fq_name,
        "",
        extra_pairs=[(cohort_mask_name, "TRUE")],
        theme_prefix=theme_prefix,
    )


def _wvl_total_countifs(
    data_name: str,
    criteria: Any,
    fq_name: str,
    theme_prefix: str,
) -> str:
    return _build_countifs_formula(
        data_name,
        criteria,
        [],
        fq_name,
        "",
        theme_prefix=theme_prefix,
    )


def _wvl_total_respondents_formula(
    fq_name: str,
    theme_prefix: str,
) -> str:
    filter_pairs = _live_filter_criteria_pairs(theme_prefix, fq_name)
    args = ",".join(f"{range_expr},{criterion_expr}" for range_expr, criterion_expr in filter_pairs)
    return f"=COUNTIFS({args})"


def _wvl_averageifs(
    data_name: str,
    fq_name: str,
    cohort_mask_name: str,
    theme_prefix: str,
) -> str:
    filter_pairs = _live_filter_criteria_pairs(theme_prefix, fq_name)
    args = ",".join(f"{range_expr},{criterion_expr}" for range_expr, criterion_expr in filter_pairs)
    return f'=IFERROR(AVERAGEIFS({data_name},{data_name},">=0",{cohort_mask_name},TRUE,{args}),"-")'


def _wvl_total_averageifs(
    data_name: str,
    fq_name: str,
    theme_prefix: str,
) -> str:
    filter_pairs = _live_filter_criteria_pairs(theme_prefix, fq_name)
    args = ",".join(f"{range_expr},{criterion_expr}" for range_expr, criterion_expr in filter_pairs)
    return f'=IFERROR(AVERAGEIFS({data_name},{data_name},">=0",{args}),"-")'


def _wvl_filtered_numeric_formula(
    metric: str,
    data_name: str,
    fq_name: str,
    cohort_mask_name: str,
    theme_prefix: str,
) -> str:
    filter_pairs = _live_filter_criteria_pairs(theme_prefix, fq_name)
    criteria = [f"({cohort_mask_name}=TRUE)", f"({data_name}>=0)"]
    criteria.extend(f"({range_expr}={criterion_expr})" for range_expr, criterion_expr in filter_pairs)
    mask_expr = "*".join(criteria)
    if metric == "Median":
        return f"=IFERROR(MEDIAN(FILTER({data_name},{mask_expr})),0)"
    if metric == "Min":
        return f"=IFERROR(MIN(FILTER({data_name},{mask_expr})),0)"
    if metric == "Max":
        return f"=IFERROR(MAX(FILTER({data_name},{mask_expr})),0)"
    return f'=IFERROR(AVERAGEIFS({data_name},{data_name},">=0",{cohort_mask_name},TRUE),0)'


def _wvl_grid_rated_source_column(question: Any, row_id: str, column_index: int) -> str:
    raw_columns = list(getattr(question, "raw_columns", ()) or ())
    matching = [column for column in raw_columns if str(column).startswith(str(row_id))]
    if matching and column_index < len(matching):
        return matching[column_index]
    if row_id in raw_columns:
        return row_id
    candidate = f"{row_id}c{column_index + 1}"
    return candidate if candidate in raw_columns else (raw_columns[column_index] if column_index < len(raw_columns) else candidate)


def _record_wvl_cohort_audit(
    log: CalculationLog,
    metric_name: str,
    question: Any,
    segment_definition: SegmentDefinition,
    cohort_n: int,
) -> None:
    log.record(
        AuditRecord(
            output_sheet="Run_Summary",
            metric_name=metric_name,
            source_question_id=question.canonical_id,
            source_columns=tuple(question.raw_columns),
            filter_expr=_wvl_segment_description(segment_definition, laggard=metric_name == "laggard_mask"),
            numerator=cohort_n,
            denominator=None,
            formula=f"{metric_name} construction from SegmentDefinition",
            value_raw=float(cohort_n),
            valid_n=cohort_n,
            missing_n=0,
            timestamp=datetime.now(timezone.utc),
        )
    )


def _record_wvl_manual_cohort_audit(
    log: CalculationLog,
    metric_name: str,
    schema: SurveySchema,
    segment_definition: SegmentDefinition,
    cohort_n: int,
) -> None:
    log.record(
        AuditRecord(
            output_sheet="Run_Summary",
            metric_name=metric_name,
            source_question_id=segment_definition.outcome_question_id,
            source_columns=(schema.respondent_id_column,),
            filter_expr=_wvl_segment_description(
                segment_definition,
                laggard=metric_name == "laggard_mask",
            ),
            numerator=cohort_n,
            denominator=None,
            formula=f"{metric_name} construction from manual respondent-id list",
            value_raw=float(cohort_n),
            valid_n=cohort_n,
            missing_n=0,
            timestamp=datetime.now(timezone.utc),
        )
    )


def _record_wvl_metric(
    log: CalculationLog,
    sheet_name: str,
    question_id: str,
    source_columns: tuple[str, ...],
    metric_name: str,
    formula: str,
    value: float | int,
) -> None:
    log.record(
        AuditRecord(
            output_sheet=sheet_name,
            metric_name=metric_name,
            source_question_id=question_id,
            source_columns=source_columns or (question_id,),
            filter_expr="workbook filters + winners/laggards mask",
            numerator=None,
            denominator=None,
            formula=formula,
            value_raw=float(value or 0),
            valid_n=0,
            missing_n=0,
            timestamp=datetime.now(timezone.utc),
        )
    )


@dataclass
class _StreamingHelperSpec:
    header: str
    data_name: str
    formula_builder: Any


_FORMULA_CACHE_ATTR = "_sie_formula_cache_values"
_RAW_STREAM_XML_ATTR = "_sie_raw_stream_sheet_xml"
_SPREADSHEET_NS = "http://schemas.openxmlformats.org/spreadsheetml/2006/main"
_RELATIONSHIP_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_PACKAGE_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
_CONTENT_TYPES_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
_CALC_CHAIN_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.spreadsheetml.calcChain+xml"
)
_CALC_CHAIN_REL_TYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/calcChain"
)
_QUESTION_PREFIX_PATTERN = re.compile(r"^\s*Q\d+[A-Za-z]*\s*[-:]?\s*", re.IGNORECASE)
_CELL_XML_RE = re.compile(rb'<c\b[^>]*\br="[^"]+"[^>]*>.*?</c>', re.DOTALL)
_XML_COPY_CHUNK_SIZE = 1024 * 1024


def _live_formula(
    worksheet: Any,
    row: int,
    column: int,
    formula: str,
    cached_value: Any,
) -> Any:
    """Write an openpyxl formula and remember the cached value to patch into XML."""

    cell = worksheet.cell(row=row, column=column, value=formula)
    _record_formula_cache(worksheet.parent, worksheet.title, cell.coordinate, cached_value)
    return cell


def _record_formula_cache(
    workbook: Any,
    sheet_name: str,
    coordinate: str,
    cached_value: Any,
) -> None:
    cache = getattr(workbook, _FORMULA_CACHE_ATTR, None)
    if cache is None:
        cache = {}
        setattr(workbook, _FORMULA_CACHE_ATTR, cache)
    cache.setdefault(sheet_name, {})[coordinate] = _normalise_formula_cache_value(
        cached_value
    )


def _normalise_formula_cache_value(value: Any) -> Any:
    if value is None:
        return 0
    if isinstance(value, bool):
        return int(value)
    return value


def _prepare_live_columns(schema: SurveySchema, context: _LiveWorkbookContext) -> None:
    columns = _live_column_specs(schema)
    context.columns = columns
    context.column_by_key = {column.key: column for column in columns}


def _raw_data_source_row_count(
    decoded_df: Any | None,
    schema: SurveySchema,
    results: list[SingleCutResult],
) -> int:
    if decoded_df is not None:
        shape = getattr(decoded_df, "shape", None)
        if shape:
            return int(shape[0])
        try:
            return len(decoded_df)
        except TypeError:
            pass
    return _synthetic_raw_row_count(schema, results)


def _synthetic_raw_row_count(
    schema: SurveySchema,
    results: list[SingleCutResult],
) -> int:
    result_counts = [result.valid_n + result.missing_n for result in results]
    return max([schema.total_respondents, *result_counts])


def _should_stream_raw_data(row_count: int) -> bool:
    return (
        row_count > _RAW_DATA_STREAMING_ROW_THRESHOLD
        and row_count <= int(RAW_DATA_SHEET_ROW_LIMIT)
    )


def _build_raw_data_sheet(
    workbook: Any,
    decoded_df: Any | None,
    schema: SurveySchema,
    results: list[SingleCutResult],
    context: _LiveWorkbookContext,
) -> None:
    worksheet = workbook.create_sheet("_RawData")
    worksheet.sheet_state = "hidden"
    if not context.columns:
        _prepare_live_columns(schema, context)
    columns = context.columns

    worksheet.cell(row=1, column=1, value="respondent_id")
    for index, column in enumerate(columns, start=2):
        worksheet.cell(row=1, column=index, value=column.header)

    rows = _raw_rows_from_dataframe(decoded_df, schema, columns)
    if not rows:
        rows = _synthetic_raw_rows(schema, results, columns)

    for row_index, row_payload in enumerate(rows, start=2):
        worksheet.cell(row=row_index, column=1, value=row_payload.get("respondent_id", row_index - 1))
        for col_index, column in enumerate(columns, start=2):
            value = row_payload.get(column.key)
            worksheet.cell(row=row_index, column=col_index, value=value)

    last_row = max(2, len(rows) + 1)
    _add_named_range(workbook, "respondent_id_data", "_RawData", f"$A$2:$A${last_row}")
    for index, column in enumerate(columns, start=2):
        col_letter = _openpyxl_column_letter(index)
        _add_named_range(
            workbook,
            column.data_name,
            "_RawData",
            f"${col_letter}$2:${col_letter}${last_row}",
        )

    for column in columns:
        values = []
        for row_payload in rows:
            value = row_payload.get(column.key)
            if _live_value_present(value):
                text = str(value)
                if text not in values:
                    values.append(text)
        if len(values) > 100:
            context.warnings.append(
                f"{column.key} has more than 100 unique values; dropdown capped at 100"
            )
            values = values[:100]
        context.option_values[column.key] = values
    worksheet.freeze_panes = "A2"


def _build_raw_data_sheet_streaming(
    workbook: Any,
    decoded_df: Any | None,
    schema: SurveySchema,
    results: list[SingleCutResult],
    context: _LiveWorkbookContext,
    sheet_filters: list[dict[str, str]],
    theme_groups: list[tuple[str, list[SingleCutResult]]],
    theme_sheet_names: dict[str, str],
) -> None:
    from openpyxl import Workbook
    from openpyxl.cell import WriteOnlyCell

    if not context.columns:
        _prepare_live_columns(schema, context)
    columns = context.columns
    base_col_count = len(columns) + 1
    helpers = _streaming_helper_specs(
        sheet_filters=sheet_filters,
        theme_groups=theme_groups,
        theme_sheet_names=theme_sheet_names,
        schema=schema,
        context=context,
        base_col_count=base_col_count,
    )

    placeholder = workbook.create_sheet("_RawData")
    placeholder.sheet_state = "hidden"

    stream_workbook = Workbook(write_only=True)
    worksheet = stream_workbook.create_sheet("_RawData")

    headers = ["respondent_id", *(column.header for column in columns)]
    headers.extend(helper.header for helper in helpers)
    worksheet.append([WriteOnlyCell(worksheet, value=value) for value in headers])

    option_value_sets: dict[str, set[str]] = {column.key: set() for column in columns}
    row_count = 0
    for row_index, row_payload in enumerate(
        _iter_raw_row_payloads(decoded_df, schema, results, columns),
        start=2,
    ):
        row_count += 1
        row_cells = [
            WriteOnlyCell(
                worksheet,
                value=row_payload.get("respondent_id", row_index - 1),
            )
        ]
        for column in columns:
            value = row_payload.get(column.key)
            row_cells.append(WriteOnlyCell(worksheet, value=value))
            if _live_value_present(value):
                option_value_sets[column.key].add(str(value))
        for col_index, helper in enumerate(helpers, start=base_col_count + 1):
            formula = helper.formula_builder(row_index)
            row_cells.append(WriteOnlyCell(worksheet, value=formula))
            _record_formula_cache(
                workbook,
                worksheet.title,
                f"{_openpyxl_column_letter(col_index)}{row_index}",
                1,
            )
        worksheet.append(row_cells)

    last_row = max(2, row_count + 1)
    context.raw_data_streamed = True
    context.raw_data_row_count = row_count
    _define_raw_data_named_ranges(workbook, columns, helpers, last_row)
    _finalise_option_values(option_value_sets, context)

    buffer = BytesIO()
    stream_workbook.save(buffer)
    buffer.seek(0)
    with ZipFile(buffer, "r") as archive:
        setattr(
            workbook,
            _RAW_STREAM_XML_ATTR,
            archive.read("xl/worksheets/sheet1.xml"),
        )


def _define_raw_data_named_ranges(
    workbook: Any,
    columns: list[_LiveColumnSpec],
    helpers: list[_StreamingHelperSpec],
    last_row: int,
) -> None:
    _add_named_range(workbook, "respondent_id_data", "_RawData", f"$A$2:$A${last_row}")
    for index, column in enumerate(columns, start=2):
        col_letter = _openpyxl_column_letter(index)
        _add_named_range(
            workbook,
            column.data_name,
            "_RawData",
            f"${col_letter}$2:${col_letter}${last_row}",
        )
    for index, helper in enumerate(helpers, start=len(columns) + 2):
        col_letter = _openpyxl_column_letter(index)
        _add_named_range(
            workbook,
            helper.data_name,
            "_RawData",
            f"${col_letter}$2:${col_letter}${last_row}",
        )


def _finalise_option_values(
    option_value_sets: dict[str, set[str]],
    context: _LiveWorkbookContext,
) -> None:
    for key, value_set in option_value_sets.items():
        values = sorted(value_set)
        if len(values) > 100:
            context.warnings.append(
                f"{key} has more than 100 unique values; dropdown capped at 100"
            )
            values = values[:100]
        context.option_values[key] = values


def _build_options_sheet(
    workbook: Any,
    schema: SurveySchema,
    context: _LiveWorkbookContext,
) -> None:
    worksheet = workbook.create_sheet("_Options")
    worksheet.sheet_state = "hidden"

    all_questions_col = 1
    worksheet.cell(row=1, column=all_questions_col, value="All_Questions")
    worksheet.cell(row=2, column=all_questions_col, value="(None)")
    question_entries = _all_questions_dropdown_entries(schema, context)
    cross_tab_entries = _cross_tab_dropdown_entries(schema, context)
    for row_index, entry in enumerate(question_entries, start=3):
        worksheet.cell(row=row_index, column=all_questions_col, value=entry["label"])
    _add_named_range(
        workbook,
        "All_Questions",
        "_Options",
        f"$A$2:$A${len(question_entries) + 2}",
    )

    worksheet.cell(row=1, column=2, value="All_Questions_Local")
    worksheet.cell(row=2, column=2, value="(Inherit)")
    worksheet.cell(row=3, column=2, value="(None)")
    for row_index, entry in enumerate(question_entries, start=4):
        worksheet.cell(row=row_index, column=2, value=entry["label"])
    _add_named_range(
        workbook,
        "All_Questions_Local",
        "_Options",
        f"$B$2:$B${len(question_entries) + 3}",
    )

    lookup_start_col = 6 + (2 * len(context.columns))
    cross_tab_col = lookup_start_col
    worksheet.cell(row=1, column=cross_tab_col, value="Cross_Tab_Questions")
    worksheet.cell(row=2, column=cross_tab_col, value="(None)")
    for row_index, entry in enumerate(cross_tab_entries, start=3):
        worksheet.cell(row=row_index, column=cross_tab_col, value=entry["label"])
    cross_tab_letter = _openpyxl_column_letter(cross_tab_col)
    _add_named_range(
        workbook,
        "Cross_Tab_Questions",
        "_Options",
        f"${cross_tab_letter}$2:${cross_tab_letter}${len(cross_tab_entries) + 2}",
    )

    cross_tab_data_col = lookup_start_col + 1
    worksheet.cell(row=1, column=cross_tab_data_col, value="Cross_Tab_Questions_Data_Names")
    worksheet.cell(
        row=2,
        column=cross_tab_data_col,
        value=_defined_name_reference(workbook, "respondent_id_data"),
    )
    for row_index, entry in enumerate(cross_tab_entries, start=3):
        worksheet.cell(
            row=row_index,
            column=cross_tab_data_col,
            value=_defined_name_reference(workbook, entry["data_name"]),
        )
    cross_tab_data_letter = _openpyxl_column_letter(cross_tab_data_col)
    _add_named_range(
        workbook,
        "Cross_Tab_Questions_Data_Names",
        "_Options",
        f"${cross_tab_data_letter}$2:${cross_tab_data_letter}${len(cross_tab_entries) + 2}",
    )

    worksheet.cell(row=1, column=5, value="None_options")
    worksheet.cell(row=2, column=5, value="(All)")
    _add_named_range(workbook, "None_options", "_Options", "$E$2:$E$2")

    option_range_refs: dict[str, str] = {
        "None_options": _sheet_range_reference("_Options", "$E$2:$E$2")
    }
    for col_index, column in enumerate(context.columns, start=6):
        worksheet.cell(row=1, column=col_index, value=f"{column.header}_options")
        worksheet.cell(row=2, column=col_index, value="(All)")
        worksheet.cell(row=1, column=col_index + len(context.columns), value=f"{column.header}_local_options")
        worksheet.cell(row=2, column=col_index + len(context.columns), value="(Inherit)")
        worksheet.cell(row=3, column=col_index + len(context.columns), value="(All)")
        values = _option_values_for_column(column, schema, context)
        for row_offset, value in enumerate(values[:100], start=3):
            worksheet.cell(row=row_offset, column=col_index, value=value)
        local_col_index = col_index + len(context.columns)
        for row_offset, value in enumerate(values[:100], start=4):
            worksheet.cell(row=row_offset, column=local_col_index, value=value)
        col_letter = _openpyxl_column_letter(col_index)
        last_row = max(2, min(102, len(values) + 2))
        _add_named_range(
            workbook,
            f"{column.header}_options",
            "_Options",
            f"${col_letter}$2:${col_letter}${last_row}",
        )
        option_range_refs[f"{column.header}_options"] = _sheet_range_reference(
            "_Options",
            f"${col_letter}$2:${col_letter}${last_row}",
        )
        local_col_letter = _openpyxl_column_letter(local_col_index)
        local_last_row = max(3, min(103, len(values) + 3))
        _add_named_range(
            workbook,
            f"{column.header}_local_options",
            "_Options",
            f"${local_col_letter}$2:${local_col_letter}${local_last_row}",
        )
        option_range_refs[f"{column.header}_local_options"] = _sheet_range_reference(
            "_Options",
            f"${local_col_letter}$2:${local_col_letter}${local_last_row}",
        )

    question_option_start_col = lookup_start_col + 3
    for offset, entry in enumerate(question_entries):
        values = [str(value) for value in entry.get("option_values", [])][:100]
        if not values:
            continue
        option_name = entry.get("options_name", "")
        if not option_name or option_name in option_range_refs:
            continue
        col_index = question_option_start_col + (offset * 2)
        worksheet.cell(row=1, column=col_index, value=option_name)
        worksheet.cell(row=2, column=col_index, value="(All)")
        for row_offset, value in enumerate(values, start=3):
            worksheet.cell(row=row_offset, column=col_index, value=value)
        col_letter = _openpyxl_column_letter(col_index)
        last_row = max(2, len(values) + 2)
        _add_named_range(
            workbook,
            option_name,
            "_Options",
            f"${col_letter}$2:${col_letter}${last_row}",
        )
        option_range_refs[option_name] = _sheet_range_reference(
            "_Options",
            f"${col_letter}$2:${col_letter}${last_row}",
        )

        local_option_name = entry.get("local_options_name", "")
        if not local_option_name:
            continue
        local_col_index = col_index + 1
        worksheet.cell(row=1, column=local_col_index, value=local_option_name)
        worksheet.cell(row=2, column=local_col_index, value="(Inherit)")
        worksheet.cell(row=3, column=local_col_index, value="(All)")
        for row_offset, value in enumerate(values, start=4):
            worksheet.cell(row=row_offset, column=local_col_index, value=value)
        local_col_letter = _openpyxl_column_letter(local_col_index)
        local_last_row = max(3, len(values) + 3)
        _add_named_range(
            workbook,
            local_option_name,
            "_Options",
            f"${local_col_letter}$2:${local_col_letter}${local_last_row}",
        )
        option_range_refs[local_option_name] = _sheet_range_reference(
            "_Options",
            f"${local_col_letter}$2:${local_col_letter}${local_last_row}",
        )

    worksheet.cell(row=1, column=3, value="All_Questions_Data_Names")
    worksheet.cell(
        row=2,
        column=3,
        value=_defined_name_reference(workbook, "respondent_id_data"),
    )
    for row_index, entry in enumerate(question_entries, start=3):
        worksheet.cell(
            row=row_index,
            column=3,
            value=_defined_name_reference(workbook, entry["data_name"]),
        )
    _add_named_range(
        workbook,
        "All_Questions_Data_Names",
        "_Options",
        f"$C$2:$C${len(question_entries) + 2}",
    )

    worksheet.cell(row=1, column=4, value="All_Questions_Options_Names")
    worksheet.cell(row=2, column=4, value=option_range_refs["None_options"])
    for row_index, entry in enumerate(question_entries, start=3):
        worksheet.cell(
            row=row_index,
            column=4,
            value=option_range_refs.get(entry["options_name"], option_range_refs["None_options"]),
        )
    _add_named_range(
        workbook,
        "All_Questions_Options_Names",
        "_Options",
        f"$D$2:$D${len(question_entries) + 2}",
    )

    cross_tab_options_col = lookup_start_col + 2
    worksheet.cell(row=1, column=cross_tab_options_col, value="Cross_Tab_Questions_Options_Names")
    worksheet.cell(row=2, column=cross_tab_options_col, value=option_range_refs["None_options"])
    for row_index, entry in enumerate(cross_tab_entries, start=3):
        worksheet.cell(
            row=row_index,
            column=cross_tab_options_col,
            value=option_range_refs.get(entry["options_name"], option_range_refs["None_options"]),
        )
    cross_tab_options_letter = _openpyxl_column_letter(cross_tab_options_col)
    _add_named_range(
        workbook,
        "Cross_Tab_Questions_Options_Names",
        "_Options",
        f"${cross_tab_options_letter}$2:${cross_tab_options_letter}${len(cross_tab_entries) + 2}",
    )

    value_lookup_rows = _filter_value_lookup_rows(question_entries, workbook)
    value_key_col = question_option_start_col + (len(question_entries) * 2) + 1
    value_data_col = value_key_col + 1
    value_criteria_col = value_key_col + 2
    worksheet.cell(row=1, column=value_key_col, value="All_Questions_Value_Keys")
    worksheet.cell(row=1, column=value_data_col, value="All_Questions_Value_Data_Names")
    worksheet.cell(row=1, column=value_criteria_col, value="All_Questions_Value_Criteria")
    for row_index, row in enumerate(value_lookup_rows, start=2):
        worksheet.cell(row=row_index, column=value_key_col, value=row["key"])
        worksheet.cell(row=row_index, column=value_data_col, value=row["data_ref"])
        worksheet.cell(row=row_index, column=value_criteria_col, value=row["criteria"])
    value_last_row = max(2, len(value_lookup_rows) + 1)
    value_key_letter = _openpyxl_column_letter(value_key_col)
    value_data_letter = _openpyxl_column_letter(value_data_col)
    value_criteria_letter = _openpyxl_column_letter(value_criteria_col)
    _add_named_range(
        workbook,
        "All_Questions_Value_Keys",
        "_Options",
        f"${value_key_letter}$2:${value_key_letter}${value_last_row}",
    )
    _add_named_range(
        workbook,
        "All_Questions_Value_Data_Names",
        "_Options",
        f"${value_data_letter}$2:${value_data_letter}${value_last_row}",
    )
    _add_named_range(
        workbook,
        "All_Questions_Value_Criteria",
        "_Options",
        f"${value_criteria_letter}$2:${value_criteria_letter}${value_last_row}",
    )


def _option_values_for_column(
    column: _LiveColumnSpec,
    schema: SurveySchema,
    context: _LiveWorkbookContext,
) -> list[str]:
    question = column.question
    if question is None:
        return []
    if column.kind == "nps_bucket":
        return [label for label, _low, _high in NPS_BUCKETS]
    if column.kind == "single":
        result = _single_select_result_for_question(context, question.canonical_id)
        if result is not None:
            ordered_payloads = sorted(
                result.distribution.values(),
                key=lambda payload: int(payload.get("count", 0)),
                reverse=True,
            )
            return [str(payload["label"]) for payload in ordered_payloads]
    if column.key in context.option_values and context.option_values[column.key]:
        return context.option_values[column.key]
    if column.kind in {"multi_select", "grid_single", "grid_binary"}:
        return ["Selected"]
    if question.option_map and column.kind == "single":
        return [str(value) for value in question.option_map.values()]
    return []


def _single_select_result_for_question(
    context: _LiveWorkbookContext,
    question_id: str,
) -> SingleSelectResult | None:
    for result in context.results:
        if (
            result.question_id == question_id
            and isinstance(result, SingleSelectResult)
        ):
            return result
    return None


def _all_questions_dropdown_entries(
    schema: SurveySchema,
    context: _LiveWorkbookContext,
) -> list[dict[str, Any]]:
    """Return filter choices and their aligned lookup metadata."""

    entries: list[dict[str, Any]] = []
    used_labels: set[str] = {"(None)", "(Inherit)"}
    result_question_ids = {result.question_id for result in context.results}
    for filter_option in filter_question_options(schema):
        question = schema.get_question(filter_option.question_id)
        if question is None:
            continue
        if question.canonical_id not in result_question_ids:
            continue
        if filter_option.question_type in {
            QuestionType.SINGLE_SELECT,
            QuestionType.DEMOGRAPHIC_OR_SEGMENT,
        }:
            column = context.column_by_key.get(question.canonical_id)
            if column is None:
                continue
            label = _dedupe_dropdown_label(
                _question_dropdown_label(question, context),
                used_labels,
            )
            entries.append(
                {
                    "label": label,
                    "data_name": column.data_name,
                    "options_name": f"{column.header}_options",
                    "local_options_name": f"{column.header}_local_options",
                    "option_values": (),
                    "value_lookup": (),
                }
            )
            continue
        if filter_option.question_type is QuestionType.NPS:
            column = context.column_by_key.get(_nps_filter_column_key(question.canonical_id))
            if column is None:
                continue
            label = _dedupe_dropdown_label(
                _question_dropdown_label(question, context),
                used_labels,
            )
            entries.append(
                {
                    "label": label,
                    "data_name": column.data_name,
                    "options_name": f"{column.header}_options",
                    "local_options_name": f"{column.header}_local_options",
                    "option_values": (),
                    "value_lookup": (),
                }
            )
            continue

        option_values: list[str] = []
        value_lookup: list[dict[str, str]] = []
        first_column: _LiveColumnSpec | None = None
        for value_option in filter_option.values:
            column = context.column_by_key.get(value_option.filter_question_id)
            if column is None:
                continue
            first_column = first_column or column
            value_label = str(value_option.label)
            if value_label not in option_values:
                option_values.append(value_label)
            value_lookup.append(
                {
                    "value_label": value_label,
                    "data_name": column.data_name,
                    "criteria": str(_filter_value_criteria(question, value_option)),
                }
            )
        if first_column is None or not value_lookup:
            continue
        label = _dedupe_dropdown_label(
            _question_dropdown_label(question, context),
            used_labels,
        )
        options_prefix = _safe_defined_name(f"{question.canonical_id}_filter_values")
        entries.append(
            {
                "label": label,
                "data_name": first_column.data_name,
                "options_name": f"{options_prefix}_options",
                "local_options_name": f"{options_prefix}_local_options",
                "option_values": tuple(option_values),
                "value_lookup": tuple(value_lookup),
            }
        )
    return entries


def _cross_tab_dropdown_entries(
    schema: SurveySchema,
    context: _LiveWorkbookContext,
) -> list[dict[str, str]]:
    """Return categorical dimensions for cross-tab controls, one row per question."""

    entries: list[dict[str, str]] = []
    used_labels: set[str] = {"(None)"}
    for question in schema.analysis_eligible_questions():
        if question.question_type not in {
            QuestionType.SINGLE_SELECT,
            QuestionType.DEMOGRAPHIC_OR_SEGMENT,
        }:
            continue
        column = context.column_by_key.get(question.canonical_id)
        if column is None:
            continue
        label = _dedupe_dropdown_label(
            _question_dropdown_label(question, context),
            used_labels,
        )
        entries.append(
            {
                "label": label,
                "data_name": column.data_name,
                "options_name": f"{column.header}_options",
            }
        )
    return entries


def _filter_value_criteria(question: Any, value_option: Any) -> Any:
    if question.question_type in {
        QuestionType.MULTI_SELECT_BINARY,
        QuestionType.GRID_BINARY_SELECT,
    }:
        return "Selected"
    if question.question_type is QuestionType.GRID_SINGLE_SELECT:
        if getattr(question, "label_to_numeric_value", None):
            return value_option.filter_value
        return question.option_map.get(value_option.filter_value, value_option.filter_value)
    return value_option.label


def _filter_value_lookup_rows(
    question_entries: list[dict[str, Any]],
    workbook: Any,
) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    for entry in question_entries:
        label = str(entry.get("label", ""))
        for value_row in entry.get("value_lookup", ()) or ():
            value_label = str(value_row.get("value_label", ""))
            if not label or not value_label:
                continue
            rows.append(
                {
                    "key": f"{label}|{value_label}",
                    "data_ref": _defined_name_reference(workbook, value_row["data_name"]),
                    "criteria": str(value_row.get("criteria", value_label)),
                }
            )
    return rows


def _is_all_questions_dropdown_eligible(question: Any) -> bool:
    return (
        getattr(question, "question_type", None) is QuestionType.SINGLE_SELECT
        and bool(getattr(question, "analysis_eligible", False))
        and len(getattr(question, "option_map", {}) or {}) >= 2
    )


def _dedupe_dropdown_label(label: str, used_labels: set[str]) -> str:
    candidate = label
    suffix = 2
    while candidate in used_labels:
        candidate = f"{label} {suffix}"
        suffix += 1
    used_labels.add(candidate)
    return candidate


def _build_filters_sheet(
    workbook: Any,
    schema: SurveySchema,
    results: list[SingleCutResult],
    demographic_questions: list[Any],
    context: _LiveWorkbookContext,
    demo_priority: dict | None,
    planned_filters: list[dict[str, str]] | None = None,
) -> list[dict[str, str]]:
    """Write the single workbook-scoped filter control panel."""

    worksheet = workbook.create_sheet("Filters")
    worksheet.sheet_state = "visible"
    worksheet.freeze_panes = "A2"

    worksheet.cell(row=1, column=1, value="DEMOGRAPHIC FILTERS").font = _live_font(bold=True)
    _live_fill_row(worksheet, 1, 6, "F2F2F2")
    worksheet.cell(row=1, column=1).fill = _live_fill("CC0000")
    worksheet.cell(row=1, column=1).font = _live_font(bold=True, color="FFFFFF")

    worksheet.cell(
        row=2,
        column=1,
        value='To filter by multiple values, separate with ", " (comma + space). Example: India, USA, UK',
    ).font = _live_font(italic=True, color="666666")

    headers = ["Filter", "Value", "Wrapped", "Available values"]
    _live_header_row(worksheet, 3, headers)
    row_index = 4
    sheet_filters = (
        list(planned_filters)
        if planned_filters is not None
        else _planned_sheet_filters(demographic_questions, context, demo_priority)
    )

    for sheet_filter in sheet_filters:
        filter_name = sheet_filter["filter_name"]
        worksheet.cell(
            row=row_index,
            column=1,
            value=sheet_filter["label"],
        )
        if sheet_filter.get("question_text"):
            _add_comment(worksheet.cell(row=row_index, column=1), sheet_filter["question_text"])
        value_cell = worksheet.cell(row=row_index, column=2, value="(All)")
        wrapped_cell = _live_formula(
            worksheet,
            row_index,
            3,
            _wrapped_formula(filter_name),
            "|(All)|",
        )
        worksheet.cell(
            row=row_index,
            column=4,
            value=_available_values_for_data_name(sheet_filter["data_name"], context),
        )
        _add_named_cell(workbook, filter_name, worksheet, value_cell.coordinate)
        _add_named_cell(workbook, f"{filter_name}_wrapped", worksheet, wrapped_cell.coordinate)
        _add_dropdown_to_cell(worksheet, value_cell.coordinate, f'={sheet_filter["options_name"].replace("_local_options", "_options")}')
        row_index += 1

    row_index += 1
    worksheet.cell(row=row_index, column=1, value="CUSTOM FILTERS").font = _live_font(bold=True)
    _live_fill_row(worksheet, row_index, 6, "F2F2F2")
    row_index += 1
    for slot in range(1, context.workbook_custom_filter_count + 1):
        _live_write_custom_filter_slot(
            workbook,
            worksheet,
            row_index,
            1,
            f"F_Custom{slot}_Q",
            f"F_Custom{slot}_V",
            slot,
        )
        row_index += 1

    _live_set_filter_column_widths(worksheet)
    return sheet_filters


def _planned_sheet_filters(
    demographic_questions: list[Any],
    context: _LiveWorkbookContext,
    demo_priority: dict | None,
) -> list[dict[str, str]]:
    sheet_filters: list[dict[str, str]] = []
    used_filter_names: set[str] = set()
    categories = (demo_priority or {}).get("categories", {})
    for question in demographic_questions[:8]:
        category = categories.get(question.canonical_id, "")
        filter_name = _demo_filter_name(question, category, used_filter_names)
        used_filter_names.add(filter_name)
        column = context.column_by_key.get(question.canonical_id)
        if column is None and question.raw_columns:
            column = context.column_by_key.get(question.raw_columns[0])
        if column is None:
            continue
        sheet_filters.append(
            {
                "filter_name": filter_name,
                "value_name": filter_name,
                "wrapped_name": f"{filter_name}_wrapped",
                "data_name": column.data_name,
                "kind": column.kind,
                "label": _filter_display_label(question, context),
                "question_text": question.question_text,
                "options_name": f"{column.header}_local_options",
            }
        )
    return sheet_filters


def _build_raw_filter_helper_columns(
    workbook: Any,
    sheet_filters: list[dict[str, str]],
    context: _LiveWorkbookContext,
) -> None:
    """Append row-wise filter masks to _RawData for COUNTIFS formulas."""

    if context.raw_data_streamed:
        return

    worksheet = workbook["_RawData"]
    last_row = max(2, worksheet.max_row)
    next_col = worksheet.max_column + 1
    match_columns: list[int] = []

    for sheet_filter in sheet_filters:
        data_name = sheet_filter.get("data_name", "")
        data_column = _raw_data_column_index_for_name(context, data_name)
        if data_column is None:
            continue

        match_name = _safe_defined_name(f"{sheet_filter['filter_name']}_match_data")
        sheet_filter["match_data_name"] = match_name
        worksheet.cell(row=1, column=next_col, value=match_name.removesuffix("_data"))
        for row_index in range(2, last_row + 1):
            data_cell = f"{_openpyxl_column_letter(data_column)}{row_index}"
            _live_formula(
                worksheet,
                row_index,
                next_col,
                _raw_filter_match_formula(
                    data_cell=data_cell,
                    value_name=sheet_filter["value_name"],
                    wrapped_name=sheet_filter["wrapped_name"],
                ),
                1,
            )
        col_letter = _openpyxl_column_letter(next_col)
        _add_named_range(
            workbook,
            match_name,
            "_RawData",
            f"${col_letter}$2:${col_letter}${last_row}",
        )
        match_columns.append(next_col)
        next_col += 1

    worksheet.cell(row=1, column=next_col, value="passes_workbook_filters")
    for row_index in range(2, last_row + 1):
        if match_columns:
            refs = [
                f"{_openpyxl_column_letter(col_index)}{row_index}"
                for col_index in match_columns
            ]
            formula = "=" + "*".join(refs + ["1"])
        else:
            formula = "=1"
        _live_formula(worksheet, row_index, next_col, formula, 1)

    pass_col_letter = _openpyxl_column_letter(next_col)
    _add_named_range(
        workbook,
        PASS_WORKBOOK_FILTERS_DATA_NAME,
        "_RawData",
        f"${pass_col_letter}$2:${pass_col_letter}${last_row}",
    )
    next_col += 1

    custom_match_columns: list[int] = []
    for slot in range(1, context.workbook_custom_filter_count + 1):
        prefix = f"F_Custom{slot}"
        match_name = f"{prefix}_match_data"
        worksheet.cell(row=1, column=next_col, value=match_name.removesuffix("_data"))
        for row_index in range(2, last_row + 1):
            _live_formula(
                worksheet,
                row_index,
                next_col,
                _raw_custom_filter_match_formula(
                    row_index=row_index,
                    question_name=f"{prefix}_Q",
                    value_name=f"{prefix}_V",
                    wrapped_name=f"{prefix}_wrapped",
                    resolved_column_name=f"{prefix}_resolved_column",
                    criteria_name=f"{prefix}_resolved_criteria",
                ),
                1,
            )
        col_letter = _openpyxl_column_letter(next_col)
        _add_named_range(
            workbook,
            match_name,
            "_RawData",
            f"${col_letter}$2:${col_letter}${last_row}",
        )
        custom_match_columns.append(next_col)
        next_col += 1

    worksheet.cell(row=1, column=next_col, value="passes_workbook_custom_filters")
    for row_index in range(2, last_row + 1):
        refs = [
            f"{_openpyxl_column_letter(col_index)}{row_index}"
            for col_index in custom_match_columns
        ]
        _live_formula(
            worksheet,
            row_index,
            next_col,
            "=" + "*".join(refs + ["1"]),
            1,
        )
    custom_pass_col_letter = _openpyxl_column_letter(next_col)
    _add_named_range(
        workbook,
        PASS_WORKBOOK_CUSTOM_FILTERS_DATA_NAME,
        "_RawData",
        f"${custom_pass_col_letter}$2:${custom_pass_col_letter}${last_row}",
    )


def _raw_data_column_index_for_name(
    context: _LiveWorkbookContext,
    data_name: str,
) -> int | None:
    for index, column in enumerate(context.columns, start=2):
        if column.data_name == data_name:
            return index
    return None


def _raw_filter_match_formula(
    data_cell: str,
    value_name: str,
    wrapped_name: str,
) -> str:
    return (
        f'=((({value_name}="(All)")+({value_name}="")+ISBLANK({value_name})+'
        f'ISNUMBER(SEARCH("|"&{data_cell}&"|",{wrapped_name})))>0)*1'
    )


def _raw_custom_filter_match_formula(
    row_index: int,
    question_name: str,
    value_name: str,
    wrapped_name: str,
    resolved_column_name: str,
    criteria_name: str | None = None,
) -> str:
    selected_value = f"INDEX(INDIRECT({resolved_column_name}),{row_index - 1})"
    criteria_wrapped = (
        f'("|" & SUBSTITUTE({criteria_name}, ", ", "|") & "|")'
        if criteria_name
        else wrapped_name
    )
    return (
        f'=IF(OR({question_name}="(None)",{question_name}="",ISBLANK({question_name}),'
        f'{value_name}="(All)",{value_name}="",ISBLANK({value_name})),1,'
        f'IFERROR(IF(ISNUMBER(SEARCH("|"&{selected_value}&"|",{criteria_wrapped})),1,0),0))'
    )


def _streaming_helper_specs(
    sheet_filters: list[dict[str, str]],
    theme_groups: list[tuple[str, list[SingleCutResult]]],
    theme_sheet_names: dict[str, str],
    schema: SurveySchema,
    context: _LiveWorkbookContext,
    base_col_count: int,
) -> list[_StreamingHelperSpec]:
    helpers: list[_StreamingHelperSpec] = []
    data_indices = _raw_data_column_indices_by_name(context)

    workbook_match_columns: list[int] = []
    next_col = base_col_count + 1
    for sheet_filter in sheet_filters:
        data_column = data_indices.get(sheet_filter.get("data_name", ""))
        if data_column is None:
            continue
        match_name = _safe_defined_name(f"{sheet_filter['filter_name']}_match_data")
        sheet_filter["match_data_name"] = match_name
        helpers.append(
            _StreamingHelperSpec(
                header=match_name.removesuffix("_data"),
                data_name=match_name,
                formula_builder=_stream_filter_match_builder(
                    data_column=data_column,
                    value_name=sheet_filter["value_name"],
                    wrapped_name=sheet_filter["wrapped_name"],
                ),
            )
        )
        workbook_match_columns.append(next_col)
        next_col += 1

    helpers.append(
        _StreamingHelperSpec(
            header="passes_workbook_filters",
            data_name=PASS_WORKBOOK_FILTERS_DATA_NAME,
            formula_builder=_stream_product_builder(workbook_match_columns),
        )
    )
    next_col += 1

    custom_match_columns: list[int] = []
    for slot in range(1, context.workbook_custom_filter_count + 1):
        prefix = f"F_Custom{slot}"
        match_name = f"{prefix}_match_data"
        helpers.append(
            _StreamingHelperSpec(
                header=match_name.removesuffix("_data"),
                data_name=match_name,
                formula_builder=_stream_custom_match_builder(prefix),
            )
        )
        custom_match_columns.append(next_col)
        next_col += 1

    helpers.append(
        _StreamingHelperSpec(
            header="passes_workbook_custom_filters",
            data_name=PASS_WORKBOOK_CUSTOM_FILTERS_DATA_NAME,
            formula_builder=_stream_product_builder(custom_match_columns),
        )
    )
    next_col += 1

    for theme_name, theme_results in theme_groups:
        theme_prefix = _theme_defined_prefix(theme_sheet_names[theme_name])
        local_match_columns: list[int] = []
        for sheet_filter in sheet_filters:
            data_column = data_indices.get(sheet_filter.get("data_name", ""))
            if data_column is None:
                continue
            local_name = f"{theme_prefix}_{sheet_filter['filter_name']}"
            match_name = _safe_defined_name(
                f"{theme_prefix}_{sheet_filter['filter_name']}_local_match_data"
            )
            helpers.append(
                _StreamingHelperSpec(
                    header=match_name.removesuffix("_data"),
                    data_name=match_name,
                    formula_builder=_stream_filter_match_builder(
                        data_column=data_column,
                        value_name=f"{local_name}_resolved",
                        wrapped_name=f"{local_name}_wrapped",
                    ),
                )
            )
            local_match_columns.append(next_col)
            next_col += 1

        pass_name = _theme_local_pass_data_name(theme_prefix)
        helpers.append(
            _StreamingHelperSpec(
                header=pass_name.removesuffix("_data"),
                data_name=pass_name,
                formula_builder=_stream_product_builder(local_match_columns),
            )
        )
        next_col += 1

        for result in theme_results:
            question = schema.get_question(result.question_id)
            if question is None:
                continue
            if context.per_question_filter_count <= 0:
                continue
            q_filter_prefix = f"{theme_prefix}_{_safe_defined_name(question.canonical_id)}_F"
            pass_name = _per_question_pass_data_name(f"{q_filter_prefix}_Q")
            helpers.append(
                _StreamingHelperSpec(
                    header=pass_name.removesuffix("_data"),
                    data_name=pass_name,
                    formula_builder=_stream_per_question_match_builder(
                        q_filter_prefix,
                        context.per_question_filter_count,
                    ),
                )
            )
            next_col += 1

    return helpers


def _raw_data_column_indices_by_name(
    context: _LiveWorkbookContext,
) -> dict[str, int]:
    return {
        column.data_name: index
        for index, column in enumerate(context.columns, start=2)
    }


def _stream_filter_match_builder(
    data_column: int,
    value_name: str,
    wrapped_name: str,
) -> Any:
    data_col_letter = _openpyxl_column_letter(data_column)

    def build(row_index: int) -> str:
        return _raw_filter_match_formula(
            data_cell=f"{data_col_letter}{row_index}",
            value_name=value_name,
            wrapped_name=wrapped_name,
        )

    return build


def _stream_product_builder(match_columns: list[int]) -> Any:
    col_letters = [_openpyxl_column_letter(col_index) for col_index in match_columns]

    def build(row_index: int) -> str:
        if not col_letters:
            return "=1"
        refs = [f"{col_letter}{row_index}" for col_letter in col_letters]
        return "=" + "*".join(refs + ["1"])

    return build


def _stream_custom_match_builder(prefix: str) -> Any:
    return lambda row_index: _raw_custom_filter_match_formula(
        row_index=row_index,
        question_name=f"{prefix}_Q",
        value_name=f"{prefix}_V",
        wrapped_name=f"{prefix}_wrapped",
        resolved_column_name=f"{prefix}_resolved_column",
        criteria_name=f"{prefix}_resolved_criteria",
    )


def _stream_per_question_match_builder(q_filter_prefix: str, slot_count: int) -> Any:
    def build(row_index: int) -> str:
        formulas = []
        for slot in range(1, max(1, int(slot_count)) + 1):
            question_name, value_name = _per_question_slot_names(q_filter_prefix, slot)
            slot_prefix = question_name[:-2] if question_name.endswith("_Q") else question_name
            formulas.append(
                _raw_custom_filter_match_formula(
                    row_index=row_index,
                    question_name=question_name,
                    value_name=value_name,
                    wrapped_name=f"{slot_prefix}_wrapped",
                    resolved_column_name=f"{slot_prefix}_resolved_column",
                    criteria_name=f"{slot_prefix}_resolved_criteria",
                ).lstrip("=")
            )
        return "=" + "*".join(formulas + ["1"])

    return build


def _build_theme_local_filter_helper_column(
    workbook: Any,
    theme_prefix: str,
    local_filters: list[dict[str, str]],
    context: _LiveWorkbookContext,
) -> str:
    """Append a per-theme local-filter pass mask to _RawData."""

    if context.raw_data_streamed:
        return _theme_local_pass_data_name(theme_prefix)

    worksheet = workbook["_RawData"]
    last_row = max(2, worksheet.max_row)
    next_col = worksheet.max_column + 1
    match_columns: list[int] = []

    for local_filter in local_filters:
        data_name = local_filter.get("data_name", "")
        data_column = _raw_data_column_index_for_name(context, data_name)
        if data_column is None:
            continue

        match_name = _safe_defined_name(
            f"{theme_prefix}_{local_filter['filter_name']}_local_match_data"
        )
        worksheet.cell(row=1, column=next_col, value=match_name.removesuffix("_data"))
        for row_index in range(2, last_row + 1):
            data_cell = f"{_openpyxl_column_letter(data_column)}{row_index}"
            _live_formula(
                worksheet,
                row_index,
                next_col,
                _raw_filter_match_formula(
                    data_cell=data_cell,
                    value_name=local_filter["value_name"],
                    wrapped_name=local_filter["wrapped_name"],
                ),
                1,
            )
        col_letter = _openpyxl_column_letter(next_col)
        _add_named_range(
            workbook,
            match_name,
            "_RawData",
            f"${col_letter}$2:${col_letter}${last_row}",
        )
        match_columns.append(next_col)
        next_col += 1

    pass_name = _theme_local_pass_data_name(theme_prefix)
    worksheet.cell(row=1, column=next_col, value=pass_name.removesuffix("_data"))
    for row_index in range(2, last_row + 1):
        if match_columns:
            refs = [
                f"{_openpyxl_column_letter(col_index)}{row_index}"
                for col_index in match_columns
            ]
            formula = "=" + "*".join(refs + ["1"])
        else:
            formula = "=1"
        _live_formula(worksheet, row_index, next_col, formula, 1)

    pass_col_letter = _openpyxl_column_letter(next_col)
    _add_named_range(
        workbook,
        pass_name,
        "_RawData",
        f"${pass_col_letter}$2:${pass_col_letter}${last_row}",
    )
    return pass_name


def _build_per_question_filter_helper_column(
    workbook: Any,
    q_filter_prefix: str,
    context: _LiveWorkbookContext,
    slot_count: int = DEFAULT_PER_QUESTION_FILTER_COUNT,
) -> str:
    """Append a per-question pass mask to _RawData for COUNTIFS filtering."""

    if context.raw_data_streamed:
        return _per_question_pass_data_name(f"{q_filter_prefix}_Q")

    worksheet = workbook["_RawData"]
    last_row = max(2, worksheet.max_row)
    next_col = worksheet.max_column + 1
    pass_name = _per_question_pass_data_name(f"{q_filter_prefix}_Q")
    worksheet.cell(row=1, column=next_col, value=pass_name.removesuffix("_data"))
    for row_index in range(2, last_row + 1):
        formulas = []
        for slot in range(1, max(1, int(slot_count)) + 1):
            question_name, value_name = _per_question_slot_names(q_filter_prefix, slot)
            slot_prefix = question_name[:-2] if question_name.endswith("_Q") else question_name
            formulas.append(
                _raw_custom_filter_match_formula(
                    row_index=row_index,
                    question_name=question_name,
                    value_name=value_name,
                    wrapped_name=f"{slot_prefix}_wrapped",
                    resolved_column_name=f"{slot_prefix}_resolved_column",
                    criteria_name=f"{slot_prefix}_resolved_criteria",
                ).lstrip("=")
            )
        _live_formula(
            worksheet,
            row_index,
            next_col,
            "=" + "*".join(formulas + ["1"]),
            1,
        )

    col_letter = _openpyxl_column_letter(next_col)
    _add_named_range(
        workbook,
        pass_name,
        "_RawData",
        f"${col_letter}$2:${col_letter}${last_row}",
    )
    return pass_name


def _live_write_run_summary(
    workbook: Any,
    schema: SurveySchema,
    quality_report: DataQualityReport,
    results: list[SingleCutResult],
    skips: list[SkipRecord],
    log: CalculationLog,
    cross_cut_results: list[CrossCutResult],
    cross_cut_skips: list[SkipRecord],
) -> None:
    worksheet = workbook.create_sheet("Run_Summary")
    rows = [
        ("Source datamap:", schema.source_datamap_path),
        ("Source raw data:", schema.source_rawdata_path),
        ("Run timestamp:", schema.parsed_at.isoformat()),
        ("Total respondents:", schema.total_respondents),
        ("Total questions:", len(schema.questions)),
        ("Results produced:", len(results)),
        ("Questions skipped:", len(skips)),
        ("Cross cuts produced:", len(cross_cut_results)),
        ("Cross cut skips:", len(cross_cut_skips)),
        ("Audit log records:", len(log)),
        ("Quality warnings:", len(quality_report.warnings)),
    ]
    for row_index, (label, value) in enumerate(rows, start=1):
        worksheet.cell(row=row_index, column=1, value=label).font = _live_font(bold=True)
        worksheet.cell(row=row_index, column=2, value=value)
    _live_autofit(worksheet)


def _live_write_question_metadata(workbook: Any, schema: SurveySchema) -> None:
    worksheet = workbook.create_sheet("Question_Metadata")
    headers = [
        "Question ID",
        "Canonical ID",
        "Question Text",
        "Type",
        "Raw Columns",
        "Options Count",
        "Analysis Eligible",
        "Demographic",
        "Parent Question",
    ]
    _live_header_row(worksheet, 1, headers)
    for row_index, spec in enumerate(schema.questions, start=2):
        values = [
            spec.question_id,
            spec.canonical_id,
            spec.question_text,
            spec.question_type.value,
            ", ".join(spec.raw_columns),
            len(spec.option_map),
            "Yes" if spec.analysis_eligible else "No",
            "Yes" if spec.is_demographic else "No",
            spec.parent_question_id or "",
        ]
        for col_index, value in enumerate(values, start=1):
            worksheet.cell(row=row_index, column=col_index, value=value)
    worksheet.freeze_panes = "A2"
    _live_autofit(worksheet)


def _live_write_single_cut_index(
    workbook: Any,
    results: list[SingleCutResult],
    result_sheet_names: dict[str, str],
    result_theme_names: dict[str, str],
) -> None:
    worksheet = workbook.create_sheet("Single_Cut_Index")
    headers = [
        "Question ID",
        "Canonical ID",
        "Type",
        "Theme",
        "Sheet Name",
        "Valid N",
        "Missing N",
        "Missing %",
        "Warnings",
    ]
    _live_header_row(worksheet, 1, headers)
    for row_index, result in enumerate(results, start=2):
        denominator = result.valid_n + result.missing_n
        missing_pct = result.missing_n / denominator if denominator else 0.0
        sheet_name = result_sheet_names.get(result.question_id, "")
        values = [
            result.question_id,
            result.question_id,
            result.question_type.value,
            result_theme_names.get(result.question_id, ""),
            sheet_name,
            result.valid_n,
            result.missing_n,
            missing_pct,
            " | ".join(result.warnings),
        ]
        for col_index, value in enumerate(values, start=1):
            cell = worksheet.cell(row=row_index, column=col_index, value=value)
            if col_index == 5 and sheet_name:
                cell.hyperlink = f"#{_quote_openpyxl_sheet(sheet_name)}!A1"
                cell.style = "Hyperlink"
            if col_index == 8:
                cell.number_format = "0.0%"
    worksheet.freeze_panes = "A2"
    _live_autofit(worksheet)


def _live_write_calculation_log(workbook: Any, log: CalculationLog) -> None:
    worksheet = workbook.create_sheet("Calculation_Log")
    headers = [
        "Output Sheet",
        "Metric Name",
        "Source Question",
        "Source Columns",
        "Filter",
        "Numerator",
        "Denominator",
        "Formula",
        "Value Raw",
        "Valid N",
        "Missing N",
        "Timestamp",
    ]
    _live_header_row(worksheet, 1, headers)
    records = log.all_records() if hasattr(log, "all_records") else tuple(log)
    for row_index, record in enumerate(records, start=2):
        values = [
            record.output_sheet,
            record.metric_name,
            record.source_question_id,
            ", ".join(record.source_columns),
            record.filter_expr or "",
            record.numerator,
            record.denominator,
            record.formula,
            record.value_raw,
            record.valid_n,
            record.missing_n,
            record.timestamp.isoformat(),
        ]
        for col_index, value in enumerate(values, start=1):
            worksheet.cell(row=row_index, column=col_index, value=value)
    worksheet.freeze_panes = "A2"
    _live_autofit(worksheet)


def _live_write_hypothesis_check_sheet(
    workbook: Any,
    schema: SurveySchema,
    hypothesis_results: list[HypothesisResult],
    log: CalculationLog,
) -> None:
    if not hypothesis_results:
        return

    worksheet = workbook.create_sheet("Hypothesis_Check")
    headers = [
        "Hypothesis title",
        "Verdict",
        "Verdict reason code",
        "Outcome question ID",
        "Outcome question text",
        "Predictor question ID",
        "Predictor question text",
        "Test used",
        "Effect size",
        "Effect size label",
        "p-value",
        "95% CI low",
        "95% CI high",
        "Sample size",
        "Cohort n",
        "Active filters summary",
        "Related cross-cut sheet",
        "AI prose explanation",
        "Warnings",
    ]
    _live_header_row(worksheet, 1, headers)
    log_row_by_metric = {
        record.metric_name: row_index
        for row_index, record in enumerate(log.all_records(), start=2)
    }
    for row_index, result in enumerate(hypothesis_results, start=2):
        outcome = schema.get_question(result.spec.outcome_question_id)
        predictor = schema.get_question(result.spec.predictor_question_id)
        statistic = result.statistic
        effect_metric = _hypothesis_metric_id(result, "effect_size")
        p_metric = _hypothesis_metric_id(result, "p_value")
        ci_low_metric = _hypothesis_metric_id(result, "ci_low")
        ci_high_metric = _hypothesis_metric_id(result, "ci_high")
        values = [
            result.spec.title,
            result.verdict,
            result.verdict_reason,
            result.spec.outcome_question_id,
            outcome.question_text if outcome is not None else "",
            result.spec.predictor_question_id,
            predictor.question_text if predictor is not None else "",
            statistic.test_name if statistic is not None else "",
            None,
            statistic.effect_size_label if statistic is not None else "",
            None,
            None,
            None,
            result.pairwise_n,
            result.cohort_n,
            result.active_filters_summary,
            result.related_cross_cut_id or "",
            result.ai_prose_explanation,
            ", ".join(result.warnings),
        ]
        for col_index, value in enumerate(values, start=1):
            worksheet.cell(row=row_index, column=col_index, value=value)
        if statistic is not None:
            _write_hypothesis_formula_reference(
                worksheet,
                row_index,
                9,
                effect_metric,
                log_row_by_metric,
                statistic.effect_size,
            )
            _write_hypothesis_formula_reference(
                worksheet,
                row_index,
                11,
                p_metric,
                log_row_by_metric,
                statistic.p_value,
            )
            if statistic.confidence_interval_low is not None:
                _write_hypothesis_formula_reference(
                    worksheet,
                    row_index,
                    12,
                    ci_low_metric,
                    log_row_by_metric,
                    statistic.confidence_interval_low,
                )
            if statistic.confidence_interval_high is not None:
                _write_hypothesis_formula_reference(
                    worksheet,
                    row_index,
                    13,
                    ci_high_metric,
                    log_row_by_metric,
                    statistic.confidence_interval_high,
                )
        if result.related_cross_cut_id:
            cell = worksheet.cell(row=row_index, column=17)
            cell.hyperlink = f"#{_quote_openpyxl_sheet(result.related_cross_cut_id)}!A1"
            cell.style = "Hyperlink"
    worksheet.freeze_panes = "A2"
    _live_autofit(worksheet)


def _hypothesis_metric_id(result: HypothesisResult, suffix: str) -> str | None:
    for entry_id in result.calculation_log_entry_ids:
        if entry_id.endswith(f":{suffix}"):
            return entry_id
    return None


def _write_hypothesis_formula_reference(
    worksheet: Any,
    row_index: int,
    column_index: int,
    metric_id: str | None,
    log_row_by_metric: dict[str, int],
    cached_value: float,
) -> None:
    log_row = log_row_by_metric.get(metric_id or "")
    if log_row is None:
        worksheet.cell(row=row_index, column=column_index, value=cached_value)
        return
    _live_formula(
        worksheet,
        row_index,
        column_index,
        f"='Calculation_Log'!I{log_row}",
        cached_value,
    )


def _live_write_filter_log(
    workbook: Any,
    cross_cut_results: list[CrossCutResult],
) -> None:
    worksheet = workbook.create_sheet("Filter_Log")
    headers = ["Cross Cut ID", "Title", "Filter Expression", "Description"]
    _live_header_row(worksheet, 1, headers)
    row_index = 2
    for result in cross_cut_results:
        filter_expr = result.result_table.get("filter_expr")
        if not filter_expr:
            continue
        values = [
            result.cross_cut_id,
            result.synthetic_question_title,
            filter_expr,
            result.result_table.get("filter_mask_description") or filter_expr,
        ]
        for col_index, value in enumerate(values, start=1):
            worksheet.cell(row=row_index, column=col_index, value=value)
        row_index += 1
    worksheet.freeze_panes = "A2"
    _live_autofit(worksheet)


def _live_write_warnings(
    workbook: Any,
    schema: SurveySchema,
    quality_report: DataQualityReport,
    results: list[SingleCutResult],
    skips: list[SkipRecord],
    cross_cut_results: list[CrossCutResult],
    cross_cut_skips: list[SkipRecord],
) -> None:
    worksheet = workbook.create_sheet("Warnings")
    _live_header_row(worksheet, 1, ["Source", "Warning"])
    row_index = 2
    for warning in quality_report.warnings:
        worksheet.cell(row=row_index, column=1, value="quality")
        worksheet.cell(row=row_index, column=2, value=warning)
        row_index += 1
    for result in results:
        for warning in result.warnings:
            worksheet.cell(row=row_index, column=1, value=f"result:{result.question_id}")
            worksheet.cell(row=row_index, column=2, value=warning)
            row_index += 1
    for result in cross_cut_results:
        for warning in result.warnings:
            worksheet.cell(row=row_index, column=1, value=f"cross_result:{result.cross_cut_id}")
            worksheet.cell(row=row_index, column=2, value=warning)
            row_index += 1
    for skip in (*skips, *cross_cut_skips):
        if skip.details:
            worksheet.cell(row=row_index, column=1, value=f"skip:{skip.question_id}")
            worksheet.cell(row=row_index, column=2, value=skip.details)
            row_index += 1
    low_confidence_questions = [
        question
        for question in schema.questions
        if getattr(question, "classification_confidence_low", False)
    ]
    if low_confidence_questions:
        row_index += 1
        headers = ["Question ID", "Question Text", "Type", "Possible Role", "Reason"]
        for col_index, header in enumerate(headers, start=1):
            worksheet.cell(row=row_index, column=col_index, value=header).font = _live_font(bold=True)
        row_index += 1
        for question in low_confidence_questions:
            values = [
                question.canonical_id,
                question.question_text,
                question.question_type.value,
                question.possible_role or "",
                "low classification confidence",
            ]
            for col_index, value in enumerate(values, start=1):
                worksheet.cell(row=row_index, column=col_index, value=value)
            row_index += 1
    decoder_warnings = tuple(getattr(quality_report, "decoder_warnings", ()) or ())
    if decoder_warnings:
        row_index = _live_write_decoder_warnings_section(
            worksheet,
            row_index + 1,
            decoder_warnings,
        )
    _live_autofit(worksheet)


def _live_write_decoder_warnings_section(
    worksheet: Any,
    row_index: int,
    decoder_warnings: tuple[dict, ...],
) -> int:
    worksheet.cell(row=row_index, column=1, value="Data Quality — Decoder Warnings").font = _live_font(bold=True)
    row_index += 1

    summary_counts: dict[tuple[str, str], int] = defaultdict(int)
    for warning in decoder_warnings:
        question_id = str(warning.get("question_id", ""))
        action = str(warning.get("action", ""))
        summary_counts[(question_id, action)] += 1

    for col_index, header in enumerate(["question_id", "action", "count"], start=1):
        worksheet.cell(row=row_index, column=col_index, value=header).font = _live_font(bold=True)
    row_index += 1
    for (question_id, action), count in sorted(summary_counts.items()):
        worksheet.cell(row=row_index, column=1, value=question_id)
        worksheet.cell(row=row_index, column=2, value=action)
        worksheet.cell(row=row_index, column=3, value=count)
        row_index += 1

    row_index += 1
    detail_headers = ["question_id", "column", "row", "raw_value", "action"]
    for col_index, header in enumerate(detail_headers, start=1):
        worksheet.cell(row=row_index, column=col_index, value=header).font = _live_font(bold=True)
    row_index += 1
    for warning in decoder_warnings:
        values = [
            warning.get("question_id", ""),
            warning.get("column", ""),
            warning.get("row", ""),
            warning.get("raw_value", ""),
            warning.get("action", ""),
        ]
        for col_index, value in enumerate(values, start=1):
            worksheet.cell(row=row_index, column=col_index, value=value)
        row_index += 1
    return row_index


def _live_write_theme_sheet(
    workbook: Any,
    theme_name: str,
    sheet_name: str,
    results: list[SingleCutResult],
    schema: SurveySchema,
    context: _LiveWorkbookContext,
    sheet_filters: list[dict[str, str]],
    short_labels: dict[str, str],
) -> None:
    worksheet = workbook.create_sheet(sheet_name)
    row_index = _live_write_theme_header(worksheet, theme_name)
    theme_prefix = _theme_defined_prefix(sheet_name)
    row_index, local_filters = _build_theme_local_filters_block(
        workbook,
        worksheet,
        row_index,
        theme_prefix,
        sheet_filters,
        context,
    )
    _build_theme_local_filter_helper_column(
        workbook,
        theme_prefix,
        local_filters,
        context,
    )
    for result in results:
        question = schema.get_question(result.question_id)
        if question is None:
            continue
        row_index = _live_write_question_block(
            workbook,
            worksheet,
            row_index,
            result,
            question,
            context,
            local_filters,
            theme_prefix,
            short_labels,
            schema,
        )
    _live_set_theme_column_widths(worksheet)
    worksheet.freeze_panes = "A4"


def _live_write_nps_sheet(
    workbook: Any,
    sheet_name: str,
    result: NPSResult,
    question: Any,
    context: _LiveWorkbookContext,
    sheet_filters: list[dict[str, str]],
) -> None:
    worksheet = workbook.create_sheet(sheet_name)
    heading_fill = _live_fill("FFCC0000")
    heading_font = _live_font(bold=True, size=12, color="FFFFFFFF")
    for col_index in range(1, 10):
        cell = worksheet.cell(row=1, column=col_index)
        cell.fill = heading_fill
        cell.font = heading_font
    worksheet.cell(
        row=1,
        column=1,
        value=f"{question.canonical_id} - {question.question_text}",
    )
    worksheet.cell(
        row=2,
        column=1,
        value="Net Promoter Score by entity",
    ).font = _live_font(italic=True, size=9, color="666666")
    ct_name = f"{_safe_defined_name(sheet_name)}_{_safe_defined_name(question.canonical_id)}_CT"
    worksheet.cell(row=3, column=1, value="Cross-tab by").font = _live_font(bold=True, size=9)
    ct_cell = worksheet.cell(row=3, column=3, value="(None)")
    _add_named_cell(workbook, ct_name, worksheet, ct_cell.coordinate)
    _add_dropdown_to_cell(worksheet, ct_cell.coordinate, "=Cross_Tab_Questions")
    table_end_row = _live_write_nps_table(
        worksheet,
        4,
        result,
        question,
        context,
        sheet_filters,
        "",
        "",
        "",
    )
    cross_tab_rows = _live_distribution_rows(result, question, context)
    if cross_tab_rows:
        _live_write_cross_tab_table(
            worksheet,
            table_end_row + 2,
            cross_tab_rows,
            sheet_filters,
            ct_name,
            "",
            "",
            "",
            context,
        )
    worksheet.freeze_panes = "A5"
    _live_autofit(worksheet)


def _live_write_theme_header(
    worksheet: Any,
    theme_name: str,
) -> int:
    from openpyxl.styles import Alignment

    worksheet.merge_cells(start_row=1, start_column=1, end_row=1, end_column=10)
    banner = worksheet.cell(row=1, column=1, value=f"THEME: {theme_name}")
    banner.font = _live_font(bold=True, size=12, color="FFFFFF")
    banner.fill = _live_fill("CC0000")
    banner.alignment = Alignment(vertical="center")
    for col_index in range(1, 11):
        worksheet.cell(row=1, column=col_index).fill = _live_fill("CC0000")

    worksheet.cell(
        row=2,
        column=1,
        value="Tip: Use the Filters sheet for workbook-wide values, or override below.",
    ).font = _live_font(italic=True, size=9, color="666666")
    return 3


def _build_theme_local_filters_block(
    workbook: Any,
    worksheet: Any,
    start_row: int,
    theme_prefix: str,
    workbook_filters: list[dict[str, str]],
    context: _LiveWorkbookContext,
) -> tuple[int, list[dict[str, str]]]:
    worksheet.cell(
        row=start_row,
        column=1,
        value="LOCAL FILTERS (override workbook defaults)",
    ).font = _live_font(bold=True)
    _live_fill_row(worksheet, start_row, 8, "F2F2F2")
    row_index = start_row + 1
    _live_header_row(
        worksheet,
        row_index,
        ["Filter", "Value", "Resolved", "Wrapped", "Available values"],
    )
    row_index += 1

    local_filters: list[dict[str, str]] = []
    for sheet_filter in workbook_filters:
        global_name = sheet_filter["filter_name"]
        local_name = f"{theme_prefix}_{global_name}"
        resolved_name = f"{local_name}_resolved"
        wrapped_name = f"{local_name}_wrapped"
        worksheet.cell(row=row_index, column=1, value=sheet_filter["label"])
        if sheet_filter.get("question_text"):
            _add_comment(worksheet.cell(row=row_index, column=1), sheet_filter["question_text"])
        value_cell = worksheet.cell(row=row_index, column=2, value="(Inherit)")
        resolved_cell = _live_formula(
            worksheet,
            row_index,
            3,
            f'=IF({local_name}="(Inherit)",{global_name},{local_name})',
            "(All)",
        )
        wrapped_cell = _live_formula(
            worksheet,
            row_index,
            4,
            _wrapped_formula(resolved_name),
            "|(All)|",
        )
        worksheet.cell(row=row_index, column=5, value=_available_values_for_data_name(sheet_filter["data_name"], context))
        _add_named_cell(workbook, local_name, worksheet, value_cell.coordinate)
        _add_named_cell(workbook, resolved_name, worksheet, resolved_cell.coordinate)
        _add_named_cell(workbook, wrapped_name, worksheet, wrapped_cell.coordinate)
        _add_dropdown_to_cell(worksheet, value_cell.coordinate, f'={sheet_filter["options_name"]}')
        local_filters.append(
            {
                **sheet_filter,
                "value_name": resolved_name,
                "wrapped_name": wrapped_name,
            }
        )
        row_index += 1

    for slot in range(1, context.workbook_custom_filter_count + 1):
        row_index = _write_theme_custom_filter(
            workbook,
            worksheet,
            row_index,
            theme_prefix,
            f"Custom {slot}",
            f"F_Custom{slot}_Q",
            f"F_Custom{slot}_V",
        )
    return row_index + 2, local_filters


def _write_theme_custom_filter(
    workbook: Any,
    worksheet: Any,
    row_index: int,
    theme_prefix: str,
    label: str,
    global_question_name: str,
    global_value_name: str,
) -> int:
    local_q = f"{theme_prefix}_{global_question_name}"
    local_v = f"{theme_prefix}_{global_value_name}"
    resolved_q = f"{local_q}_resolved"
    resolved_v = f"{local_v}_resolved"
    resolved_column = f"{theme_prefix}_{global_question_name[:-2]}_resolved_column"
    wrapped_name = f"{theme_prefix}_{global_question_name[:-2]}_wrapped"

    worksheet.cell(row=row_index, column=1, value=f"{label} question")
    q_cell = worksheet.cell(row=row_index, column=2, value="(Inherit)")
    q_resolved_cell = _live_formula(
        worksheet,
        row_index,
        3,
        f'=IF({local_q}="(Inherit)",{global_question_name},{local_q})',
        "(None)",
    )
    worksheet.cell(row=row_index, column=5, value="Workbook question or local override")
    _add_named_cell(workbook, local_q, worksheet, q_cell.coordinate)
    _add_named_cell(workbook, resolved_q, worksheet, q_resolved_cell.coordinate)
    _add_dropdown_to_cell(worksheet, q_cell.coordinate, "=All_Questions_Local")
    row_index += 1

    worksheet.cell(row=row_index, column=1, value=f"{label} value")
    v_cell = worksheet.cell(row=row_index, column=2, value="(Inherit)")
    v_resolved_cell = _live_formula(
        worksheet,
        row_index,
        3,
        f'=IF({local_v}="(Inherit)",{global_value_name},{local_v})',
        "(All)",
    )
    wrapped_cell = _live_formula(
        worksheet,
        row_index,
        4,
        _wrapped_formula(resolved_v),
        "|(All)|",
    )
    worksheet.cell(row=row_index, column=5, value="Choose one value or type comma-separated values")
    resolved_column_cell = _live_formula(
        worksheet,
        row_index,
        6,
        f"={_question_data_name_formula(resolved_q, resolved_v)}",
        "respondent_id_data",
    )
    resolved_criteria_cell = _live_formula(
        worksheet,
        row_index,
        7,
        f"={_question_filter_criteria_formula(resolved_q, resolved_v)}",
        "(All)",
    )
    _add_named_cell(workbook, local_v, worksheet, v_cell.coordinate)
    _add_named_cell(workbook, resolved_v, worksheet, v_resolved_cell.coordinate)
    _add_named_cell(workbook, wrapped_name, worksheet, wrapped_cell.coordinate)
    _add_named_cell(workbook, resolved_column, worksheet, resolved_column_cell.coordinate)
    _add_named_cell(workbook, f"{theme_prefix}_{global_question_name[:-2]}_resolved_criteria", worksheet, resolved_criteria_cell.coordinate)
    _add_dropdown_to_cell(
        worksheet,
        v_cell.coordinate,
        f"=INDIRECT({_question_options_name_formula(resolved_q)})",
    )
    return row_index + 1


def _live_write_custom_filter_slot(
    workbook: Any,
    worksheet: Any,
    row_index: int,
    col_index: int,
    question_name: str,
    value_name: str,
    slot: int | None = None,
) -> None:
    prefix = question_name[:-2] if question_name.endswith("_Q") else question_name
    slot_label = slot if slot is not None else 1
    worksheet.cell(row=row_index, column=col_index, value=f"Filter {slot_label} question")
    q_cell = worksheet.cell(row=row_index, column=col_index + 1, value="(None)")
    worksheet.cell(row=row_index, column=col_index + 2, value="Filter value")
    v_cell = worksheet.cell(row=row_index, column=col_index + 3, value="(All)")
    resolved_column = _live_formula(
        worksheet,
        row_index,
        col_index + 4,
        f"={_question_data_name_formula(question_name, value_name)}",
        "respondent_id_data",
    )
    wrapped_cell = _live_formula(
        worksheet,
        row_index,
        col_index + 5,
        _wrapped_formula(value_name),
        "|(All)|",
    )
    resolved_criteria = _live_formula(
        worksheet,
        row_index,
        col_index + 6,
        f"={_question_filter_criteria_formula(question_name, value_name)}",
        "(All)",
    )
    _add_named_cell(workbook, question_name, worksheet, q_cell.coordinate)
    _add_named_cell(workbook, value_name, worksheet, v_cell.coordinate)
    _add_named_cell(workbook, f"{prefix}_resolved_column", worksheet, resolved_column.coordinate)
    _add_named_cell(workbook, f"{prefix}_wrapped", worksheet, wrapped_cell.coordinate)
    _add_named_cell(workbook, f"{prefix}_resolved_criteria", worksheet, resolved_criteria.coordinate)
    _add_dropdown_to_cell(worksheet, q_cell.coordinate, "=All_Questions")
    _add_dropdown_to_cell(
        worksheet,
        v_cell.coordinate,
        f"=INDIRECT({_question_options_name_formula(question_name)})",
    )


def _per_question_slot_names(q_filter_prefix: str, slot: int) -> tuple[str, str]:
    suffix = "" if slot == 1 else f"_{slot}"
    return f"{q_filter_prefix}{suffix}_Q", f"{q_filter_prefix}{suffix}_V"


def _live_write_per_question_filter_slot(
    workbook: Any,
    worksheet: Any,
    row_index: int,
    q_filter_prefix: str,
    slot: int,
) -> tuple[str, str]:
    fq_name, fv_name = _per_question_slot_names(q_filter_prefix, slot)
    label = "Per-question filter" if slot == 1 else f"Per-question filter {slot}"
    worksheet.cell(row=row_index, column=1, value=label).font = _live_font(bold=True, size=9)
    worksheet.cell(row=row_index, column=2, value="Filter Q")
    fq_cell = worksheet.cell(row=row_index, column=3, value="(None)")
    worksheet.cell(row=row_index, column=4, value="Value")
    fv_cell = worksheet.cell(row=row_index, column=5, value="(All)")
    resolved_column = _live_formula(
        worksheet,
        row_index,
        6,
        f"={_question_data_name_formula(fq_name, fv_name)}",
        "respondent_id_data",
    )
    wrapped_cell = _live_formula(
        worksheet,
        row_index,
        7,
        _wrapped_formula(fv_name),
        "|(All)|",
    )
    resolved_criteria = _live_formula(
        worksheet,
        row_index,
        8,
        f"={_question_filter_criteria_formula(fq_name, fv_name)}",
        "(All)",
    )
    slot_prefix = fq_name[:-2] if fq_name.endswith("_Q") else fq_name
    _add_named_cell(workbook, fq_name, worksheet, fq_cell.coordinate)
    _add_named_cell(workbook, fv_name, worksheet, fv_cell.coordinate)
    _add_named_cell(workbook, f"{slot_prefix}_resolved_column", worksheet, resolved_column.coordinate)
    _add_named_cell(workbook, f"{slot_prefix}_wrapped", worksheet, wrapped_cell.coordinate)
    _add_named_cell(workbook, f"{slot_prefix}_resolved_criteria", worksheet, resolved_criteria.coordinate)
    _add_dropdown_to_cell(worksheet, fq_cell.coordinate, "=All_Questions")
    _add_dropdown_to_cell(
        worksheet,
        fv_cell.coordinate,
        f"=INDIRECT({_question_options_name_formula(fq_name)})",
    )
    return fq_name, fv_name


def _live_write_question_block(
    workbook: Any,
    worksheet: Any,
    start_row: int,
    result: SingleCutResult,
    question: Any,
    context: _LiveWorkbookContext,
    sheet_filters: list[dict[str, str]],
    theme_prefix: str,
    short_labels: dict[str, str],
    schema: SurveySchema,
) -> int:
    from openpyxl.styles import Border, Side

    side = Side(style="thin", color="BFBFBF")
    for col_index in range(1, 11):
        worksheet.cell(row=start_row, column=col_index).border = Border(bottom=side)
    row_index = start_row + 1

    heading_fill = _live_fill("FFCC0000")
    heading_font = _live_font(bold=True, size=11, color="FFFFFFFF")
    heading_last_col = 14
    for col_index in range(1, heading_last_col + 1):
        cell = worksheet.cell(row=row_index, column=col_index)
        cell.fill = heading_fill
        cell.font = heading_font
    heading_text = _question_heading_text(question, short_labels)
    title_cell = worksheet.cell(
        row=row_index,
        column=1,
        value=heading_text,
    )
    title_cell.font = heading_font
    title_cell.fill = heading_fill
    row_index += 1

    q_prefix = f"{theme_prefix}_{_safe_defined_name(question.canonical_id)}"
    q_filter_prefix = f"{q_prefix}_F"
    fq_name, fv_name = _per_question_slot_names(q_filter_prefix, 1)
    ct_name = f"{q_prefix}_CT"

    if context.per_question_filter_count > 0:
        for slot in range(1, context.per_question_filter_count + 1):
            slot_fq_name, slot_fv_name = _live_write_per_question_filter_slot(
                workbook,
                worksheet,
                row_index,
                q_filter_prefix,
                slot,
            )
            if slot == 1:
                fq_name, fv_name = slot_fq_name, slot_fv_name
            row_index += 1
        _build_per_question_filter_helper_column(
            workbook,
            q_filter_prefix,
            context,
            context.per_question_filter_count,
        )
    else:
        fq_name = ""
        fv_name = ""
    cross_tab_enabled = not _question_is_grid_target(question)
    if cross_tab_enabled:
        worksheet.cell(row=row_index, column=1, value="Cross-tab by").font = _live_font(bold=True, size=9)
        ct_cell = worksheet.cell(row=row_index, column=3, value="(None)")
        _add_named_cell(workbook, ct_name, worksheet, ct_cell.coordinate)
        _add_dropdown_to_cell(worksheet, ct_cell.coordinate, "=Cross_Tab_Questions")
        row_index += 1
    rank_points_name = ""
    if isinstance(result, RankOrderResult):
        rank_points_name = _safe_defined_name(f"{q_prefix}_PTS")
        row_index = _live_write_rank_points_block(
            workbook,
            worksheet,
            row_index,
            rank_points_name,
            result.K,
            context,
            result.question_id,
        )
    note = _subset_denominator_note(result, question, schema)
    if note:
        note_cell = worksheet.cell(row=row_index, column=1, value=note)
        note_cell.font = _live_font(italic=True, size=9, color="808080")
        row_index += 1
    row_index += 1
    distribution_rows = _live_distribution_rows(result, question, context)
    cross_tab_rows = distribution_rows if cross_tab_enabled else []
    if isinstance(result, RankOrderResult):
        next_row = _live_write_rank_order_table(
            worksheet,
            row_index,
            result,
            question,
            context,
            sheet_filters,
            fq_name,
            fv_name,
            theme_prefix,
            rank_points_name,
        )
    elif isinstance(result, GridRatedResult):
        next_row = _live_write_grid_rated_result_table(
            worksheet,
            row_index,
            result,
            question,
            context,
            sheet_filters,
            fq_name,
            fv_name,
            theme_prefix,
        )
    elif isinstance(result, GridBinaryPivotResult):
        next_row = _live_write_grid_binary_pivot_result_table(worksheet, row_index, result)
    elif isinstance(result, NPSResult):
        next_row = _live_write_nps_table(
            worksheet,
            row_index,
            result,
            question,
            context,
            sheet_filters,
            fq_name,
            fv_name,
            theme_prefix,
        )
    elif isinstance(result, NumericResult):
        next_row = _live_write_numeric_table(
            worksheet,
            row_index,
            result,
            question,
            context,
            sheet_filters,
            fq_name,
            fv_name,
            theme_prefix,
        )
    elif isinstance(result, GridSingleSelectResult) and _grid_render_subtype(question, result) == GRID_RATED:
        next_row = _live_write_grid_rated_table(
            worksheet,
            row_index,
            result,
            question,
            context,
            sheet_filters,
            fq_name,
            fv_name,
            theme_prefix,
        )
    elif isinstance(result, GridSingleSelectResult) and _grid_render_subtype(question, result) == GRID_CATEGORICAL:
        next_row = _live_write_grid_categorical_table(
            worksheet,
            row_index,
            result,
            question,
            context,
            sheet_filters,
            fq_name,
            fv_name,
            theme_prefix,
        )
    else:
        rows = distribution_rows
        total_respondents = _categorical_total_respondents(result)
        total_responses = _categorical_total_responses(result)
        next_row = _live_write_categorical_table(
            worksheet,
            row_index,
            rows,
            sheet_filters,
            fq_name,
            fv_name,
            theme_prefix,
            total_respondents=total_respondents,
            total_responses=total_responses,
        )
    if cross_tab_rows:
        cross_tab_end_row = _live_write_cross_tab_table(
            worksheet,
            next_row + 2,
            cross_tab_rows,
            sheet_filters,
            ct_name,
            fq_name,
            fv_name,
            theme_prefix,
            context,
        )
        next_row = max(next_row, cross_tab_end_row)
    return next_row + 2


def _question_is_grid_target(question: Any | None) -> bool:
    return getattr(question, "question_type", None) in {
        QuestionType.GRID_RATED,
        QuestionType.GRID_SINGLE_SELECT,
        QuestionType.GRID_BINARY_SELECT,
    }


def _grid_spec_subtype(question: Any | None) -> str:
    role = str(getattr(question, "possible_role", "") or "").upper()
    if role in {GRID_RATED, GRID_CATEGORICAL, GRID_BINARY_SELECT}:
        return role
    if question is not None:
        try:
            from src.question_classifier import classify_grid_subtype

            subtype, _confidence = classify_grid_subtype(
                question,
                set(getattr(question, "raw_columns", ()) or ()),
            )
            if subtype in {GRID_RATED, GRID_CATEGORICAL, GRID_BINARY_SELECT}:
                return subtype
        except Exception:
            pass
    return GRID_CATEGORICAL


def _grid_render_subtype(question: Any, result: GridSingleSelectResult) -> str:
    subtype = _grid_spec_subtype(question)
    if subtype != GRID_BINARY_SELECT:
        return subtype
    if _grid_result_values_look_binary(result):
        return GRID_BINARY_SELECT
    return GRID_BINARY_SELECT


def _cross_tab_can_render_side_by_side(result: SingleCutResult, question: Any) -> bool:
    del question
    return isinstance(result, (SingleSelectResult, MultiSelectResult))


def _grid_result_values_look_binary(result: GridSingleSelectResult) -> bool:
    values: list[Any] = []
    for row_result in result.rows.values():
        values.extend(row_result.distribution.keys())
        values.extend(
            payload.get("label")
            for payload in row_result.distribution.values()
            if isinstance(payload, dict)
        )
    if not values:
        return True
    text = " ".join(str(value).strip().lower() for value in values)
    return any(token in text for token in ("selected", "checked", "unchecked"))


def _live_write_rank_points_block(
    workbook: Any,
    worksheet: Any,
    start_row: int,
    points_name: str,
    rank_k: int,
    context: _LiveWorkbookContext,
    question_id: str,
) -> int:
    settings = _normalise_rank_cross_tab_settings(
        context.rank_cross_tab_settings.get(question_id),
        rank_k,
    )
    points = settings["points"]
    worksheet.cell(row=start_row, column=1, value="Rank points").font = _live_font(bold=True, size=9)
    worksheet.cell(row=start_row, column=2, value="Rank").font = _live_font(bold=True, size=9)
    worksheet.cell(row=start_row, column=3, value="Points").font = _live_font(bold=True, size=9)
    point_start_row = start_row + 1
    for offset, point_value in enumerate(points):
        row_index = point_start_row + offset
        worksheet.cell(row=row_index, column=2, value=f"Rank {offset + 1}")
        point_cell = worksheet.cell(row=row_index, column=3, value=point_value)
        point_cell.number_format = "0.00"
    _add_named_range(
        workbook,
        points_name,
        worksheet.title,
        f"$C${point_start_row}:$C${point_start_row + rank_k - 1}",
    )
    for row_index in range(start_row, point_start_row + rank_k):
        for col_index in range(1, 4):
            worksheet.cell(row=row_index, column=col_index).fill = _live_fill("F7F7F7")
    return point_start_row + rank_k


def _live_write_rank_order_table(
    worksheet: Any,
    start_row: int,
    result: RankOrderResult,
    question: Any,
    context: _LiveWorkbookContext,
    sheet_filters: list[dict[str, str]],
    fq_name: str,
    fv_name: str,
    theme_prefix: str,
    rank_points_name: str,
) -> int:
    del question
    header_row = start_row
    subheader_row = start_row + 1
    worksheet.cell(row=header_row, column=1, value="Option ID").font = _live_font(bold=True)
    worksheet.cell(row=header_row, column=2, value="Option").font = _live_font(bold=True)
    worksheet.cell(row=header_row, column=3, value="Weighted Average").font = _live_font(bold=True)
    worksheet.cell(row=subheader_row, column=3, value="0-K").font = _live_font(bold=True)
    col_index = 4
    percent_columns: list[int] = []
    for rank in range(1, result.K + 1):
        worksheet.cell(row=header_row, column=col_index, value=str(rank)).font = _live_font(bold=True)
        worksheet.merge_cells(
            start_row=header_row,
            start_column=col_index,
            end_row=header_row,
            end_column=col_index + 1,
        )
        worksheet.cell(row=subheader_row, column=col_index, value="# of respondents").font = _live_font(bold=True)
        worksheet.cell(row=subheader_row, column=col_index + 1, value="% of respondents").font = _live_font(bold=True)
        percent_columns.append(col_index + 1)
        col_index += 2
    last_col = col_index - 1
    for row in (header_row, subheader_row):
        for col in range(1, last_col + 1):
            worksheet.cell(row=row, column=col).fill = _live_fill("F2F2F2")

    data_start = start_row + 2
    data_end = data_start + len(result.rows) - 1
    total_row = max(data_start, data_end + 1)
    count_cell_refs: list[str] = []
    denominator_data_names = [
        context.column_by_key[row.option_id].data_name
        for row in result.rows
        if row.option_id in context.column_by_key
    ]
    rank_points = _normalise_rank_cross_tab_settings(
        context.rank_cross_tab_settings.get(result.question_id),
        result.K,
    )["points"]
    for offset, row in enumerate(result.rows):
        excel_row = data_start + offset
        worksheet.cell(row=excel_row, column=1, value=row.option_id).font = _live_font(color="808080")
        worksheet.cell(row=excel_row, column=2, value=row.option_label)
        column = context.column_by_key.get(row.option_id)
        if column is None:
            raise ValueError(f"Missing live data named range for rank-order option {row.option_id!r}")
        col = 4
        row_count_refs: list[tuple[int, str]] = []
        for rank_number, (count, pct) in enumerate(zip(row.counts_per_rank, row.pcts_per_rank), start=1):
            _live_formula(
                worksheet,
                excel_row,
                col,
                _build_countifs_formula(
                    column.data_name,
                    str(rank_number),
                    sheet_filters,
                    fq_name,
                    fv_name,
                    theme_prefix=theme_prefix,
                ),
                count,
            ).number_format = "#,##0"
            count_ref = f"{_openpyxl_column_letter(col)}{excel_row}"
            row_count_refs.append((rank_number, count_ref))
            count_cell_refs.append(count_ref)
            _live_formula(
                worksheet,
                excel_row,
                col + 1,
                f'=IFERROR({count_ref}/$B${total_row},0)',
                pct,
            ).number_format = "0.0%"
            col += 2
        points_terms = []
        for rank_number, _count_ref in row_count_refs:
            count_expr = _build_countifs_formula(
                column.data_name,
                str(rank_number),
                sheet_filters,
                fq_name,
                fv_name,
                theme_prefix=theme_prefix,
            ).lstrip("=")
            points_terms.append(f"INDEX({rank_points_name},{rank_number})*{count_expr}")
        score_formula = (
            f'=IFERROR(({" + ".join(points_terms)})/$B${total_row},0)'
            if points_terms
            else "=0"
        )
        _live_formula(
            worksheet,
            excel_row,
            3,
            score_formula,
            _rank_weighted_average_from_counts(
                row.counts_per_rank,
                rank_points,
                result.total_respondents,
            ),
        ).number_format = "0.00"

    if result.rows:
        for percent_col in percent_columns:
            _apply_color_scale_range(worksheet, data_start, percent_col, data_end, percent_col)
        _apply_color_scale_range(worksheet, data_start, 3, data_end, 3)
    _write_total_respondents_row(worksheet, total_row, result.total_respondents, last_col)
    if not denominator_data_names:
        raise ValueError(f"Missing live data named range for rank-order question {result.question_id!r}")
    _live_formula(
        worksheet,
        total_row,
        2,
        "=" + _build_multi_select_respondent_count_formula(
            denominator_data_names,
            sheet_filters,
            fq_name,
            fv_name,
            theme_prefix=theme_prefix,
        ),
        result.total_respondents,
    ).number_format = "#,##0"
    responses_row = total_row + 1
    _write_total_responses_row(worksheet, responses_row, result.total_responses, last_col)
    count_sum_formula = f'=SUM({",".join(count_cell_refs)})' if count_cell_refs else "=0"
    _live_formula(
        worksheet,
        responses_row,
        2,
        count_sum_formula,
        result.total_responses,
    ).number_format = "#,##0"
    qc_row = responses_row + 1
    worksheet.cell(row=qc_row, column=2, value="QC check").font = _live_font(bold=True)
    _live_formula(
        worksheet,
        qc_row,
        3,
        f"=B{responses_row}={count_sum_formula.lstrip('=')}",
        sum(sum(row.counts_per_rank) for row in result.rows) == result.total_responses,
    )
    return qc_row


def _live_write_grid_rated_result_table(
    worksheet: Any,
    start_row: int,
    result: GridRatedResult,
    question: Any,
    context: _LiveWorkbookContext,
    sheet_filters: list[dict[str, str]],
    fq_name: str,
    fv_name: str,
    theme_prefix: str,
) -> int:
    header_row = start_row
    worksheet.cell(row=header_row, column=1, value="Sub-question ID").font = _live_font(bold=True)
    worksheet.cell(row=header_row, column=2, value="Sub-question").font = _live_font(bold=True)
    col_index = 3
    metric_columns: list[int] = []
    for header in result.column_headers:
        worksheet.cell(row=header_row, column=col_index, value=header).font = _live_font(bold=True)
        metric_columns.append(col_index)
        col_index += 1
    delta_col = None
    if result.show_delta:
        delta_col = col_index
        worksheet.cell(row=header_row, column=delta_col, value="Delta").font = _live_font(bold=True)
        col_index += 1
    last_col = col_index - 1
    for col in range(1, last_col + 1):
        worksheet.cell(row=header_row, column=col).fill = _live_fill("F2F2F2")

    data_start = start_row + 1
    row_sources = _grid_rated_result_source_columns(result, question, context)
    for offset, row in enumerate(result.rows):
        excel_row = data_start + offset
        worksheet.cell(row=excel_row, column=1, value=row.row_id).font = _live_font(color="808080")
        worksheet.cell(row=excel_row, column=2, value=row.row_label)
        col = 3
        mean_cell_refs: list[str] = []
        for column_offset, mean in enumerate(row.means_per_column):
            source_column = row_sources[(row.row_id, column_offset)]
            _live_formula(
                worksheet,
                excel_row,
                col,
                _build_averageifs_nonnegative_formula(
                    source_column.data_name,
                    fq_name,
                    theme_prefix,
                ),
                mean,
            ).number_format = "0.00"
            mean_cell_refs.append(f"{_openpyxl_column_letter(col)}{excel_row}")
            col += 1
        if delta_col is not None:
            delta_formula = (
                f"={mean_cell_refs[0]}-{mean_cell_refs[1]}"
                if len(mean_cell_refs) >= 2
                else "=0"
            )
            _live_formula(
                worksheet,
                excel_row,
                delta_col,
                delta_formula,
                row.delta or 0.0,
            ).number_format = "0.00"

    data_end = data_start + len(result.rows) - 1
    if result.rows:
        for col in metric_columns:
            _apply_color_scale_range(worksheet, data_start, col, data_end, col)
    total_row = max(data_start, data_end + 1)
    _write_total_respondents_row(worksheet, total_row, result.total_respondents, last_col)
    first_source = next(iter(row_sources.values()), None)
    if first_source is None:
        raise ValueError(f"Missing live data named range for rated grid {result.question_id!r}")
    _live_formula(
        worksheet,
        total_row,
        2,
        _build_countifs_formula(
            first_source.data_name,
            "<>",
            sheet_filters,
            fq_name,
            fv_name,
            theme_prefix=theme_prefix,
        ),
        result.total_respondents,
    ).number_format = "#,##0"
    responses_row = total_row + 1
    _write_total_responses_row(worksheet, responses_row, result.total_responses, last_col)
    response_formula_parts = [
        _build_countifs_formula(
            source_column.data_name,
            "<>",
            sheet_filters,
            fq_name,
            fv_name,
            theme_prefix=theme_prefix,
        ).lstrip("=")
        for source_column in row_sources.values()
    ]
    response_formula = f"=SUM({','.join(response_formula_parts)})" if response_formula_parts else "=0"
    _live_formula(
        worksheet,
        responses_row,
        2,
        response_formula,
        result.total_responses,
    ).number_format = "#,##0"
    qc_row = responses_row + 1
    worksheet.cell(row=qc_row, column=2, value="QC check").font = _live_font(bold=True)
    _live_formula(
        worksheet,
        qc_row,
        3,
        f"=B{responses_row}={response_formula.lstrip('=')}",
        True,
    )
    return qc_row


def _grid_rated_result_source_columns(
    result: GridRatedResult,
    question: Any,
    context: _LiveWorkbookContext,
) -> dict[tuple[str, int], _LiveColumnSpec]:
    mapping: dict[tuple[str, int], _LiveColumnSpec] = {}
    raw_columns = list(getattr(question, "raw_columns", ()) or ())
    for row in result.rows:
        candidates = [
            source_column
            for source_column in raw_columns
            if _grid_row_and_group_ids(source_column)[0] == row.row_id
        ]
        if not candidates and row.row_id in raw_columns:
            candidates = [row.row_id]
        used: set[str] = set()
        for column_index, header in enumerate(result.column_headers):
            source_column = _grid_rated_source_for_header(
                question,
                candidates,
                header,
                column_index,
                used,
            )
            column = context.column_by_key.get(source_column)
            if column is None:
                raise ValueError(
                    f"Missing live data named range for rated grid source column {source_column!r}"
                )
            mapping[(row.row_id, column_index)] = column
            used.add(source_column)
    return mapping


def _grid_rated_source_for_header(
    question: Any,
    candidates: list[str],
    header: str,
    column_index: int,
    used: set[str],
) -> str:
    target = _normalise_header_label(header)
    for source_column in candidates:
        if source_column in used:
            continue
        _row_id, group_key = _grid_row_and_group_ids(source_column)
        label = _grid_group_label(question, group_key)
        if _normalise_header_label(label) == target:
            return source_column
    remaining = [source_column for source_column in candidates if source_column not in used]
    if column_index < len(candidates) and candidates[column_index] not in used:
        return candidates[column_index]
    if remaining:
        return remaining[0]
    raise ValueError(
        f"Could not resolve rated grid source column for header {header!r}"
    )


def _normalise_header_label(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", str(value).lower())


def _build_averageifs_nonnegative_formula(
    data_name: str,
    fq_name: str,
    theme_prefix: str,
) -> str:
    filter_args = _live_filter_criteria_args(theme_prefix, fq_name)
    return f'=IFERROR(AVERAGEIFS({data_name},{data_name},">=0",{filter_args}),"-")'


def _live_write_grid_binary_pivot_result_table(
    worksheet: Any,
    start_row: int,
    result: GridBinaryPivotResult,
) -> int:
    header_row = start_row
    subheader_row = start_row + 1
    worksheet.cell(row=header_row, column=1, value="Sub-question ID").font = _live_font(bold=True)
    worksheet.cell(row=header_row, column=2, value="Sub-question").font = _live_font(bold=True)
    worksheet.cell(row=subheader_row, column=1, value="")
    worksheet.cell(row=subheader_row, column=2, value="")
    col_index = 3
    percent_columns: list[int] = []
    for header in result.column_headers:
        worksheet.cell(row=header_row, column=col_index, value=header).font = _live_font(bold=True)
        worksheet.merge_cells(
            start_row=header_row,
            start_column=col_index,
            end_row=header_row,
            end_column=col_index + 1,
        )
        worksheet.cell(row=subheader_row, column=col_index, value="Count").font = _live_font(bold=True)
        worksheet.cell(row=subheader_row, column=col_index + 1, value="%").font = _live_font(bold=True)
        percent_columns.append(col_index + 1)
        col_index += 2
    last_col = col_index - 1
    for row in (header_row, subheader_row):
        for col in range(1, last_col + 1):
            worksheet.cell(row=row, column=col).fill = _live_fill("F2F2F2")

    data_start = start_row + 2
    for offset, row in enumerate(result.rows):
        excel_row = data_start + offset
        worksheet.cell(row=excel_row, column=1, value=row.row_id).font = _live_font(color="808080")
        worksheet.cell(row=excel_row, column=2, value=row.row_label)
        col = 3
        for count, pct in zip(row.counts_per_column, row.pcts_per_column):
            worksheet.cell(row=excel_row, column=col, value=count).number_format = "#,##0"
            worksheet.cell(row=excel_row, column=col + 1, value=pct).number_format = "0.0%"
            col += 2

    data_end = data_start + len(result.rows) - 1
    if result.rows:
        for percent_col in percent_columns:
            _apply_color_scale_range(worksheet, data_start, percent_col, data_end, percent_col)
    total_row = max(data_start, data_end + 1)
    _write_total_respondents_row(worksheet, total_row, result.total_respondents, last_col)
    responses_row = total_row + 1
    _write_total_responses_row(worksheet, responses_row, result.total_responses, last_col)
    qc_row = responses_row + 1
    worksheet.cell(row=qc_row, column=2, value="QC check").font = _live_font(bold=True)
    worksheet.cell(
        row=qc_row,
        column=3,
        value=sum(sum(row.counts_per_column) for row in result.rows) == result.total_responses,
    )
    return qc_row


def _live_write_nps_table(
    worksheet: Any,
    start_row: int,
    result: NPSResult,
    question: Any,
    context: _LiveWorkbookContext,
    sheet_filters: list[dict[str, str]],
    fq_name: str,
    fv_name: str,
    theme_prefix: str,
) -> int:
    del sheet_filters, fq_name, fv_name, theme_prefix
    headers = [
        "Entity",
        "Promoters (n)",
        "Promoters %",
        "Passives (n)",
        "Passives %",
        "Detractors (n)",
        "Detractors %",
        "Valid N",
        "NPS",
    ]
    _live_header_row(worksheet, start_row, headers)
    row_index = start_row + 1
    for entity, source_column in zip(result.entities, question.raw_columns):
        column = context.column_by_key.get(source_column)
        if column is None:
            continue
        worksheet.cell(row=row_index, column=1, value=entity.entity_label)
        worksheet.cell(row=row_index, column=2, value=entity.promoters).number_format = "#,##0"
        worksheet.cell(row=row_index, column=3, value=entity.pct_promoters).number_format = "0.0%"
        worksheet.cell(row=row_index, column=4, value=entity.passives).number_format = "#,##0"
        worksheet.cell(row=row_index, column=5, value=entity.pct_passives).number_format = "0.0%"
        worksheet.cell(row=row_index, column=6, value=entity.detractors).number_format = "#,##0"
        worksheet.cell(row=row_index, column=7, value=entity.pct_detractors).number_format = "0.0%"
        worksheet.cell(row=row_index, column=8, value=entity.valid_n).number_format = "#,##0"
        worksheet.cell(row=row_index, column=9, value=round(entity.nps_score)).number_format = "0"
        row_index += 1

    data_start = start_row + 1
    data_end = row_index - 1
    if data_end >= data_start:
        _apply_color_scale_range(worksheet, data_start, 9, data_end, 9)
    _openpyxl_add_table(worksheet, start_row, 1, max(start_row + 1, data_end), 9)
    return max(start_row + 1, data_end)


def _live_write_grid_rated_table(
    worksheet: Any,
    start_row: int,
    result: GridSingleSelectResult,
    question: Any,
    context: _LiveWorkbookContext,
    sheet_filters: list[dict[str, str]],
    fq_name: str,
    fv_name: str,
    theme_prefix: str,
) -> int:
    row_entries, group_labels = _grid_rated_row_entries(result, question)
    if not group_labels:
        group_labels = ["All"]

    has_delta = len(group_labels) >= 2
    header_row = start_row
    subheader_row = start_row + 1
    left_headers = [
        "Sub-question ID",
        "Sub-question",
        "# of respondents",
        "% of respondents",
    ]
    for col_index, header in enumerate(left_headers, start=1):
        worksheet.cell(row=header_row, column=col_index, value=header).font = _live_font(bold=True)
        worksheet.cell(row=subheader_row, column=col_index, value="")

    right_start_col = 6
    col_index = right_start_col
    metric_columns: list[int] = []
    for group_label in group_labels:
        worksheet.cell(row=header_row, column=col_index, value=group_label).font = _live_font(bold=True)
        worksheet.merge_cells(
            start_row=header_row,
            start_column=col_index,
            end_row=header_row,
            end_column=col_index + 1,
        )
        worksheet.cell(row=subheader_row, column=col_index, value="Mean").font = _live_font(bold=True)
        worksheet.cell(row=subheader_row, column=col_index + 1, value="Median").font = _live_font(bold=True)
        metric_columns.extend([col_index, col_index + 1])
        col_index += 2

    delta_start_col = None
    if has_delta:
        delta_start_col = col_index
        worksheet.cell(row=header_row, column=col_index, value="Delta").font = _live_font(bold=True)
        worksheet.merge_cells(
            start_row=header_row,
            start_column=col_index,
            end_row=header_row,
            end_column=col_index + 1,
        )
        worksheet.cell(row=subheader_row, column=col_index, value="Mean").font = _live_font(bold=True)
        worksheet.cell(row=subheader_row, column=col_index + 1, value="Median").font = _live_font(bold=True)
        col_index += 2

    last_col = col_index - 1
    for row in (header_row, subheader_row):
        for col in range(1, last_col + 1):
            worksheet.cell(row=row, column=col).fill = _live_fill("F2F2F2")

    data_start = start_row + 2
    for offset, entry in enumerate(row_entries):
        excel_row = data_start + offset
        worksheet.cell(row=excel_row, column=1, value=entry["row_id"]).font = _live_font(color="808080")
        worksheet.cell(row=excel_row, column=2, value=entry["label"])
        response_count, response_formula = _grid_rated_response_count(
            entry,
            context,
            sheet_filters,
            fq_name,
            fv_name,
            theme_prefix,
        )
        _live_formula(
            worksheet,
            excel_row,
            3,
            response_formula,
            response_count,
        )
        total_formula = _build_countifs_formula(
            "respondent_id_data",
            "<>",
            sheet_filters,
            fq_name,
            fv_name,
            theme_prefix=theme_prefix,
        )
        _live_formula(
            worksheet,
            excel_row,
            4,
            f'=IFERROR(C{excel_row}/{total_formula.lstrip("=")},0)',
            response_count / result.valid_n if result.valid_n else 0,
        ).number_format = "0.0%"

        col = right_start_col
        group_cells: list[tuple[str, str]] = []
        for group_label in group_labels:
            stats = entry["groups"].get(group_label, {})
            mean_value = float(stats.get("mean", 0.0))
            median_value = float(stats.get("median", 0.0))
            column = context.column_by_key.get(str(stats.get("source_column", "")))
            if column is not None:
                _live_formula(
                    worksheet,
                    excel_row,
                    col,
                    _build_numeric_formula(
                        "Mean",
                        column.data_name,
                        sheet_filters,
                        fq_name,
                        fv_name,
                        theme_prefix,
                    ),
                    mean_value,
                ).number_format = "0.0"
            else:
                worksheet.cell(row=excel_row, column=col, value=mean_value).number_format = "0.0"
            worksheet.cell(row=excel_row, column=col + 1, value=median_value).number_format = "0.0"
            group_cells.append(
                (
                    f"{_openpyxl_column_letter(col)}{excel_row}",
                    f"{_openpyxl_column_letter(col + 1)}{excel_row}",
                )
            )
            col += 2
        if has_delta and delta_start_col is not None and len(group_cells) >= 2:
            mean_formula = f"=ROUND({group_cells[0][0]}-{group_cells[1][0]},1)"
            median_formula = f"=ROUND({group_cells[0][1]}-{group_cells[1][1]},1)"
            first_stats = entry["groups"].get(group_labels[0], {})
            second_stats = entry["groups"].get(group_labels[1], {})
            _live_formula(
                worksheet,
                excel_row,
                delta_start_col,
                mean_formula,
                round(
                    float(first_stats.get("mean", 0.0))
                    - float(second_stats.get("mean", 0.0)),
                    1,
                ),
            ).number_format = "0.0"
            _live_formula(
                worksheet,
                excel_row,
                delta_start_col + 1,
                median_formula,
                round(
                    float(first_stats.get("median", 0.0))
                    - float(second_stats.get("median", 0.0)),
                    1,
                ),
            ).number_format = "0.0"

    data_end = data_start + len(row_entries) - 1
    if row_entries:
        _apply_color_scale_range(worksheet, data_start, 4, data_end, 4)
        for col in metric_columns:
            _apply_color_scale_range(worksheet, data_start, col, data_end, col)
    total_row = max(data_start, data_end + 1)
    _write_total_respondents_row(worksheet, total_row, result.valid_n, max(4, last_col))
    return total_row


def _live_write_grid_categorical_table(
    worksheet: Any,
    start_row: int,
    result: GridSingleSelectResult,
    question: Any,
    context: _LiveWorkbookContext,
    sheet_filters: list[dict[str, str]],
    fq_name: str,
    fv_name: str,
    theme_prefix: str,
) -> int:
    pivot_entries, pivot_categories = _grid_categorical_c_column_entries(result, question)
    if pivot_entries:
        return _live_write_grid_categorical_c_column_table(
            worksheet,
            start_row,
            pivot_entries,
            pivot_categories,
            context,
            sheet_filters,
            fq_name,
            fv_name,
            theme_prefix,
            result.valid_n,
        )

    categories = _grid_categorical_categories(result)
    header_row = start_row
    subheader_row = start_row + 1
    worksheet.cell(row=header_row, column=1, value="Sub-question ID").font = _live_font(bold=True)
    worksheet.cell(row=header_row, column=2, value="Sub-question").font = _live_font(bold=True)
    worksheet.cell(row=subheader_row, column=1, value="")
    worksheet.cell(row=subheader_row, column=2, value="")
    category_columns: list[tuple[int, int]] = []
    col_index = 3
    for category in categories:
        worksheet.cell(row=header_row, column=col_index, value=category).font = _live_font(bold=True)
        worksheet.merge_cells(
            start_row=header_row,
            start_column=col_index,
            end_row=header_row,
            end_column=col_index + 1,
        )
        worksheet.cell(row=subheader_row, column=col_index, value="Count").font = _live_font(bold=True)
        worksheet.cell(row=subheader_row, column=col_index + 1, value="%").font = _live_font(bold=True)
        category_columns.append((col_index, col_index + 1))
        col_index += 2
    last_col = max(4, col_index - 1)
    for row in (header_row, subheader_row):
        for col in range(1, last_col + 1):
            worksheet.cell(row=row, column=col).fill = _live_fill("F2F2F2")

    row_labels = getattr(question, "grid_row_labels", None) or {}
    data_start = start_row + 2
    for offset, (sub_column_id, row_result) in enumerate(result.rows.items()):
        excel_row = data_start + offset
        worksheet.cell(row=excel_row, column=1, value=sub_column_id).font = _live_font(color="808080")
        worksheet.cell(row=excel_row, column=2, value=row_labels.get(sub_column_id, sub_column_id))
        column = context.column_by_key.get(sub_column_id)
        for category, (category_index, percent_index) in zip(categories, category_columns):
            cached_count = _grid_category_count(row_result, category, categories)
            if column is None:
                worksheet.cell(row=excel_row, column=category_index, value=cached_count)
            else:
                if category == "Other":
                    top_categories = [item for item in categories if item != "Other"]
                    top_refs = ",".join(
                        f"{_openpyxl_column_letter(col)}{excel_row}"
                        for col, _percent_col in category_columns[: len(top_categories)]
                    )
                    total_formula = _build_countifs_formula(
                        column.data_name,
                        "<>",
                        sheet_filters,
                        fq_name,
                        fv_name,
                        theme_prefix=theme_prefix,
                    ).lstrip("=")
                    formula = f"={total_formula}-SUM({top_refs})"
                else:
                    formula = _build_countifs_formula(
                        column.data_name,
                        category,
                        sheet_filters,
                        fq_name,
                        fv_name,
                        theme_prefix=theme_prefix,
                    )
                _live_formula(
                    worksheet,
                    excel_row,
                    category_index,
                    formula,
                    cached_count,
                )
            count_ref = f"{_openpyxl_column_letter(category_index)}{excel_row}"
            row_total = max(1, row_result.valid_n)
            _live_formula(
                worksheet,
                excel_row,
                percent_index,
                f'=IFERROR({count_ref}/{row_total},0)',
                cached_count / row_total if row_total else 0,
            ).number_format = "0.0%"

    data_end = data_start + len(result.rows) - 1
    if result.rows:
        for _count_col, percent_col in category_columns:
            _apply_color_scale_range(worksheet, data_start, percent_col, data_end, percent_col)
    total_row = max(data_start, data_end + 1)
    _write_total_respondents_row(worksheet, total_row, result.valid_n, last_col)
    total_responses_row = total_row + 1
    _write_total_responses_row(
        worksheet,
        total_responses_row,
        _grid_total_responses(result, categories),
        last_col,
    )
    return total_responses_row


def _live_write_grid_categorical_c_column_table(
    worksheet: Any,
    start_row: int,
    row_entries: list[dict[str, Any]],
    categories: list[str],
    context: _LiveWorkbookContext,
    sheet_filters: list[dict[str, str]],
    fq_name: str,
    fv_name: str,
    theme_prefix: str,
    total_respondents: int,
) -> int:
    header_row = start_row
    subheader_row = start_row + 1
    worksheet.cell(row=header_row, column=1, value="Sub-question ID").font = _live_font(bold=True)
    worksheet.cell(row=header_row, column=2, value="Sub-question").font = _live_font(bold=True)
    worksheet.cell(row=subheader_row, column=1, value="")
    worksheet.cell(row=subheader_row, column=2, value="")
    category_columns: list[tuple[int, int]] = []
    col_index = 3
    for category in categories:
        worksheet.cell(row=header_row, column=col_index, value=category).font = _live_font(bold=True)
        worksheet.merge_cells(
            start_row=header_row,
            start_column=col_index,
            end_row=header_row,
            end_column=col_index + 1,
        )
        worksheet.cell(row=subheader_row, column=col_index, value="Count").font = _live_font(bold=True)
        worksheet.cell(row=subheader_row, column=col_index + 1, value="%").font = _live_font(bold=True)
        category_columns.append((col_index, col_index + 1))
        col_index += 2
    last_col = max(4, col_index - 1)
    for row in (header_row, subheader_row):
        for col in range(1, last_col + 1):
            worksheet.cell(row=row, column=col).fill = _live_fill("F2F2F2")

    data_start = start_row + 2
    for offset, entry in enumerate(row_entries):
        excel_row = data_start + offset
        worksheet.cell(row=excel_row, column=1, value=entry["row_id"]).font = _live_font(color="808080")
        worksheet.cell(row=excel_row, column=2, value=entry["label"])
        row_total = max(1, int(entry.get("respondent_total", 0)))
        for category, (count_col, percent_col) in zip(categories, category_columns):
            sources = entry["categories"].get(category, [])
            cached_count = sum(int(source.get("count", 0)) for source in sources)
            formula = _grid_categorical_source_sum_formula(
                sources,
                context,
                sheet_filters,
                fq_name,
                fv_name,
                theme_prefix,
            )
            _live_formula(
                worksheet,
                excel_row,
                count_col,
                formula,
                cached_count,
            )
            count_ref = f"{_openpyxl_column_letter(count_col)}{excel_row}"
            _live_formula(
                worksheet,
                excel_row,
                percent_col,
                f'=IFERROR({count_ref}/{row_total},0)',
                cached_count / row_total if row_total else 0,
            ).number_format = "0.0%"

    data_end = data_start + len(row_entries) - 1
    if row_entries:
        for count_col, percent_col in category_columns:
            _apply_color_scale_range(worksheet, data_start, count_col, data_end, count_col)
            _apply_color_scale_range(worksheet, data_start, percent_col, data_end, percent_col)
    total_row = max(data_start, data_end + 1)
    _write_total_respondents_row(worksheet, total_row, total_respondents, last_col)
    total_responses_row = total_row + 1
    _write_total_responses_row(
        worksheet,
        total_responses_row,
        sum(
            int(source.get("count", 0))
            for entry in row_entries
            for sources in entry["categories"].values()
            for source in sources
        ),
        last_col,
    )
    return total_responses_row


def _live_write_categorical_table(
    worksheet: Any,
    start_row: int,
    rows: list[dict[str, Any]],
    sheet_filters: list[dict[str, str]],
    fq_name: str,
    fv_name: str,
    theme_prefix: str,
    total_respondents: int | None = None,
    total_responses: int | None = None,
) -> int:
    headers = ["Option", "Count", "%", "Denominator"]
    _live_header_row(worksheet, start_row, headers)
    data_start = start_row + 1
    total_count = sum(int(row.get("sort_count", 0)) for row in rows)
    is_multi_select = total_responses is not None
    respondent_denominator = (
        int(total_respondents)
        if total_respondents is not None
        else total_count
    )
    denominator_data_names = (
        list(rows[0].get("denominator_data_names", []))
        if rows and is_multi_select
        else []
    )
    live_respondent_denominator_formula = (
        _build_multi_select_respondent_count_formula(
            denominator_data_names,
            sheet_filters,
            fq_name,
            fv_name,
            theme_prefix,
        )
        if denominator_data_names
        else None
    )
    for offset, row in enumerate(rows):
        excel_row = data_start + offset
        worksheet.cell(row=excel_row, column=1, value=row["label"])
        count_formula = _build_countifs_formula(
            row["data_name"],
            row["criteria"],
            sheet_filters,
            fq_name,
            fv_name,
            theme_prefix=theme_prefix,
        )
        _live_formula(
            worksheet,
            excel_row,
            2,
            count_formula,
            int(row.get("sort_count", 0)),
        )
    data_end = max(data_start, data_start + len(rows) - 1)
    if rows:
        for row_index in range(data_start, data_end + 1):
            count_value = int(rows[row_index - data_start].get("sort_count", 0))
            if is_multi_select:
                denominator_formula = (
                    live_respondent_denominator_formula
                    or f"SUBTOTAL(9,B{data_start}:B{data_end})"
                )
                _live_formula(
                    worksheet,
                    row=row_index,
                    column=3,
                    formula=f"=IFERROR(B{row_index}/{denominator_formula},0)",
                    cached_value=(
                        count_value / respondent_denominator
                        if respondent_denominator
                        else 0
                    ),
                ).number_format = "0.0%"
                _live_formula(
                    worksheet,
                    row=row_index,
                    column=4,
                    formula=f"={denominator_formula}",
                    cached_value=respondent_denominator,
                ).number_format = "#,##0"
            else:
                _live_formula(
                    worksheet,
                    row=row_index,
                    column=3,
                    formula=f"=IFERROR(B{row_index}/SUBTOTAL(9,B{data_start}:B{data_end}),0)",
                    cached_value=(count_value / total_count if total_count else 0),
                ).number_format = "0.0%"
                _live_formula(
                    worksheet,
                    row=row_index,
                    column=4,
                    formula=f"=SUBTOTAL(9,B{data_start}:B{data_end})",
                    cached_value=total_count,
                ).number_format = "#,##0"
        _apply_distribution_conditional_formatting(worksheet, data_start, data_end)
    else:
        data_end = start_row

    total_row = data_end + 1
    respondent_total = total_respondents if total_respondents is not None else total_count
    _write_total_respondents_row(worksheet, total_row, respondent_total, 4)
    table_end_row = total_row
    if total_responses is not None:
        table_end_row = total_row + 1
        _write_total_responses_row(worksheet, table_end_row, total_responses, 4)
    _openpyxl_add_table(worksheet, start_row, 1, table_end_row, 4)
    return table_end_row


def _numeric_table_is_filtered_view(
    sheet_filters: list[dict[str, str]],
    fq_name: str,
    fv_name: str,
    theme_prefix: str,
) -> bool:
    del fq_name, fv_name, theme_prefix
    inactive_values = {"", "(All)", "(Inherit)", "(None)", None}
    for sheet_filter in sheet_filters:
        for key in ("value", "filter_value", "selected_value"):
            if key in sheet_filter and sheet_filter.get(key) not in inactive_values:
                return True
    return False


def _live_write_numeric_table(
    worksheet: Any,
    start_row: int,
    result: NumericResult,
    question: Any,
    context: _LiveWorkbookContext,
    sheet_filters: list[dict[str, str]],
    fq_name: str,
    fv_name: str,
    theme_prefix: str,
) -> int:
    if result.question_type is QuestionType.NUMERIC_ALLOCATION:
        return _live_write_numeric_allocation_table(
            worksheet,
            start_row,
            result,
            question,
            context,
            sheet_filters,
            fq_name,
            fv_name,
            theme_prefix,
        )

    headers = ["Metric", "Value", "Note", "Denominator"]
    _live_header_row(worksheet, start_row, headers)
    row_index = start_row + 1
    column = context.column_by_key.get(question.canonical_id)
    if column is None:
        return start_row
    is_filtered_view = _numeric_table_is_filtered_view(
        sheet_filters,
        fq_name,
        fv_name,
        theme_prefix,
    )
    for metric, static_value in (
        ("Mean", None),
        ("Min", None),
        ("Max", None),
        ("Std", result.std),
    ):
        worksheet.cell(row=row_index, column=1, value=metric)
        if static_value is None:
            _live_formula(
                worksheet,
                row=row_index,
                column=2,
                formula=_build_numeric_formula(
                    metric,
                    column.data_name,
                    sheet_filters,
                    fq_name,
                    fv_name,
                    theme_prefix,
                ),
                cached_value=_numeric_result_cache_value(metric, result),
            ).number_format = "0.00"
        else:
            worksheet.cell(row=row_index, column=2, value=static_value).number_format = "0.00"
            worksheet.cell(row=row_index, column=3, value="static baseline")
        _live_formula(
            worksheet,
            row_index,
            4,
            _build_numeric_count_formula(column.data_name, sheet_filters, fq_name, fv_name, theme_prefix),
            result.valid_n,
        )
        row_index += 1
    if is_filtered_view:
        worksheet.cell(
            row=row_index,
            column=1,
            value="Median not available in filtered view - see static baseline in audit log.",
        ).font = _live_font(italic=True, size=9, color="666666")
        table_end = max(start_row + 1, row_index - 1)
    else:
        worksheet.cell(row=row_index, column=1, value="Median")
        worksheet.cell(row=row_index, column=2, value=result.median).number_format = "0.00"
        worksheet.cell(row=row_index, column=3, value="static baseline")
        worksheet.cell(row=row_index, column=4, value=result.valid_n).number_format = "#,##0"
        table_end = row_index
    _openpyxl_add_table(worksheet, start_row, 1, table_end, 4)
    return row_index


def _live_write_numeric_allocation_table(
    worksheet: Any,
    start_row: int,
    result: NumericResult,
    question: Any,
    context: _LiveWorkbookContext,
    sheet_filters: list[dict[str, str]],
    fq_name: str,
    fv_name: str,
    theme_prefix: str,
) -> int:
    headers = ["Option", "Mean", "Median", "Denominator"]
    _live_header_row(worksheet, start_row, headers)
    row_index = start_row + 1
    for option_id, payload in (result.per_option_stats or {}).items():
        column = context.column_by_key.get(str(option_id))
        if column is None:
            continue
        worksheet.cell(
            row=row_index,
            column=1,
            value=_numeric_allocation_option_label(question, str(option_id)),
        )
        _live_formula(
            worksheet,
            row=row_index,
            column=2,
            formula=_build_numeric_formula(
                "Mean",
                column.data_name,
                sheet_filters,
                fq_name,
                fv_name,
                theme_prefix,
            ),
            cached_value=_numeric_formula_cache_value("Mean", payload),
        ).number_format = "0.00"
        worksheet.cell(
            row=row_index,
            column=3,
            value=_numeric_formula_cache_value("Median", payload),
        ).number_format = "0.00"
        _live_formula(
            worksheet,
            row_index,
            4,
            _build_numeric_count_formula(
                column.data_name,
                sheet_filters,
                fq_name,
                fv_name,
                theme_prefix,
            ),
            int(payload.get("valid_n", result.valid_n)),
        )
        row_index += 1

    data_start = start_row + 1
    data_end = row_index - 1
    if data_end >= data_start:
        _apply_color_scale_range(worksheet, data_start, 2, data_end, 2)
        _apply_color_scale_range(worksheet, data_start, 3, data_end, 3)
    total_row = row_index
    _write_total_respondents_row(worksheet, total_row, result.valid_n, 4)
    _openpyxl_add_table(worksheet, start_row, 1, total_row, 4)
    return total_row


def _live_write_cross_tab_table(
    worksheet: Any,
    start_row: int,
    rows: list[dict[str, Any]],
    sheet_filters: list[dict[str, str]],
    ct_name: str,
    fq_name: str,
    fv_name: str,
    theme_prefix: str,
    context: _LiveWorkbookContext,
) -> int:
    if not rows:
        return start_row
    metric_kind = str(rows[0].get("metric_kind") or "count")
    rank_controls: dict[str, Any] | None = None
    first_col = 1
    max_groups = int(CROSS_TAB_MAX_GROUPS)
    if max_groups < 2:
        max_groups = 2
    top_groups = max_groups - 1
    total_col = first_col + max_groups + 1
    header_row = start_row
    subheader_row = start_row + 1
    data_start = start_row + 2

    worksheet.cell(row=header_row, column=first_col, value="Option").font = _live_font(bold=True)
    worksheet.cell(row=subheader_row, column=first_col, value="").font = _live_font(bold=True)
    options_formula = _cross_tab_options_name_formula(ct_name)
    ct_range = f"INDIRECT({_cross_tab_data_name_formula(ct_name)})"
    filtered_total_formula = _build_countifs_formula(
        "respondent_id_data",
        "<>",
        sheet_filters,
        fq_name,
        fv_name,
        theme_prefix=theme_prefix,
    )
    if metric_kind == "rank_weighted_average" and int(rows[0].get("rank_k", 0) or 0) > 0:
        rank_controls = _live_write_rank_cross_tab_controls(
            worksheet,
            header_row,
            total_col + 2,
            rows,
            ct_name,
            context,
        )

    for offset in range(max_groups):
        group_col = first_col + 1 + offset
        if offset < top_groups:
            header_formula = (
                f'=IF({ct_name}="(None)","",'
                f'IFERROR(INDEX(INDIRECT({options_formula}),{offset + 2}),""))'
            )
        else:
            header_formula = (
                f'=IF({ct_name}="(None)","",'
                f'IF(COUNTA(INDIRECT({options_formula}))-1>{top_groups},'
                f'"Other",IFERROR(INDEX(INDIRECT({options_formula}),{offset + 2}),"")))'
            )
        _live_formula(
            worksheet,
            header_row,
            group_col,
            header_formula,
            "",
        )
        if rank_controls:
            subheader_cell = _live_formula(
                worksheet,
                subheader_row,
                group_col,
                _rank_metric_subheader_formula(rank_controls["metric_name"]),
                rank_controls["metric"],
            )
            subheader_cell.font = _live_font(bold=True)
        else:
            worksheet.cell(
                row=subheader_row,
                column=group_col,
                value=_cross_tab_metric_subheader(metric_kind),
            ).font = _live_font(bold=True)

    worksheet.cell(row=header_row, column=total_col, value="Total").font = _live_font(bold=True)
    if rank_controls:
        total_subheader_cell = _live_formula(
            worksheet,
            subheader_row,
            total_col,
            _rank_metric_subheader_formula(rank_controls["metric_name"]),
            rank_controls["metric"],
        )
        total_subheader_cell.font = _live_font(bold=True)
    else:
        worksheet.cell(
            row=subheader_row,
            column=total_col,
            value=_cross_tab_metric_subheader(metric_kind),
        ).font = _live_font(bold=True)
    for col_index in range(first_col, total_col + 1):
        worksheet.cell(row=header_row, column=col_index).fill = _live_fill("F2F2F2")
        worksheet.cell(row=subheader_row, column=col_index).fill = _live_fill("F2F2F2")

    for row_offset, row in enumerate(rows):
        excel_row = data_start + row_offset
        worksheet.cell(row=excel_row, column=first_col, value=row["label"])
        total_expression = _live_cross_tab_metric_expression(
            row,
            metric_kind,
            sheet_filters,
            fq_name,
            theme_prefix,
            extra_pairs=[],
            header_ref=None,
            rank_controls=rank_controls,
        )
        for offset in range(max_groups):
            group_col = first_col + 1 + offset
            header_ref = f"{_openpyxl_column_letter(group_col)}${header_row}"
            if offset < top_groups:
                count_expression = _live_cross_tab_metric_expression(
                    row,
                    metric_kind,
                    sheet_filters,
                    fq_name,
                    theme_prefix,
                    extra_pairs=[(ct_range, header_ref)],
                    header_ref=header_ref,
                    rank_controls=rank_controls,
                )
            else:
                first_group_col = first_col + 1
                last_top_count_col = first_col + top_groups
                top_count_refs = ",".join(
                    f"{_openpyxl_column_letter(col)}{excel_row}"
                    for col in range(first_group_col, last_top_count_col + 1)
                )
                if metric_kind == "count":
                    normal_count_formula = _live_cross_tab_metric_expression(
                        row,
                        metric_kind,
                        sheet_filters,
                        fq_name,
                        theme_prefix,
                        extra_pairs=[(ct_range, header_ref)],
                        header_ref=header_ref,
                        rank_controls=rank_controls,
                    )
                    count_expression = (
                        f'IF({header_ref}="Other",'
                        f"{total_expression}-SUM({top_count_refs}),"
                        f"{normal_count_formula})"
                    )
                else:
                    count_expression = _live_cross_tab_metric_expression(
                        row,
                        metric_kind,
                        sheet_filters,
                        fq_name,
                        theme_prefix,
                        extra_pairs=[(ct_range, header_ref)],
                        header_ref=header_ref,
                        rank_controls=rank_controls,
                    )
            cell = _live_formula(
                worksheet,
                excel_row,
                group_col,
                f'=IF(OR({ct_name}="(None)",{header_ref}=""),"",{count_expression})',
                row.get("cached_value", 0),
            )
            _apply_cross_tab_metric_format(cell, metric_kind, rank_selectable=bool(rank_controls))
            formula = f'=IF(OR({ct_name}="(None)",{header_ref}=""),"",{count_expression})'
            if rank_controls:
                _record_live_rank_cross_tab_audits(
                    context,
                    worksheet.title,
                    row,
                    formula,
                    f"{ct_name} = {header_ref}",
                )
            else:
                _record_live_cross_tab_audit(
                    context,
                    worksheet.title,
                    row,
                    metric_kind,
                    formula,
                    f"{ct_name} = {header_ref}",
                )

        total_cell = _live_formula(
            worksheet,
            excel_row,
            total_col,
            f'=IF({ct_name}="(None)","",{total_expression})',
            row.get("cached_total", row.get("cached_value", 0)),
        )
        _apply_cross_tab_metric_format(total_cell, metric_kind, rank_selectable=bool(rank_controls))

    total_row = data_start + len(rows)
    worksheet.cell(row=total_row, column=first_col, value="Total respondents").font = _live_font(bold=True)
    _live_formula(
        worksheet,
        total_row,
        total_col,
        f'=IF({ct_name}="(None)","",{filtered_total_formula.lstrip("=")})',
        0,
    )
    for col_index in range(first_col, total_col + 1):
        worksheet.cell(row=total_row, column=col_index).fill = _live_fill("F2F2F2")
    if rows:
        data_end = data_start + len(rows) - 1
        if not rank_controls:
            for offset in range(max_groups):
                count_col = first_col + 1 + offset
                _apply_color_scale_range(worksheet, data_start, count_col, data_end, count_col)
    return total_row


def _live_write_rank_cross_tab_controls(
    worksheet: Any,
    header_row: int,
    start_col: int,
    rows: list[dict[str, Any]],
    ct_name: str,
    context: _LiveWorkbookContext,
) -> dict[str, Any] | None:
    workbook = worksheet.parent
    rank_k = max(int(row.get("rank_k", 0) or 0) for row in rows)
    if rank_k <= 0:
        return None
    question_id = str(rows[0].get("question_id") or "")
    settings = _normalise_rank_cross_tab_settings(
        context.rank_cross_tab_settings.get(question_id),
        rank_k,
    )
    base_name = ct_name[:-3] if ct_name.endswith("_CT") else ct_name
    metric_name = _safe_defined_name(f"{base_name}_RM")
    points_name = _safe_defined_name(f"{base_name}_PTS")
    position_name = _safe_defined_name(f"{base_name}_RP")

    label_col = start_col
    value_col = start_col + 1
    worksheet.cell(row=header_row, column=label_col, value="Rank metric").font = _live_font(bold=True, size=9)
    metric_cell = worksheet.cell(row=header_row, column=value_col, value=settings["metric"])
    _add_named_cell(workbook, metric_name, worksheet, metric_cell.coordinate)
    _add_dropdown_to_cell(
        worksheet,
        metric_cell.coordinate,
        '"' + ",".join(RANK_CROSS_TAB_METRICS) + '"',
        allow_blank=False,
    )

    worksheet.cell(row=header_row + 1, column=label_col, value="Rank position").font = _live_font(bold=True, size=9)
    position_cell = worksheet.cell(row=header_row + 1, column=value_col, value=settings["rank_position"])
    _add_named_cell(workbook, position_name, worksheet, position_cell.coordinate)
    _add_dropdown_to_cell(
        worksheet,
        position_cell.coordinate,
        '"' + ",".join(str(rank) for rank in range(1, rank_k + 1)) + '"',
        allow_blank=False,
    )

    for col_index in range(label_col, value_col + 1):
        for row_index in range(header_row, header_row + 2):
            worksheet.cell(row=row_index, column=col_index).fill = _live_fill("F7F7F7")
    return {
        "metric_name": metric_name,
        "points_name": points_name,
        "position_name": position_name,
        "metric": settings["metric"],
        "points": settings["points"],
        "rank_position": settings["rank_position"],
    }


def _normalise_rank_cross_tab_settings(raw: Any, rank_k: int) -> dict[str, Any]:
    raw = raw if isinstance(raw, dict) else {}
    metric = str(raw.get("metric") or "Weighted Average")
    if metric == "Weighted average":
        metric = "Weighted Average"
    if metric not in RANK_CROSS_TAB_METRICS:
        metric = "Weighted Average"
    points_raw = raw.get("points", raw.get("weights"))
    points: list[float] = []
    if isinstance(points_raw, (list, tuple)):
        for value in points_raw[:rank_k]:
            try:
                points.append(float(value))
            except (TypeError, ValueError):
                points.append(0.0)
    if len(points) != rank_k:
        points = [float(rank_k - offset) for offset in range(rank_k)]
    try:
        rank_position = int(raw.get("rank_position") or 1)
    except (TypeError, ValueError):
        rank_position = 1
    rank_position = max(1, min(rank_k, rank_position))
    return {
        "metric": metric,
        "points": points,
        "rank_position": rank_position,
    }


def _rank_weighted_average_from_counts(
    counts_per_rank: list[int],
    points: list[float],
    denominator: int,
) -> float:
    if denominator <= 0:
        return 0.0
    numerator = sum(
        int(count) * float(point)
        for count, point in zip(counts_per_rank, points)
    )
    return float(numerator / denominator)


def _rank_metric_subheader_formula(metric_name: str) -> str:
    return (
        f'=IF({metric_name}="Sum of ranks","Sum of ranks",'
        f'IF({metric_name}="Rank position count","Rank position count","Weighted Average"))'
    )


def _cross_tab_metric_subheader(metric_kind: str) -> str:
    if metric_kind == "mean":
        return "Mean"
    if metric_kind == "rank_mean":
        return "Mean rank"
    if metric_kind == "rank_weighted_average":
        return "Weighted Average"
    if metric_kind == "selection_rate":
        return "% selected"
    if metric_kind == "nps":
        return "NPS"
    return "# of resp"


def _apply_cross_tab_metric_format(
    cell: Any,
    metric_kind: str,
    *,
    rank_selectable: bool = False,
) -> None:
    if rank_selectable:
        cell.number_format = "General"
        return
    if metric_kind == "selection_rate":
        cell.number_format = "0.0%"
    elif metric_kind in {"mean", "rank_mean"}:
        cell.number_format = "0.00"
    elif metric_kind == "rank_weighted_average":
        cell.number_format = "0.00"
    elif metric_kind == "nps":
        cell.number_format = "0"
    else:
        cell.number_format = "#,##0"


def _live_cross_tab_metric_expression(
    row: dict[str, Any],
    metric_kind: str,
    sheet_filters: list[dict[str, str]],
    fq_name: str,
    theme_prefix: str,
    *,
    extra_pairs: list[tuple[str, str]],
    header_ref: str | None,
    rank_controls: dict[str, Any] | None = None,
) -> str:
    if header_ref is not None and metric_kind != "count":
        other_guard_prefix = f'IF({header_ref}="Other","",'
        other_guard_suffix = ")"
    else:
        other_guard_prefix = ""
        other_guard_suffix = ""

    filter_pairs = _live_filter_criteria_pairs(theme_prefix, fq_name)
    all_pairs = [*filter_pairs, *extra_pairs]
    data_name = str(row["data_name"])

    if metric_kind == "count":
        return _build_countifs_formula(
            data_name,
            row["criteria"],
            sheet_filters,
            fq_name,
            "",
            extra_pairs=extra_pairs,
            theme_prefix=theme_prefix,
        ).lstrip("=")

    if metric_kind in {"mean", "rank_mean"}:
        row_criteria = row.get("criteria")
        row_pairs = (
            [(data_name, _countifs_criteria(row_criteria))]
            if row_criteria not in {None, ""}
            else []
        )
        expression = _averageifs_expression(data_name, [*row_pairs, *all_pairs])
        return f"{other_guard_prefix}{expression}{other_guard_suffix}"

    if metric_kind == "rank_weighted_average":
        rank_k = int(row.get("rank_k", 0) or 0)
        if rank_k <= 0:
            return '""'
        rank_sum = _sumifs_expression(
            data_name,
            [(data_name, _countifs_criteria(">=1")), *all_pairs],
        )
        answered_count = _build_multi_select_respondent_count_formula(
            row.get("denominator_data_names") or [data_name],
            sheet_filters,
            fq_name,
            "",
            theme_prefix,
            extra_pairs=extra_pairs,
        )
        points_name = (
            str(rank_controls["points_name"])
            if rank_controls
            else str(row.get("points_name") or "")
        )
        if not points_name:
            return '""'
        weighted_terms = [
            (
                f"INDEX({points_name},{rank})*"
                f"{_countifs_expression([(data_name, str(rank)), *all_pairs])}"
            )
            for rank in range(1, rank_k + 1)
        ]
        expression = (
            f'IFERROR(({" + ".join(weighted_terms)})/{answered_count},"")'
        )
        if rank_controls:
            sum_expression = f'IFERROR({rank_sum},"")'
            position_expression = (
                f'IFERROR({_countifs_expression([(data_name, rank_controls["position_name"]), *all_pairs])},"")'
            )
            metric_name = rank_controls["metric_name"]
            expression = (
                f'IF({metric_name}="Sum of ranks",{sum_expression},'
                f'IF({metric_name}="Rank position count",{position_expression},'
                f"{expression}))"
            )
        return f"{other_guard_prefix}{expression}{other_guard_suffix}"

    if metric_kind == "selection_rate":
        numerator = _countifs_expression(
            [(data_name, _countifs_criteria(row.get("criteria", "Selected"))), *all_pairs]
        )
        denominator = _build_multi_select_respondent_count_formula(
            row.get("denominator_data_names") or [data_name],
            sheet_filters,
            fq_name,
            "",
            theme_prefix,
            extra_pairs=extra_pairs,
        )
        expression = f'IFERROR({numerator}/{denominator},"")'
        return f"{other_guard_prefix}{expression}{other_guard_suffix}"

    if metric_kind == "nps":
        promoters = _countifs_expression(
            [
                (data_name, _countifs_criteria(">=9")),
                (data_name, _countifs_criteria("<=10")),
                *all_pairs,
            ]
        )
        detractors = _countifs_expression(
            [
                (data_name, _countifs_criteria(">=0")),
                (data_name, _countifs_criteria("<=6")),
                *all_pairs,
            ]
        )
        valid_n = _countifs_expression(
            [
                (data_name, _countifs_criteria(">=0")),
                (data_name, _countifs_criteria("<=10")),
                *all_pairs,
            ]
        )
        expression = f'IFERROR((({promoters}/{valid_n})-({detractors}/{valid_n}))*100,"")'
        return f"{other_guard_prefix}{expression}{other_guard_suffix}"

    return "0"


def _averageifs_expression(data_name: str, pairs: list[tuple[str, str]]) -> str:
    args = ",".join([data_name, *(f"{range_expr},{criterion_expr}" for range_expr, criterion_expr in pairs)])
    return f'IFERROR(AVERAGEIFS({args}),"")'


def _countifs_expression(pairs: list[tuple[str, str]]) -> str:
    args = ",".join(f"{range_expr},{criterion_expr}" for range_expr, criterion_expr in pairs)
    return f"COUNTIFS({args})"


def _sumifs_expression(
    sum_range: str,
    pairs: list[tuple[str, str]],
) -> str:
    args = ",".join(f"{range_expr},{criterion_expr}" for range_expr, criterion_expr in pairs)
    return f"SUMIFS({sum_range},{args})"


def _record_live_cross_tab_audit(
    context: _LiveWorkbookContext,
    sheet_name: str,
    row: dict[str, Any],
    metric_kind: str,
    formula: str,
    filter_expr: str,
) -> None:
    if context.log is None or metric_kind == "count":
        return
    context.log.record(
        AuditRecord(
            output_sheet=sheet_name,
            metric_name=f"cross_tab_by_{metric_kind}",
            source_question_id=str(row.get("question_id") or ""),
            source_columns=(str(row.get("source_column") or row.get("data_name") or ""),),
            filter_expr=filter_expr,
            numerator=None,
            denominator=None,
            formula=formula,
            value_raw=float(row.get("cached_value", 0) or 0),
            valid_n=0,
            missing_n=0,
            timestamp=datetime.now(timezone.utc),
        )
    )


def _record_live_rank_cross_tab_audits(
    context: _LiveWorkbookContext,
    sheet_name: str,
    row: dict[str, Any],
    formula: str,
    filter_expr: str,
) -> None:
    for metric_kind in (
        "rank_weighted_average",
        "rank_sum_of_ranks",
        "rank_position_count",
    ):
        _record_live_cross_tab_audit(
            context,
            sheet_name,
            row,
            metric_kind,
            formula,
            filter_expr,
        )


def _legacy_unused_cross_tab_table_marker() -> None:
    """Kept as a separator after the live 2D cross-tab writer."""


def _live_distribution_rows(
    result: SingleCutResult,
    question: Any,
    context: _LiveWorkbookContext,
) -> list[dict[str, Any]]:
    if isinstance(result, SingleSelectResult):
        column = context.column_by_key.get(question.canonical_id)
        if column is None:
            return []
        rows = [
            {
                "label": str(payload["label"]),
                "data_name": column.data_name,
                "criteria": str(payload["label"]),
                "metric_kind": "count",
                "question_id": result.question_id,
                "source_column": question.canonical_id,
                "sort_count": int(payload["count"]),
            }
            for payload in result.distribution.values()
        ]
    elif isinstance(result, MultiSelectResult):
        rows = []
        denominator_data_names = [
            context.column_by_key[sub_column_id].data_name
            for sub_column_id in result.selections
            if sub_column_id in context.column_by_key
            and not _is_computed_multi_select_column(question, sub_column_id)
        ]
        for sub_column_id, payload in result.selections.items():
            if _is_computed_multi_select_column(question, sub_column_id):
                continue
            column = context.column_by_key.get(sub_column_id)
            if column is None:
                continue
            rows.append(
                {
                    "label": str(payload["label"]),
                    "data_name": column.data_name,
                    "criteria": "Selected",
                    "metric_kind": "selection_rate",
                    "question_id": result.question_id,
                    "source_column": sub_column_id,
                    "cached_value": float(payload.get("selection_rate", 0.0) or 0.0),
                    "sort_count": int(payload["count"]),
                    "denominator_data_names": denominator_data_names,
                }
            )
    elif isinstance(result, NumericResult):
        rows = []
        if result.question_type is QuestionType.NUMERIC_ALLOCATION:
            for option_id, payload in (result.per_option_stats or {}).items():
                column = context.column_by_key.get(option_id)
                if column is None:
                    continue
                rows.append(
                    {
                        "label": _numeric_allocation_option_label(question, option_id),
                        "data_name": column.data_name,
                        "criteria": "<>",
                        "metric_kind": "mean",
                        "question_id": result.question_id,
                        "source_column": option_id,
                        "cached_value": float(payload.get("mean", 0.0) or 0.0),
                        "sort_count": int(payload.get("valid_n", 0) or 0),
                    }
                )
        else:
            column = context.column_by_key.get(question.canonical_id)
            if column is None:
                return []
            rows = [
                {
                    "label": str(question.question_text),
                    "data_name": column.data_name,
                    "criteria": "<>",
                    "metric_kind": "mean",
                    "question_id": result.question_id,
                    "source_column": question.canonical_id,
                    "cached_value": float(result.mean),
                    "sort_count": int(result.valid_n),
                }
            ]
    elif isinstance(result, GridRatedResult):
        rows = []
        for row in result.rows:
            source_column = next(
                (
                    column_id
                    for column_id in question.raw_columns
                    if _grid_row_and_group_ids(column_id)[0] == row.row_id
                ),
                row.row_id,
            )
            column = context.column_by_key.get(source_column)
            if column is None:
                continue
            rows.append(
                {
                    "label": row.row_label,
                    "data_name": column.data_name,
                    "criteria": ">=0",
                    "metric_kind": "mean",
                    "question_id": result.question_id,
                    "source_column": source_column,
                    "cached_value": float(row.means_per_column[0] if row.means_per_column else 0.0),
                    "sort_count": int(row.valid_n_per_column[0] if row.valid_n_per_column else 0),
                }
            )
    elif (
        isinstance(result, GridSingleSelectResult)
        and _grid_render_subtype(question, result) == GRID_RATED
    ):
        rows = []
        row_entries, group_labels = _grid_rated_row_entries(result, question)
        if not group_labels:
            group_labels = ["All"]
        for entry in row_entries:
            stats = next(
                (
                    entry.get("groups", {}).get(group_label)
                    for group_label in group_labels
                    if entry.get("groups", {}).get(group_label)
                ),
                None,
            )
            if not isinstance(stats, dict):
                continue
            source_column = str(stats.get("source_column", ""))
            column = context.column_by_key.get(source_column)
            if column is None:
                continue
            rows.append(
                {
                    "label": str(entry.get("label", source_column)),
                    "data_name": column.data_name,
                    "criteria": ">=0",
                    "metric_kind": "mean",
                    "question_id": result.question_id,
                    "source_column": source_column,
                    "cached_value": float(stats.get("mean", 0.0) or 0.0),
                    "sort_count": int(stats.get("valid_n", 0) or 0),
                }
            )
    elif isinstance(result, RankOrderResult):
        rows = []
        denominator_data_names = [
            context.column_by_key[row.option_id].data_name
            for row in result.rows
            if row.option_id in context.column_by_key
        ]
        rank_points = _normalise_rank_cross_tab_settings(
            context.rank_cross_tab_settings.get(result.question_id),
            result.K,
        )["points"]
        for row in result.rows:
            column = context.column_by_key.get(row.option_id)
            if column is None:
                continue
            weighted_average = _rank_weighted_average_from_counts(
                row.counts_per_rank,
                rank_points,
                result.total_respondents,
            )
            rows.append(
                {
                    "label": row.option_label,
                    "data_name": column.data_name,
                    "criteria": ">=1",
                    "metric_kind": "rank_weighted_average",
                    "question_id": result.question_id,
                    "source_column": row.option_id,
                    "cached_value": weighted_average,
                    "cached_total": weighted_average,
                    "sort_count": int(round(weighted_average * 1000)),
                    "denominator_data_names": denominator_data_names,
                    "rank_k": int(getattr(result, "K", 0) or len(row.counts_per_rank)),
                }
            )
    elif isinstance(result, GridBinaryPivotResult):
        rows = []
        row_labels = question.grid_row_labels or {}
        denominator_data_names = [
            context.column_by_key[source_column].data_name
            for source_column in question.raw_columns
            if source_column in context.column_by_key
        ]
        for source_column in question.raw_columns:
            column = context.column_by_key.get(source_column)
            if column is None:
                continue
            rows.append(
                {
                    "label": _grid_base_row_label(str(row_labels.get(source_column, source_column)), "All"),
                    "data_name": column.data_name,
                    "criteria": "Selected",
                    "metric_kind": "selection_rate",
                    "question_id": result.question_id,
                    "source_column": source_column,
                    "cached_value": 0.0,
                    "sort_count": 0,
                    "denominator_data_names": denominator_data_names,
                }
            )
    elif isinstance(result, NPSResult):
        rows = []
        for entity, source_column in zip(result.entities, question.raw_columns):
            column = context.column_by_key.get(source_column)
            if column is None:
                continue
            rows.append(
                {
                    "label": entity.entity_label,
                    "data_name": column.data_name,
                    "criteria": "<>",
                    "metric_kind": "nps",
                    "question_id": result.question_id,
                    "source_column": source_column,
                    "cached_value": round(entity.nps_score),
                    "sort_count": int(entity.valid_n),
                }
            )
    elif isinstance(result, GridSingleSelectResult):
        rows = []
        grid_labels = question.grid_row_labels or {}
        for sub_column_id, row_result in result.rows.items():
            column = context.column_by_key.get(sub_column_id)
            if column is None:
                continue
            count = sum(int(payload["count"]) for payload in row_result.distribution.values())
            if count == 0:
                continue
            rows.append(
                {
                    "label": str(grid_labels.get(sub_column_id, sub_column_id)),
                    "data_name": column.data_name,
                    "criteria": "Selected",
                    "metric_kind": "count",
                    "question_id": result.question_id,
                    "source_column": sub_column_id,
                    "sort_count": count,
                }
            )
    else:
        return []
    rows.sort(key=lambda item: int(item["sort_count"]), reverse=True)
    return rows


def _rank_order_mean(row: Any) -> float:
    total = int(sum(row.counts_per_rank))
    if total <= 0:
        return 0.0
    weighted_sum = sum(
        rank * int(count)
        for rank, count in enumerate(row.counts_per_rank, start=1)
    )
    return float(weighted_sum / total)


def _categorical_total_respondents(result: SingleCutResult) -> int | None:
    if isinstance(result, MultiSelectResult):
        return int(result.respondents_who_answered_any)
    if isinstance(result, GridSingleSelectResult):
        return int(result.valid_n)
    if isinstance(result, SingleSelectResult):
        return int(result.valid_n)
    return None


def _categorical_total_responses(result: SingleCutResult) -> int | None:
    if isinstance(result, MultiSelectResult):
        return sum(int(payload.get("count", 0)) for payload in result.selections.values())
    if isinstance(result, GridSingleSelectResult):
        return sum(
            int(payload.get("count", 0))
            for row_result in result.rows.values()
            for payload in row_result.distribution.values()
        )
    return None


def _grid_total_responses(
    result: GridSingleSelectResult,
    visible_categories: list[str],
) -> int:
    return sum(
        _grid_category_count(row_result, category, visible_categories)
        for row_result in result.rows.values()
        for category in visible_categories
    )


def _grid_rated_response_count(
    entry: dict[str, Any],
    context: _LiveWorkbookContext,
    sheet_filters: list[dict[str, str]],
    fq_name: str,
    fv_name: str,
    theme_prefix: str,
) -> tuple[int, str]:
    groups = entry.get("groups", {})
    source_columns = [
        str(payload.get("source_column", ""))
        for payload in groups.values()
        if isinstance(payload, dict) and payload.get("source_column")
    ]
    cached_count = max(
        (
            int(payload.get("valid_n", 0))
            for payload in groups.values()
            if isinstance(payload, dict)
        ),
        default=0,
    )
    first_column = next(
        (
            context.column_by_key.get(source_column)
            for source_column in source_columns
            if context.column_by_key.get(source_column) is not None
        ),
        None,
    )
    if first_column is None:
        return cached_count, "=0"
    return (
        cached_count,
        _build_countifs_formula(
            first_column.data_name,
            "<>",
            sheet_filters,
            fq_name,
            fv_name,
            theme_prefix=theme_prefix,
        ),
    )


def _grid_rated_row_entries(
    result: GridSingleSelectResult,
    question: Any,
) -> tuple[list[dict[str, Any]], list[str]]:
    row_labels = getattr(question, "grid_row_labels", None) or {}
    grouped: dict[str, dict[str, Any]] = {}
    group_labels: list[str] = []
    grouped_by_column_pattern = False

    for sub_column_id, row_result in result.rows.items():
        base_row_id, group_key = _grid_row_and_group_ids(sub_column_id)
        if group_key is not None:
            grouped_by_column_pattern = True
        group_label = _grid_group_label(question, group_key) if group_key is not None else "All"
        if group_label not in group_labels:
            group_labels.append(group_label)
        entry = grouped.setdefault(
            base_row_id,
            {
                "row_id": base_row_id,
                "label": _grid_base_row_label(row_labels.get(sub_column_id, sub_column_id), group_label),
                "groups": {},
            },
        )
        if not grouped_by_column_pattern:
            entry["label"] = row_labels.get(sub_column_id, sub_column_id)
        entry["groups"][group_label] = {
            **_weighted_numeric_stats(row_result),
            "source_column": sub_column_id,
        }

    if not grouped_by_column_pattern:
        group_labels = ["All"]
    return list(grouped.values()), group_labels


def _grid_row_and_group_ids(sub_column_id: str) -> tuple[str, str | None]:
    match = re.match(r"^(?P<row>.+?r\d+)(?:c(?P<group>\d+))$", str(sub_column_id))
    if match is None:
        return str(sub_column_id), None
    return match.group("row"), match.group("group")


def _grid_group_label(question: Any, group_key: str | None) -> str:
    if group_key is None:
        return "All"
    option_map = getattr(question, "option_map", None) or {}
    for candidate in (group_key, int(group_key) if str(group_key).isdigit() else group_key):
        if candidate in option_map:
            label = str(option_map[candidate])
            if not _grid_group_label_is_numeric_scale(label, group_key):
                return label
    return _default_grid_group_label(group_key)


def _grid_group_label_is_numeric_scale(label: str, group_key: str) -> bool:
    text = str(label).strip()
    try:
        return float(text) == float(group_key)
    except (TypeError, ValueError):
        return False


def _default_grid_group_label(group_key: str) -> str:
    if str(group_key) == "1":
        return "Winner - All"
    if str(group_key) == "2":
        return "Other considered vendor"
    return f"Group {group_key}"


def _grid_base_row_label(label: str, group_label: str) -> str:
    text = str(label)
    if group_label and group_label in text:
        text = text.replace(group_label, "")
    text = re.sub(r"\s{2,}", " ", text).strip(" -:|")
    return text or str(label)


def _weighted_numeric_stats(row_result: SingleSelectResult) -> dict[str, float]:
    values: list[float] = []
    for code, payload in row_result.distribution.items():
        numeric_value = _numeric_distribution_value(code, payload)
        if numeric_value is None:
            continue
        values.extend([numeric_value] * int(payload.get("count", 0)))
    if not values:
        return {"mean": 0.0, "median": 0.0, "valid_n": 0}
    values.sort()
    count = len(values)
    middle = count // 2
    if count % 2:
        median = values[middle]
    else:
        median = (values[middle - 1] + values[middle]) / 2.0
    return {
        "mean": round(float(sum(values) / count), 1),
        "median": round(float(median), 1),
        "valid_n": count,
    }


def _numeric_distribution_value(code: Any, payload: dict[str, Any]) -> float | None:
    label = payload.get("label") if isinstance(payload, dict) else None
    for candidate in (label, code):
        try:
            return float(str(candidate).strip())
        except (TypeError, ValueError):
            match = re.match(r"^\s*(-?\d+(?:\.\d+)?)", str(candidate))
            if match is not None:
                try:
                    return float(match.group(1))
                except (TypeError, ValueError):
                    continue
    return None


def _grid_categorical_categories(result: GridSingleSelectResult) -> list[str]:
    totals: dict[str, int] = defaultdict(int)
    for row_result in result.rows.values():
        for payload in row_result.distribution.values():
            label = str(payload.get("label", "")).strip()
            if not label:
                continue
            totals[label] += int(payload.get("count", 0))
    sorted_labels = [
        label for label, _count in sorted(
            totals.items(),
            key=lambda item: item[1],
            reverse=True,
        )
    ]
    if len(sorted_labels) <= 8:
        return sorted_labels
    return [*sorted_labels[:7], "Other"]


def _grid_categorical_c_column_entries(
    result: GridSingleSelectResult,
    question: Any,
) -> tuple[list[dict[str, Any]], list[str]]:
    row_labels = getattr(question, "grid_row_labels", None) or {}
    grouped: dict[str, dict[str, Any]] = {}
    category_totals: dict[str, int] = defaultdict(int)
    category_order: list[str] = []
    saw_c_column = False
    for sub_column_id, row_result in result.rows.items():
        base_row_id, group_key = _grid_row_and_group_ids(sub_column_id)
        if group_key is None:
            continue
        saw_c_column = True
        category = _grid_group_label(question, group_key)
        if category not in category_order:
            category_order.append(category)
        count = _grid_selected_response_count(row_result)
        category_totals[category] += count
        entry = grouped.setdefault(
            base_row_id,
            {
                "row_id": base_row_id,
                "label": _grid_base_row_label(
                    row_labels.get(sub_column_id, base_row_id),
                    category,
                ),
                "categories": defaultdict(list),
                "respondent_total": 0,
            },
        )
        entry["categories"][category].append(
            {
                "source_column": sub_column_id,
                "row_result": row_result,
                "count": count,
            }
        )
        entry["respondent_total"] = max(
            int(entry.get("respondent_total", 0)),
            int(row_result.valid_n),
        )

    if not saw_c_column:
        return [], []

    if len(category_order) > 8:
        top_categories = [
            category
            for category, _count in sorted(
                category_totals.items(),
                key=lambda item: item[1],
                reverse=True,
            )
        ][:7]
        visible_categories = [*top_categories, "Other"]
        other_categories = set(category_order) - set(top_categories)
        for entry in grouped.values():
            other_sources = []
            for category in list(entry["categories"]):
                if category in other_categories:
                    other_sources.extend(entry["categories"].pop(category))
            entry["categories"]["Other"].extend(other_sources)
    else:
        visible_categories = category_order
    return list(grouped.values()), visible_categories


def _grid_selected_response_count(row_result: SingleSelectResult) -> int:
    return sum(int(payload.get("count", 0)) for payload in row_result.distribution.values())


def _grid_categorical_source_sum_formula(
    sources: list[dict[str, Any]],
    context: _LiveWorkbookContext,
    sheet_filters: list[dict[str, str]],
    fq_name: str,
    fv_name: str,
    theme_prefix: str,
) -> str:
    formulas: list[str] = []
    for source in sources:
        column = context.column_by_key.get(str(source.get("source_column", "")))
        row_result = source.get("row_result")
        if column is None or row_result is None:
            continue
        criteria = _grid_selected_criteria(row_result)
        if not criteria:
            formulas.append(
                _build_countifs_formula(
                    column.data_name,
                    "<>",
                    sheet_filters,
                    fq_name,
                    fv_name,
                    theme_prefix=theme_prefix,
                ).lstrip("=")
            )
            continue
        for criterion in criteria:
            formulas.append(
                _build_countifs_formula(
                    column.data_name,
                    criterion,
                    sheet_filters,
                    fq_name,
                    fv_name,
                    theme_prefix=theme_prefix,
                ).lstrip("=")
            )
    if not formulas:
        return "=0"
    return "=" + "+".join(formulas)


def _grid_selected_criteria(row_result: SingleSelectResult) -> list[Any]:
    criteria: list[Any] = []
    for code, payload in row_result.distribution.items():
        label = payload.get("label") if isinstance(payload, dict) else None
        criteria.append(label if label not in {None, ""} else code)
    return criteria


def _grid_category_count(
    row_result: SingleSelectResult,
    category: str,
    visible_categories: list[str],
) -> int:
    if category == "Other":
        top = set(item for item in visible_categories if item != "Other")
        return sum(
            int(payload.get("count", 0))
            for payload in row_result.distribution.values()
            if str(payload.get("label", "")).strip() not in top
        )
    return sum(
        int(payload.get("count", 0))
        for payload in row_result.distribution.values()
        if str(payload.get("label", "")).strip() == category
    )


def _build_countifs_formula(
    data_name: str,
    criteria: Any,
    sheet_filters: list[dict[str, str]],
    fq_name: str,
    fv_name: str,
    extra_pairs: list[tuple[str, str]] | None = None,
    theme_prefix: str = "",
) -> str:
    del sheet_filters, fv_name
    pairs = [
        (data_name, _countifs_criteria(criteria)),
    ]
    pairs.extend(_live_filter_criteria_pairs(theme_prefix, fq_name))
    for range_expr, criterion_expr in extra_pairs or []:
        pairs.append((range_expr, criterion_expr))

    args = ",".join(
        f"{range_expr},{criterion_expr}"
        for range_expr, criterion_expr in pairs
    )
    return f"=COUNTIFS({args})"


def _build_multi_select_respondent_count_formula(
    data_names: list[str],
    sheet_filters: list[dict[str, str]],
    fq_name: str,
    fv_name: str,
    theme_prefix: str = "",
    extra_pairs: list[tuple[str, str]] | None = None,
) -> str:
    del sheet_filters, fv_name
    if not data_names:
        return "0"
    answered_terms = "+".join(f"--({data_name}<>\"\")" for data_name in data_names)
    factors = [f"--(({answered_terms})>0)"]
    factors.extend(
        f"--({range_expr}={criterion_expr})"
        for range_expr, criterion_expr in _live_filter_criteria_pairs(theme_prefix, fq_name)
    )
    factors.extend(
        f"--({range_expr}={criterion_expr})"
        for range_expr, criterion_expr in extra_pairs or []
    )
    return f"SUMPRODUCT({','.join(factors)})"


def _build_numeric_formula(
    metric: str,
    data_name: str,
    sheet_filters: list[dict[str, str]],
    fq_name: str,
    fv_name: str,
    theme_prefix: str = "",
) -> str:
    del sheet_filters, fv_name
    filter_args = _live_filter_criteria_args(theme_prefix, fq_name)
    if metric == "Mean":
        return (
            f"=IFERROR(AVERAGEIFS({data_name},"
            f"{filter_args}),0)"
        )
    if metric == "Min":
        return (
            f"=IFERROR(MINIFS({data_name},"
            f"{filter_args}),0)"
        )
    if metric == "Max":
        return (
            f"=IFERROR(MAXIFS({data_name},"
            f"{filter_args}),0)"
        )
    return "0"


def _numeric_formula_cache_value(metric: str, payload: dict[str, Any]) -> float:
    key = metric.lower()
    if metric == "Min":
        key = "min_val"
    elif metric == "Max":
        key = "max_val"
    try:
        return float(payload.get(key, 0.0))
    except (TypeError, ValueError):
        return 0.0


def _numeric_allocation_option_label(question: Any, option_id: str) -> str:
    grid_row_labels = getattr(question, "grid_row_labels", None) or {}
    if option_id in grid_row_labels:
        return str(grid_row_labels[option_id])
    option_map = getattr(question, "option_map", None) or {}
    if option_id in option_map:
        return str(option_map[option_id])
    for key, value in option_map.items():
        if str(key) == option_id:
            return str(value)
    return option_id


def _numeric_result_cache_value(metric: str, result: NumericResult) -> float:
    if metric == "Mean":
        return float(result.mean)
    if metric == "Min":
        return float(result.min_val)
    if metric == "Max":
        return float(result.max_val)
    return 0.0


def _build_numeric_count_formula(
    data_name: str,
    sheet_filters: list[dict[str, str]],
    fq_name: str,
    fv_name: str,
    theme_prefix: str = "",
) -> str:
    del sheet_filters, fv_name
    filter_args = _live_filter_criteria_args(theme_prefix, fq_name)
    return (
        f'=COUNTIFS({data_name},"<>",'
        f"{filter_args})"
    )


def _live_filter_criteria_pairs(
    theme_prefix: str,
    fq_name: str,
) -> list[tuple[str, str]]:
    pairs = [
        (PASS_WORKBOOK_FILTERS_DATA_NAME, "1"),
        (PASS_WORKBOOK_CUSTOM_FILTERS_DATA_NAME, "1"),
    ]
    if theme_prefix:
        pairs.append((_theme_local_pass_data_name(theme_prefix), "1"))
    if fq_name:
        pairs.append((_per_question_pass_data_name(fq_name), "1"))
    return pairs


def _live_filter_criteria_args(theme_prefix: str, fq_name: str) -> str:
    return ",".join(
        f"{range_expr},{criterion_expr}"
        for range_expr, criterion_expr in _live_filter_criteria_pairs(
            theme_prefix,
            fq_name,
        )
    )


def _theme_local_pass_data_name(theme_prefix: str) -> str:
    return _safe_defined_name(f"{theme_prefix}_passes_local_filters_data")


def _per_question_pass_data_name(fq_name: str) -> str:
    prefix = fq_name[:-2] if fq_name.endswith("_Q") else fq_name
    return _safe_defined_name(f"{prefix}_passes_per_q_filter_data")


def _countifs_criteria(criteria: Any) -> str:
    if criteria == "<>":
        return '"<>"'
    return _excel_criteria(criteria)


def _live_column_specs(schema: SurveySchema) -> list[_LiveColumnSpec]:
    columns: list[_LiveColumnSpec] = []
    used_headers: set[str] = {"respondent_id"}
    for question in schema.analysis_eligible_questions():
        if question.question_type in {
            QuestionType.SINGLE_SELECT,
            QuestionType.DEMOGRAPHIC_OR_SEGMENT,
        }:
            header = _unique_live_header(question.canonical_id, used_headers)
            columns.append(
                _LiveColumnSpec(
                    key=question.canonical_id,
                    header=header,
                    data_name=_column_data_name(header),
                    question=question,
                    source_column=question.raw_columns[0] if question.raw_columns else question.canonical_id,
                    kind="single",
                )
            )
        elif question.question_type is QuestionType.MULTI_SELECT_BINARY:
            for source_column in question.raw_columns:
                header = _unique_live_header(source_column, used_headers)
                columns.append(
                    _LiveColumnSpec(
                        key=source_column,
                        header=header,
                        data_name=_column_data_name(header),
                        question=question,
                        source_column=source_column,
                        kind="multi_select",
                    )
                )
        elif question.question_type in {
            QuestionType.GRID_SINGLE_SELECT,
            QuestionType.GRID_RATED,
            QuestionType.GRID_BINARY_SELECT,
        }:
            for source_column in question.raw_columns:
                header = _unique_live_header(source_column, used_headers)
                columns.append(
                    _LiveColumnSpec(
                        key=source_column,
                        header=header,
                        data_name=_column_data_name(header),
                        question=question,
                        source_column=source_column,
                        kind=(
                            "grid_binary"
                            if question.question_type is QuestionType.GRID_BINARY_SELECT
                            else "grid_single"
                        ),
                    )
                )
        elif question.question_type in {
            QuestionType.NUMERIC_ALLOCATION,
            QuestionType.RANK_ORDER,
            QuestionType.NPS,
        }:
            if question.question_type is QuestionType.NPS:
                header = _unique_live_header(f"{question.canonical_id}_NPS_Bucket", used_headers)
                columns.append(
                    _LiveColumnSpec(
                        key=_nps_filter_column_key(question.canonical_id),
                        header=header,
                        data_name=_column_data_name(header),
                        question=question,
                        source_column=None,
                        kind="nps_bucket",
                    )
                )
            for source_column in question.raw_columns:
                header = _unique_live_header(source_column, used_headers)
                columns.append(
                    _LiveColumnSpec(
                        key=source_column,
                        header=header,
                        data_name=_column_data_name(header),
                        question=question,
                        source_column=source_column,
                        kind="numeric",
                    )
                )
        elif question.question_type is QuestionType.DIRECT_NUMERIC:
            header = _unique_live_header(question.canonical_id, used_headers)
            columns.append(
                _LiveColumnSpec(
                    key=question.canonical_id,
                    header=header,
                    data_name=_column_data_name(header),
                    question=question,
                    source_column=question.raw_columns[0] if question.raw_columns else question.canonical_id,
                    kind="numeric",
                )
            )
    return columns


_RAW_DATA_NUMERIC_METRIC_TYPES = {
    QuestionType.DIRECT_NUMERIC,
    QuestionType.GRID_RATED,
    QuestionType.NUMERIC_ALLOCATION,
    QuestionType.RANK_ORDER,
}


def _raw_data_column_requires_numeric_cells(column: _LiveColumnSpec) -> bool:
    question = column.question
    if question is None:
        return False
    question_type = getattr(question, "question_type", None)
    if question_type in _RAW_DATA_NUMERIC_METRIC_TYPES:
        return True
    return (
        question_type is QuestionType.GRID_SINGLE_SELECT
        and str(getattr(question, "possible_role", "") or "").upper() == GRID_RATED
    )


def _raw_data_numeric_metric_value(value: Any) -> int | float | None:
    if not _live_value_present(value) or isinstance(value, bool):
        return None
    if isinstance(value, (int, float)):
        number = float(value)
    else:
        text = str(value).strip().replace(",", "")
        try:
            number = float(text)
        except (TypeError, ValueError):
            return None
    if not math.isfinite(number):
        return None
    if number.is_integer():
        return int(number)
    return number


def _raw_rows_from_dataframe(
    decoded_df: Any | None,
    schema: SurveySchema,
    columns: list[_LiveColumnSpec],
) -> list[dict[str, Any]]:
    if decoded_df is None:
        return []
    df_columns = set(getattr(decoded_df, "columns", []))
    rows: list[dict[str, Any]] = []
    respondent_column = schema.respondent_id_column
    for row_index, (_idx, source_row) in enumerate(decoded_df.iterrows(), start=1):
        row_payload: dict[str, Any] = {
            "respondent_id": (
                source_row[respondent_column]
                if respondent_column in df_columns
                else row_index
            )
        }
        for column in columns:
            question = column.question
            source_col = column.source_column
            if question is None:
                row_payload[column.key] = None
                continue
            if column.kind == "nps_bucket":
                row_payload[column.key] = _nps_bucket_for_row(source_row, question)
                continue
            if source_col not in df_columns:
                row_payload[column.key] = None
                continue
            raw_value = source_row[source_col]
            if _raw_data_column_requires_numeric_cells(column):
                row_payload[column.key] = _raw_data_numeric_metric_value(raw_value)
            elif column.kind == "single":
                if getattr(question, "label_to_numeric_value", None):
                    row_payload[column.key] = None if not _live_value_present(raw_value) else raw_value
                else:
                    row_payload[column.key] = _decode_option_value(raw_value, question.option_map)
            elif column.kind == "multi_select":
                row_payload[column.key] = "Selected" if _is_selected_value(raw_value) else None
            elif column.kind == "grid_binary":
                row_payload[column.key] = "Selected" if _is_selected_value(raw_value) else None
            elif column.kind == "grid_single":
                if _grid_spec_subtype(question) == GRID_BINARY_SELECT:
                    row_payload[column.key] = "Selected" if _is_selected_value(raw_value) else None
                elif getattr(question, "label_to_numeric_value", None):
                    row_payload[column.key] = None if not _live_value_present(raw_value) else raw_value
                else:
                    row_payload[column.key] = _decode_option_value(raw_value, question.option_map)
            else:
                row_payload[column.key] = None if not _live_value_present(raw_value) else raw_value
        rows.append(row_payload)
    return rows


def _iter_raw_row_payloads(
    decoded_df: Any | None,
    schema: SurveySchema,
    results: list[SingleCutResult],
    columns: list[_LiveColumnSpec],
) -> Any:
    if decoded_df is None:
        yield from _synthetic_raw_rows(schema, results, columns)
        return

    df_columns = set(getattr(decoded_df, "columns", []))
    respondent_column = schema.respondent_id_column
    for row_index, (_idx, source_row) in enumerate(decoded_df.iterrows(), start=1):
        row_payload: dict[str, Any] = {
            "respondent_id": (
                source_row[respondent_column]
                if respondent_column in df_columns
                else row_index
            )
        }
        for column in columns:
            question = column.question
            source_col = column.source_column
            if question is None:
                row_payload[column.key] = None
                continue
            if column.kind == "nps_bucket":
                row_payload[column.key] = _nps_bucket_for_row(source_row, question)
                continue
            if source_col not in df_columns:
                row_payload[column.key] = None
                continue
            raw_value = source_row[source_col]
            if _raw_data_column_requires_numeric_cells(column):
                row_payload[column.key] = _raw_data_numeric_metric_value(raw_value)
            elif column.kind == "single":
                if getattr(question, "label_to_numeric_value", None):
                    row_payload[column.key] = None if not _live_value_present(raw_value) else raw_value
                else:
                    row_payload[column.key] = _decode_option_value(raw_value, question.option_map)
            elif column.kind == "multi_select":
                row_payload[column.key] = "Selected" if _is_selected_value(raw_value) else None
            elif column.kind == "grid_binary":
                row_payload[column.key] = "Selected" if _is_selected_value(raw_value) else None
            elif column.kind == "grid_single":
                if _grid_spec_subtype(question) == GRID_BINARY_SELECT:
                    row_payload[column.key] = "Selected" if _is_selected_value(raw_value) else None
                elif getattr(question, "label_to_numeric_value", None):
                    row_payload[column.key] = None if not _live_value_present(raw_value) else raw_value
                else:
                    row_payload[column.key] = _decode_option_value(raw_value, question.option_map)
            else:
                row_payload[column.key] = None if not _live_value_present(raw_value) else raw_value
        yield row_payload


def _synthetic_raw_rows(
    schema: SurveySchema,
    results: list[SingleCutResult],
    columns: list[_LiveColumnSpec],
) -> list[dict[str, Any]]:
    total_rows = max(schema.total_respondents, *(result.valid_n + result.missing_n for result in results))
    rows = [{"respondent_id": index + 1} for index in range(total_rows)]
    for result in results:
        if isinstance(result, SingleSelectResult):
            values: list[Any] = []
            for payload in result.distribution.values():
                values.extend([payload["label"]] * int(payload["count"]))
            _fill_synthetic_column(rows, result.question_id, values)
        elif isinstance(result, MultiSelectResult):
            for sub_column_id, payload in result.selections.items():
                values = ["Selected"] * int(payload["count"])
                _fill_synthetic_column(rows, sub_column_id, values)
        elif isinstance(result, GridSingleSelectResult):
            question = schema.get_question(result.question_id)
            subtype = _grid_spec_subtype(question)
            for sub_column_id, row_result in result.rows.items():
                if subtype == GRID_BINARY_SELECT:
                    count = sum(
                        int(payload["count"])
                        for payload in row_result.distribution.values()
                    )
                    _fill_synthetic_column(rows, sub_column_id, ["Selected"] * count)
                else:
                    values: list[Any] = []
                    for payload in row_result.distribution.values():
                        values.extend([payload["label"]] * int(payload["count"]))
                    _fill_synthetic_column(rows, sub_column_id, values)
        elif isinstance(result, NumericResult):
            if result.question_type is QuestionType.NUMERIC_ALLOCATION:
                for option_id, payload in (result.per_option_stats or {}).items():
                    valid_n = int(payload.get("valid_n", result.valid_n))
                    _fill_synthetic_column(rows, option_id, [payload.get("mean", 0.0)] * valid_n)
            else:
                _fill_synthetic_column(rows, result.question_id, [result.mean] * result.valid_n)
    return rows


def _fill_synthetic_column(rows: list[dict[str, Any]], key: str, values: list[Any]) -> None:
    for row_index, value in enumerate(values[: len(rows)]):
        rows[row_index][key] = value


def _ordered_demographic_questions(
    schema: SurveySchema,
    demo_priority: dict | None,
) -> list[Any]:
    demographics = list(schema.demographic_questions())
    if not demo_priority:
        return demographics
    by_id = {question.canonical_id: question for question in demographics}
    ordered = [
        by_id[question_id]
        for question_id in demo_priority.get("priority_ordered", [])
        if question_id in by_id
    ]
    ordered_ids = {question.canonical_id for question in ordered}
    ordered.extend(
        question for question in demographics if question.canonical_id not in ordered_ids
    )
    return ordered


def _live_font(
    bold: bool = False,
    italic: bool = False,
    size: int = 10,
    color: str | None = None,
) -> Any:
    from openpyxl.styles import Font

    return Font(bold=bold, italic=italic, size=size, color=color, name="Arial")


def _live_fill(color: str) -> Any:
    from openpyxl.styles import PatternFill

    return PatternFill(start_color=color, end_color=color, fill_type="solid")


def _live_fill_row(worksheet: Any, row_index: int, n_cols: int, color: str) -> None:
    fill = _live_fill(color)
    for col_index in range(1, n_cols + 1):
        worksheet.cell(row=row_index, column=col_index).fill = fill


def _live_header_row(worksheet: Any, row_index: int, headers: list[str]) -> None:
    for col_index, header in enumerate(headers, start=1):
        cell = worksheet.cell(row=row_index, column=col_index, value=header)
        cell.font = _live_font(bold=True)
        cell.fill = _live_fill("F2F2F2")


def _apply_distribution_conditional_formatting(
    worksheet: Any,
    data_start: int,
    data_end: int,
) -> None:
    _apply_color_scale_range(worksheet, data_start, 2, data_end, 2)
    _apply_color_scale_range(worksheet, data_start, 3, data_end, 3)


def _apply_color_scale_range(
    worksheet: Any,
    start_row: int,
    start_col: int,
    end_row: int,
    end_col: int,
) -> None:
    if end_row < start_row:
        return
    from openpyxl.formatting.rule import ColorScaleRule

    cell_range = (
        f"{_openpyxl_column_letter(start_col)}{start_row}:"
        f"{_openpyxl_column_letter(end_col)}{end_row}"
    )
    worksheet.conditional_formatting.add(
        cell_range,
        ColorScaleRule(
            start_type="min",
            start_color="F8696B",
            mid_type="percentile",
            mid_value=50,
            mid_color="FFEB84",
            end_type="max",
            end_color="63BE7B",
        ),
    )


def _write_total_respondents_row(
    worksheet: Any,
    row_index: int,
    denominator: int,
    n_cols: int,
) -> None:
    from openpyxl.styles import Border, Side

    top_border = Border(top=Side(style="thin", color="BFBFBF"))
    for col_index in range(1, n_cols + 1):
        cell = worksheet.cell(row=row_index, column=col_index)
        cell.fill = _live_fill("F2F2F2")
        cell.font = _live_font(bold=True)
        cell.border = top_border
    worksheet.cell(row=row_index, column=1, value="Total respondents")
    worksheet.cell(row=row_index, column=2, value=denominator)


def _write_total_responses_row(
    worksheet: Any,
    row_index: int,
    total_responses: int,
    n_cols: int,
) -> None:
    for col_index in range(1, n_cols + 1):
        cell = worksheet.cell(row=row_index, column=col_index)
        cell.fill = _live_fill("F2F2F2")
        cell.font = _live_font(bold=True)
    worksheet.cell(row=row_index, column=1, value="Total responses")
    worksheet.cell(row=row_index, column=2, value=total_responses)


def _live_autofit(worksheet: Any, max_width: int = 60) -> None:
    for column_cells in worksheet.columns:
        max_len = 0
        column_letter = column_cells[0].column_letter
        for cell in column_cells:
            if cell.value is not None:
                max_len = max(max_len, len(str(cell.value)))
        worksheet.column_dimensions[column_letter].width = min(max_len + 2, max_width)


def _live_set_theme_column_widths(worksheet: Any) -> None:
    widths = {
        "A": 38,
        "B": 18,
        "C": 12,
        "D": 14,
        "E": 4,
        "F": 22,
        "G": 14,
        "H": 14,
        "I": 14,
        "J": 14,
        "K": 14,
        "L": 14,
        "M": 14,
        "N": 14,
        "O": 14,
    }
    for column_letter, width in widths.items():
        worksheet.column_dimensions[column_letter].width = width


def _live_set_filter_column_widths(worksheet: Any) -> None:
    widths = {
        "A": 52,
        "B": 28,
        "C": 28,
        "D": 28,
        "E": 34,
        "F": 28,
        "G": 28,
    }
    for column_letter, width in widths.items():
        worksheet.column_dimensions[column_letter].width = width


def _wrapped_formula(value_name: str) -> str:
    return f'="|" & SUBSTITUTE({value_name}, ", ", "|") & "|"'


def _available_values_for_column(
    column: _LiveColumnSpec,
    context: _LiveWorkbookContext,
    max_values: int = 15,
) -> str:
    return ", ".join(str(value) for value in _option_values_for_column(column, context.schema, context)[:max_values])


def _available_values_for_data_name(
    data_name: str,
    context: _LiveWorkbookContext,
    max_values: int = 15,
) -> str:
    for column in context.columns:
        if column.data_name == data_name:
            return _available_values_for_column(column, context, max_values=max_values)
    return ""


def _filter_display_label(question: Any, context: _LiveWorkbookContext) -> str:
    label = context.short_labels_map.get(getattr(question, "canonical_id", ""), "")
    question_text = str(getattr(question, "question_text", question))
    if isinstance(label, str) and label.strip():
        cleaned = _strip_question_prefix(label.strip())
        if _label_drops_terminal_question_word(cleaned, question_text):
            return _strip_question_prefix(question_text).title()
        if cleaned and not _looks_like_fallback_question_label(cleaned):
            return cleaned.title()

    theme = getattr(question, "theme", None)
    if isinstance(theme, str) and 1 <= len(theme.split()) <= 3 and len(theme) <= 30:
        return theme.strip().title()

    text = _strip_question_prefix(question_text)
    lowered = text.lower()
    if "country" in lowered:
        return "Country"
    if "industry" in lowered or "sector" in lowered:
        return "Industry"
    if "region" in lowered or "geograph" in lowered:
        return "Region"
    if (
        "company size" in lowered
        or "organization size" in lowered
        or "organisation size" in lowered
        or "employees" in lowered
        or "headcount" in lowered
        or ("size" in lowered and ("organization" in lowered or "organisation" in lowered))
    ):
        return "Company Size"
    if "revenue" in lowered:
        return "Revenue"
    if "function" in lowered or "department" in lowered:
        return "Function"
    if "role" in lowered or "seniority" in lowered:
        return "Role"

    words = re.findall(r"[A-Za-z0-9&%]+", text)
    if not words:
        return "Filter"
    return " ".join(words[:5]).title()


def _looks_like_fallback_question_label(label: str) -> bool:
    return bool(
        re.match(
            r"^\s*(which|if|approximately|how|what|during|in\s+which)\b",
            label,
            flags=re.IGNORECASE,
        )
    )


def _unique_filter_option_labels(
    columns: list[_LiveColumnSpec],
    context: _LiveWorkbookContext,
) -> dict[str, str]:
    labels: dict[str, str] = {}
    used: set[str] = {"(None)", "(Inherit)"}
    for column in columns:
        base = _filter_option_label_for_column(column, context)
        candidate = base
        suffix = 2
        while candidate in used:
            candidate = f"{base} {suffix}"
            suffix += 1
        labels[column.key] = candidate
        used.add(candidate)
    return labels


def _filter_option_label_for_column(
    column: _LiveColumnSpec,
    context: _LiveWorkbookContext,
) -> str:
    question = column.question
    if question is None:
        return _strip_question_prefix(column.header).replace("_", " ").title()
    if column.kind in {"multi_select", "grid_single", "numeric"} and column.source_column:
        if question.option_map and column.source_column in question.option_map:
            return str(question.option_map[column.source_column])
        if question.grid_row_labels and column.source_column in question.grid_row_labels:
            return str(question.grid_row_labels[column.source_column])
    return _filter_display_label(question, context)


def _question_dropdown_label(question: Any, context: _LiveWorkbookContext) -> str:
    return f"{_question_display_id(question)} - {_short_label_for_question(question, context)}"


def _question_data_name_formula(question_name: str, value_name: str | None = None) -> str:
    fallback = (
        f'IFERROR(INDEX(All_Questions_Data_Names,'
        f'MATCH({question_name},All_Questions,0)),'
        f'"{_sheet_range_reference("_RawData", "$A:$A")}")'
    )
    if not value_name:
        return fallback
    return (
        f'IFERROR(INDEX(All_Questions_Value_Data_Names,'
        f'MATCH({question_name}&"|"&{value_name},All_Questions_Value_Keys,0)),'
        f"{fallback})"
    )


def _question_filter_criteria_formula(question_name: str, value_name: str) -> str:
    return (
        f'IFERROR(INDEX(All_Questions_Value_Criteria,'
        f'MATCH({question_name}&"|"&{value_name},All_Questions_Value_Keys,0)),'
        f"{value_name})"
    )


def _cross_tab_data_name_formula(question_name: str) -> str:
    return (
        f'IFERROR(INDEX(Cross_Tab_Questions_Data_Names,'
        f'MATCH({question_name},Cross_Tab_Questions,0)),'
        f'"{_sheet_range_reference("_RawData", "$A:$A")}")'
    )


def _question_options_name_formula(question_name: str) -> str:
    return (
        f'IFERROR(INDEX(All_Questions_Options_Names,'
        f'MATCH({question_name},All_Questions,0)),'
        f'"{_sheet_range_reference("_Options", "$E$2:$E$2")}")'
    )


def _cross_tab_options_name_formula(question_name: str) -> str:
    return (
        f'IFERROR(INDEX(Cross_Tab_Questions_Options_Names,'
        f'MATCH({question_name},Cross_Tab_Questions,0)),'
        f'"{_sheet_range_reference("_Options", "$E$2:$E$2")}")'
    )


def _question_heading_text(question: Any, short_labels: dict[str, str]) -> str:
    del short_labels
    text = str(getattr(question, "question_text", "")).strip()
    display_id = _question_display_id(question)
    if not text:
        return display_id
    if re.match(rf"^\s*{re.escape(display_id)}\b", text, flags=re.IGNORECASE):
        return text
    return f"{display_id} - {text}"


def _short_label_for_question(question: Any, context: _LiveWorkbookContext) -> str:
    question_text = str(getattr(question, "question_text", ""))
    cleaned_question_text = _strip_question_prefix(question_text)
    canonical_id = getattr(question, "canonical_id", "")
    explicit_label = context.explicit_short_labels.get(canonical_id, "")
    if isinstance(explicit_label, str) and explicit_label.strip():
        cleaned = _strip_question_prefix(explicit_label.strip())
        if cleaned:
            return cleaned.title()
    label = context.short_labels_map.get(canonical_id, "")
    if isinstance(label, str) and label.strip():
        cleaned = _strip_question_prefix(label.strip())
        if _label_drops_terminal_question_word(cleaned, question_text):
            return cleaned_question_text.title()
        if cleaned and _label_is_contiguous_question_phrase(cleaned, cleaned_question_text):
            return cleaned.title()
    return (cleaned_question_text or _short_question_label(question_text)).title()


def _label_drops_terminal_question_word(label: str, question_text: str) -> bool:
    original = _strip_question_prefix(question_text)
    original_norm = _normalised_label_words(original)
    label_norm = _normalised_label_words(label)
    return bool(
        original_norm.endswith(" question")
        and label_norm
        and label_norm == original_norm[: -len(" question")]
    )


def _label_diverges_from_question_text(label: str, question_text: str) -> bool:
    label_words = _normalised_label_words(label).split()
    text_words = _normalised_label_words(question_text).split()
    if not label_words or not text_words:
        return False
    if len(label_words) <= 3:
        return False
    text_set = set(text_words)
    return not all(word in text_set for word in label_words)


def _label_is_contiguous_question_phrase(label: str, question_text: str) -> bool:
    label_words = _normalised_label_words(label).split()
    text_words = _normalised_label_words(question_text).split()
    if not label_words or len(label_words) > len(text_words):
        return False
    label_len = len(label_words)
    return any(
        text_words[index : index + label_len] == label_words
        for index in range(0, len(text_words) - label_len + 1)
    )


def _normalised_label_words(text: str) -> str:
    return " ".join(re.findall(r"[A-Za-z0-9&%]+", str(text).lower()))


def _question_display_id(question: Any) -> str:
    raw_id = str(
        getattr(
            question,
            "question_id_raw",
            getattr(question, "question_id", getattr(question, "canonical_id", "")),
        )
    ).strip()
    if raw_id.startswith("[") and raw_id.endswith("]"):
        raw_id = raw_id[1:-1].strip()
    return raw_id or str(getattr(question, "canonical_id", "Question"))


def _strip_question_prefix(text: str) -> str:
    return _QUESTION_PREFIX_PATTERN.sub("", text).strip()


def _subset_denominator_note(
    result: SingleCutResult,
    question: Any,
    schema: SurveySchema,
) -> str:
    denominator = int(getattr(result, "valid_n", 0))
    if denominator >= 0.9 * schema.total_respondents:
        return ""

    conditional_on = getattr(question, "conditional_on", None)
    if conditional_on:
        referenced = schema.get_question(conditional_on)
        if referenced is not None:
            label = _short_question_label(referenced.question_text)
            return (
                "Note: This question was shown only to respondents who answered "
                f"{label}. Total respondents shown: {denominator:,}."
            )
    return (
        "Note: This question was shown to a subset. "
        f"Total respondents shown: {denominator:,}."
    )


def _short_question_label(question_text: str, max_words: int = 5) -> str:
    text = _strip_question_prefix(question_text)
    words = re.findall(r"[A-Za-z0-9&%]+", text.lower())
    if not words:
        return "prior question"
    return " ".join(words[:max_words])


def _theme_defined_prefix(sheet_name: str) -> str:
    return _safe_defined_name(sheet_name)[:80]


def _add_comment(cell: Any, text: str) -> None:
    if not text:
        return
    from openpyxl.comments import Comment

    cell.comment = Comment(text=str(text), author="Survey Analysis Engine")


def _add_dropdown_to_cell(
    worksheet: Any,
    cell_ref: str,
    source_range: str,
    allow_blank: bool = True,
) -> None:
    from openpyxl.worksheet.datavalidation import DataValidation

    validation = DataValidation(
        type="list",
        formula1=source_range,
        allow_blank=allow_blank,
    )
    validation.add(cell_ref)
    worksheet.add_data_validation(validation)


def _add_named_cell(workbook: Any, name: str, worksheet: Any, cell_ref: str) -> None:
    from openpyxl.utils.cell import absolute_coordinate

    _add_named_range(workbook, name, worksheet.title, absolute_coordinate(cell_ref))


def _add_named_range(workbook: Any, name: str, sheet_name: str, range_ref: str) -> None:
    from openpyxl.workbook.defined_name import DefinedName

    safe_name = _safe_defined_name(name)
    attr_text = f"{_quote_openpyxl_sheet(sheet_name)}!{range_ref}"
    if safe_name in workbook.defined_names:
        del workbook.defined_names[safe_name]
    workbook.defined_names.add(DefinedName(safe_name, attr_text=attr_text))


def _sheet_range_reference(sheet_name: str, range_ref: str) -> str:
    return f"{_quote_openpyxl_sheet(sheet_name)}!{range_ref}"


def _defined_name_reference(workbook: Any, name: str) -> str:
    safe_name = _safe_defined_name(name)
    if safe_name in workbook.defined_names:
        return str(workbook.defined_names[safe_name].attr_text)
    return name


def _defined_name_keys(workbook: Any) -> set[str]:
    defined: set[str] = set()
    keys = getattr(getattr(workbook, "defined_names", None), "keys", None)
    if callable(keys):
        defined.update(str(name) for name in keys())
    for worksheet in getattr(workbook, "worksheets", []):
        try:
            local_keys = getattr(getattr(worksheet, "defined_names", None), "keys", None)
            if callable(local_keys):
                defined.update(str(name) for name in local_keys())
        except Exception:
            continue
    return defined


def _formula_without_string_literals(formula: str) -> str:
    return _FORMULA_STRING_LITERAL_PATTERN.sub("", formula)


def _validate_no_undefined_names(workbook: Any, *, strict: bool = False) -> list[str]:
    defined = _defined_name_keys(workbook)
    missing: set[str] = set()
    for worksheet in getattr(workbook, "worksheets", []):
        for row in worksheet.iter_rows():
            for cell in row:
                value = cell.value
                if not (isinstance(value, str) and value.startswith("=")):
                    continue
                formula = _formula_without_string_literals(value)
                for token in _FORMULA_DEFINED_NAME_PATTERN.findall(formula):
                    if token not in defined:
                        missing.add(token)
    missing_names = sorted(missing)
    if strict and missing_names:
        raise ValueError(
            "Workbook formulas reference undefined named ranges: "
            + ", ".join(missing_names)
        )
    return missing_names


def _append_formula_validation_warnings(workbook: Any, missing_names: list[str]) -> None:
    if "Warnings" in workbook.sheetnames:
        worksheet = workbook["Warnings"]
    else:
        worksheet = workbook.create_sheet("Warnings")
        _live_header_row(worksheet, 1, ["Source", "Warning"])
    row_index = max(worksheet.max_row + 1, 2)
    for name in missing_names:
        worksheet.cell(row=row_index, column=1, value="formula_validation")
        worksheet.cell(
            row=row_index,
            column=2,
            value=f"Undefined named range referenced by formulas: {name}",
        )
        row_index += 1


def _validate_formula_names_for_export(workbook: Any, *, strict: bool = False) -> list[str]:
    missing_names = _validate_no_undefined_names(workbook, strict=strict)
    if missing_names:
        logging.warning(
            "Workbook formulas reference undefined named ranges: %s",
            ", ".join(missing_names),
        )
        _append_formula_validation_warnings(workbook, missing_names)
    return missing_names


def _replace_streamed_raw_data_sheet(output_path: str, workbook: Any) -> None:
    raw_sheet_xml = getattr(workbook, _RAW_STREAM_XML_ATTR, None)
    if not raw_sheet_xml:
        return

    path = Path(output_path)
    if not path.exists():
        return

    with ZipFile(path, "r") as archive:
        names = set(archive.namelist())
        if "xl/workbook.xml" not in names or "xl/_rels/workbook.xml.rels" not in names:
            return
        sheet_paths = _sheet_xml_paths(
            archive.read("xl/workbook.xml"),
            archive.read("xl/_rels/workbook.xml.rels"),
        )
        raw_path = sheet_paths.get("_RawData", ("", ""))[1]
        if not raw_path:
            return
        archive_entries = [
            (item, archive.read(item.filename))
            for item in archive.infolist()
        ]

    with ZipFile(path, "w", ZIP_DEFLATED) as patched_archive:
        for item, original_data in archive_entries:
            data = raw_sheet_xml if item.filename == raw_path else original_data
            patched_archive.writestr(item, data)


def _write_formula_caches(output_path: str, workbook: Any) -> None:
    """Patch cached formula results and calcChain into openpyxl output."""

    from src import memory_profiler as export_memory

    formula_cache: dict[str, dict[str, Any]] = getattr(
        workbook, _FORMULA_CACHE_ATTR, {}
    )
    path = Path(output_path)
    if not path.exists():
        return

    ET.register_namespace("", _SPREADSHEET_NS)
    ET.register_namespace("r", _RELATIONSHIP_NS)

    temp_xlsx = tempfile.NamedTemporaryFile(
        suffix=".xlsx",
        dir=str(path.parent),
        delete=False,
    )
    temp_xlsx_path = Path(temp_xlsx.name)
    temp_xlsx.close()
    calc_refs_file = tempfile.NamedTemporaryFile(
        suffix=".calc_refs",
        dir=str(path.parent),
        delete=False,
    )
    calc_refs_path = Path(calc_refs_file.name)

    try:
        with ZipFile(path, "r") as archive:
            names = set(archive.namelist())
            if "xl/workbook.xml" not in names or "xl/_rels/workbook.xml.rels" not in names:
                return

            workbook_xml = archive.read("xl/workbook.xml")
            workbook_rels_xml = archive.read("xl/_rels/workbook.xml.rels")
            sheet_paths = _sheet_xml_paths(workbook_xml, workbook_rels_xml)
            path_to_sheet = {
                sheet_path: (sheet_name, sheet_id)
                for sheet_name, (sheet_id, sheet_path) in sheet_paths.items()
            }

            with ZipFile(temp_xlsx_path, "w", ZIP_DEFLATED) as patched_archive:
                with export_memory.memory_step("patch_formula_caches"):
                    for item in archive.infolist():
                        filename = item.filename
                        if filename == "xl/calcChain.xml":
                            continue
                        if filename == "[Content_Types].xml":
                            patched_archive.writestr(
                                item,
                                _ensure_calc_chain_content_type(
                                    archive.read(filename)
                                ),
                            )
                            continue
                        if filename == "xl/_rels/workbook.xml.rels":
                            patched_archive.writestr(
                                item,
                                _ensure_calc_chain_relationship(workbook_rels_xml),
                            )
                            continue
                        if filename in path_to_sheet:
                            sheet_name, sheet_id = path_to_sheet[filename]
                            sheet_cache = formula_cache.pop(sheet_name, {})
                            with archive.open(item, "r") as source:
                                with patched_archive.open(item, "w") as target:
                                    _patch_sheet_formula_cache_values_stream(
                                        source,
                                        target,
                                        sheet_cache,
                                        sheet_id,
                                        calc_refs_file,
                                    )
                            del sheet_cache
                            continue

                        _copy_zip_entry_stream(archive, patched_archive, item)

                    formula_cache.clear()
                    del formula_cache
                    gc.collect()

                with export_memory.memory_step("write_calc_chain"):
                    calc_refs_file.flush()
                    calc_refs_file.close()
                    _write_calc_chain_xml_stream(patched_archive, calc_refs_path)

        _replace_output_file(temp_xlsx_path, path)
    finally:
        calc_refs_file.close()
        for temp_path in (temp_xlsx_path, calc_refs_path):
            try:
                if temp_path.exists():
                    temp_path.unlink()
            except OSError:
                pass


def _replace_output_file(temp_path: Path, output_path: Path) -> None:
    try:
        os.replace(temp_path, output_path)
        return
    except PermissionError:
        # Some synced Windows folders deny atomic replace even after both zip
        # handles are closed. Keep the normal atomic path above, but fall back
        # to a streamed overwrite so local exports and tests can still finish.
        pass

    with temp_path.open("rb") as source:
        with output_path.open("wb") as target:
            while True:
                chunk = source.read(_XML_COPY_CHUNK_SIZE)
                if not chunk:
                    break
                target.write(chunk)


def _write_memory_report_if_enabled(output_path: str, profiler: Any) -> None:
    if not profiler.is_profiling_enabled():
        return

    report = profiler.get_report()
    report_dir = Path("outputs")
    report_dir.mkdir(parents=True, exist_ok=True)
    report_path = report_dir / f"{Path(output_path).stem}.memory_report.txt"
    report_path.write_text(report, encoding="utf-8")
    print(report, file=sys.stderr)


def _copy_zip_entry_stream(source_archive: ZipFile, target_archive: ZipFile, item: Any) -> None:
    with source_archive.open(item, "r") as source:
        with target_archive.open(item, "w") as target:
            while True:
                chunk = source.read(_XML_COPY_CHUNK_SIZE)
                if not chunk:
                    break
                target.write(chunk)


def _sheet_xml_paths(
    workbook_xml: bytes,
    workbook_rels_xml: bytes,
) -> dict[str, tuple[str, str]]:
    workbook_root = ET.fromstring(workbook_xml)
    rels_root = ET.fromstring(workbook_rels_xml)
    rel_targets = {
        rel.attrib["Id"]: rel.attrib.get("Target", "")
        for rel in rels_root.findall(f"{{{_PACKAGE_REL_NS}}}Relationship")
    }
    sheet_paths: dict[str, tuple[str, str]] = {}
    for sheet in workbook_root.findall(f".//{{{_SPREADSHEET_NS}}}sheet"):
        rel_id = sheet.attrib.get(f"{{{_RELATIONSHIP_NS}}}id", "")
        target = rel_targets.get(rel_id, "")
        if not target:
            continue
        if target.startswith("/"):
            sheet_path = target.lstrip("/")
        elif target.startswith("xl/"):
            sheet_path = target
        else:
            sheet_path = f"xl/{target}"
        sheet_paths[sheet.attrib["name"]] = (
            str(sheet.attrib.get("sheetId", len(sheet_paths) + 1)),
            sheet_path,
        )
    return sheet_paths


def _patch_sheet_formula_cache_values_stream(
    source: Any,
    target: Any,
    cache_values: dict[str, Any],
    sheet_id: str,
    calc_refs_file: Any,
) -> None:
    buffer = b""
    while True:
        chunk = source.read(_XML_COPY_CHUNK_SIZE)
        if not chunk:
            break
        buffer += chunk
        last_cell_end = buffer.rfind(b"</c>")
        if last_cell_end == -1:
            continue
        cutoff = last_cell_end + len(b"</c>")
        target.write(
            _patch_formula_cells_in_xml_chunk(
                buffer[:cutoff],
                cache_values,
                sheet_id,
                calc_refs_file,
            )
        )
        buffer = buffer[cutoff:]

    if buffer:
        target.write(
            _patch_formula_cells_in_xml_chunk(
                buffer,
                cache_values,
                sheet_id,
                calc_refs_file,
            )
        )


def _patch_formula_cells_in_xml_chunk(
    xml_chunk: bytes,
    cache_values: dict[str, Any],
    sheet_id: str,
    calc_refs_file: Any,
) -> bytes:
    def replace_cell(match: re.Match[bytes]) -> bytes:
        cell_xml = match.group(0)
        if b"<f" not in cell_xml:
            return cell_xml
        start_match = re.match(rb"<c\b[^>]*>", cell_xml, re.DOTALL)
        if start_match is None:
            return cell_xml
        ref_match = re.search(rb'\br="([^"]+)"', start_match.group(0))
        if ref_match is None:
            return cell_xml

        ref = ref_match.group(1).decode("ascii", errors="ignore")
        calc_refs_file.write(f"{sheet_id}\t{ref}\n".encode("ascii"))
        return _patch_formula_cell_cached_value(cell_xml, cache_values.get(ref, 0))

    return _CELL_XML_RE.sub(replace_cell, xml_chunk)


def _patch_formula_cell_cached_value(cell_xml: bytes, value: Any) -> bytes:
    cached = _normalise_formula_cache_value(value)
    start_match = re.match(rb"<c\b[^>]*>", cell_xml, re.DOTALL)
    if start_match is None:
        return cell_xml

    start_tag = start_match.group(0)
    cached_text = _formula_cache_xml_text(cached)
    if isinstance(cached, str):
        if re.search(rb'\st="[^"]*"', start_tag):
            start_tag = re.sub(rb'\st="[^"]*"', b' t="str"', start_tag, count=1)
        else:
            start_tag = start_tag[:-1] + b' t="str">'
    else:
        start_tag = re.sub(rb'\st="[^"]*"', b"", start_tag, count=1)

    patched = start_tag + cell_xml[start_match.end():]
    v_xml = b"<v>" + cached_text + b"</v>"
    if re.search(rb"<v\b[^>]*/>", patched):
        return re.sub(rb"<v\b[^>]*/>", v_xml, patched, count=1)
    if re.search(rb"<v\b[^>]*>.*?</v>", patched, re.DOTALL):
        return re.sub(rb"<v\b[^>]*>.*?</v>", v_xml, patched, count=1, flags=re.DOTALL)
    return patched.replace(b"</c>", v_xml + b"</c>", 1)


def _formula_cache_xml_text(value: Any) -> bytes:
    if isinstance(value, str):
        return escape(value).encode("utf-8")
    if isinstance(value, float) and not math.isfinite(value):
        value = 0
    return str(value).encode("ascii", errors="xmlcharrefreplace")


def _write_calc_chain_xml_stream(target_archive: ZipFile, calc_refs_path: Path) -> None:
    with target_archive.open("xl/calcChain.xml", "w") as target:
        target.write(
            (
                '<?xml version="1.0" encoding="utf-8"?>'
                f'<calcChain xmlns="{_SPREADSHEET_NS}">'
            ).encode("utf-8")
        )
        if calc_refs_path.exists():
            with calc_refs_path.open("rb") as refs:
                for line in refs:
                    line = line.strip()
                    if not line:
                        continue
                    sheet_id, ref = line.split(b"\t", 1)
                    target.write(b'<c r="')
                    target.write(ref)
                    target.write(b'" i="')
                    target.write(sheet_id)
                    target.write(b'"/>')
        target.write(b"</calcChain>")


def _ensure_calc_chain_content_type(content_types_xml: bytes) -> bytes:
    root = ET.fromstring(content_types_xml)
    override_tag = f"{{{_CONTENT_TYPES_NS}}}Override"
    for override in root.findall(override_tag):
        if override.attrib.get("PartName") == "/xl/calcChain.xml":
            return content_types_xml
    root.append(
        ET.Element(
            override_tag,
            {
                "PartName": "/xl/calcChain.xml",
                "ContentType": _CALC_CHAIN_CONTENT_TYPE,
            },
        )
    )
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def _ensure_calc_chain_relationship(workbook_rels_xml: bytes) -> bytes:
    root = ET.fromstring(workbook_rels_xml)
    rel_tag = f"{{{_PACKAGE_REL_NS}}}Relationship"
    existing_ids: set[str] = set()
    for rel in root.findall(rel_tag):
        existing_ids.add(rel.attrib.get("Id", ""))
        if rel.attrib.get("Type") == _CALC_CHAIN_REL_TYPE:
            return workbook_rels_xml

    next_index = 1
    while f"rId{next_index}" in existing_ids:
        next_index += 1
    root.append(
        ET.Element(
            rel_tag,
            {
                "Id": f"rId{next_index}",
                "Type": _CALC_CHAIN_REL_TYPE,
                "Target": "calcChain.xml",
            },
        )
    )
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def _openpyxl_add_table(
    worksheet: Any,
    start_row: int,
    start_col: int,
    end_row: int,
    end_col: int,
) -> None:
    from openpyxl.worksheet.table import Table, TableStyleInfo

    if end_row <= start_row:
        return
    table_name = _safe_defined_name(
        f"T_{worksheet.title}_{start_row}_{start_col}"
    )
    table_ref = (
        f"{_openpyxl_column_letter(start_col)}{start_row}:"
        f"{_openpyxl_column_letter(end_col)}{end_row}"
    )
    table = Table(displayName=table_name, ref=table_ref)
    table.tableStyleInfo = TableStyleInfo(
        name="TableStyleLight9",
        showFirstColumn=False,
        showLastColumn=False,
        showRowStripes=True,
        showColumnStripes=False,
    )
    worksheet.add_table(table)


def _openpyxl_column_letter(index: int) -> str:
    from openpyxl.utils import get_column_letter

    return get_column_letter(index)


def _quote_openpyxl_sheet(sheet_name: str) -> str:
    escaped = sheet_name.replace("'", "''")
    return f"'{escaped}'"


def _demo_filter_name(
    question: Any,
    category: str,
    used_names: set[str],
) -> str:
    if category:
        base = "F_" + "_".join(part.capitalize() for part in category.split("_"))
    else:
        base = "F_" + _safe_defined_name(question.canonical_id)
    candidate = _safe_defined_name(base)
    suffix = 2
    while candidate in used_names:
        candidate = _safe_defined_name(f"{base}_{suffix}")
        suffix += 1
    return candidate


def _safe_defined_name(name: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9_]", "_", str(name))
    cleaned = re.sub(r"_+", "_", cleaned).strip("_") or "Name"
    if not re.match(r"^[A-Za-z_]", cleaned):
        cleaned = f"N_{cleaned}"
    return cleaned[:200]


def _unique_live_header(value: str, used_headers: set[str]) -> str:
    base = _safe_excel_column_name(value)
    candidate = base
    suffix = 2
    while candidate in used_headers:
        suffix_text = f"_{suffix}"
        candidate = f"{base[:31 - len(suffix_text)]}{suffix_text}"
        suffix += 1
    used_headers.add(candidate)
    return candidate


def _safe_excel_column_name(value: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9_]", "_", str(value))
    cleaned = re.sub(r"_+", "_", cleaned).strip("_") or "Column"
    if not re.match(r"^[A-Za-z_]", cleaned):
        cleaned = f"Q_{cleaned}"
    return cleaned[:31]


def _decode_option_value(value: Any, option_map: dict[int | str, str]) -> Any:
    if not _live_value_present(value):
        return None
    candidates = [value, str(value)]
    try:
        numeric_value = float(value)
        if numeric_value.is_integer():
            candidates.append(int(numeric_value))
    except (TypeError, ValueError):
        pass
    for candidate in candidates:
        if candidate in option_map:
            return option_map[candidate]
    return value


def _nps_bucket_for_row(source_row: Any, question: Any) -> str | None:
    for source_column in getattr(question, "raw_columns", ()):
        try:
            value = source_row[source_column]
        except Exception:
            continue
        score = _coerce_nps_score(value)
        if score is None:
            continue
        if score >= 9:
            return "Promoter"
        if score >= 7:
            return "Passive"
        return "Detractor"
    return None


def _coerce_nps_score(value: Any) -> int | None:
    if not _live_value_present(value) or isinstance(value, bool):
        return None
    try:
        numeric = float(str(value).strip())
    except (TypeError, ValueError):
        return None
    if not numeric.is_integer():
        return None
    score = int(numeric)
    if 0 <= score <= 10:
        return score
    return None


def _is_selected_value(value: Any) -> bool:
    if not _live_value_present(value):
        return False
    if isinstance(value, str):
        return value.strip().lower() not in {"", "0", "false", "no", "unchecked"}
    try:
        return float(value) != 0.0
    except (TypeError, ValueError):
        return True


def _live_value_present(value: Any) -> bool:
    if value is None:
        return False
    try:
        import pandas as pd

        if pd.isna(value):
            return False
    except Exception:
        pass
    if isinstance(value, str) and not value.strip():
        return False
    return True


def _excel_criteria(value: Any) -> str:
    if isinstance(value, str):
        escaped = value.replace('"', '""')
        return f'"{escaped}"'
    return f'"{value}"'


def export_cross_cuts_only(
    cross_cut_results: list[CrossCutResult],
    schema: SurveySchema,
    log: CalculationLog,
    output_path: str,
    embed_input_files: bool = False,
    input_file_sources: dict[str, Any] | None = None,
) -> None:
    """Write a workbook containing only selected cross-cut outputs."""

    workbook = _create_workbook(output_path)
    formats = _make_formats(workbook)
    used_sheet_names: set[str] = set()
    cc_sheet_names = {
        result.cross_cut_id: _unique_sheet_name(
            f"CC_{result.cross_cut_id}", used_sheet_names
        )
        for result in cross_cut_results
    }

    _write_cross_cut_index(workbook, cross_cut_results, cc_sheet_names, schema, formats)
    for result in cross_cut_results:
        _write_cc_sheet(
            workbook,
            result,
            schema,
            formats,
            cc_sheet_names[result.cross_cut_id],
        )
    _write_calculation_log(workbook, log, formats)
    _write_filter_log(workbook, cross_cut_results, cc_sheet_names, formats)
    _embed_input_file_sheets_xlsxwriter(
        workbook,
        schema,
        formats,
        embed_input_files=embed_input_files,
        input_file_sources=input_file_sources,
    )
    workbook.close()


def export_filtered_single_cuts(
    filtered_results: list[FilteredSingleCutResult],
    schema: SurveySchema,
    log: CalculationLog,
    output_path: str,
    embed_input_files: bool = False,
    input_file_sources: dict[str, Any] | None = None,
) -> None:
    """Write a workbook of filtered single-cut analyses."""

    workbook = _create_workbook(output_path)
    formats = _make_formats(workbook)
    used_sheet_names: set[str] = set()
    fsc_sheet_names = _filtered_sheet_names(filtered_results, used_sheet_names)

    _write_filtered_run_summary(workbook, filtered_results, schema, log, formats)
    _write_filtered_cut_index(workbook, filtered_results, fsc_sheet_names, schema, formats)
    for index, result in enumerate(filtered_results):
        _write_fsc_sheet(workbook, result, schema, formats, fsc_sheet_names[index])
    _write_filtered_calculation_log(workbook, filtered_results, log, formats)
    _write_filtered_filter_log(workbook, filtered_results, fsc_sheet_names, schema, formats)
    _embed_input_file_sheets_xlsxwriter(
        workbook,
        schema,
        formats,
        embed_input_files=embed_input_files,
        input_file_sources=input_file_sources,
    )
    workbook.close()


def _create_workbook(output_path: str) -> Any:
    if xlsxwriter is not None:
        return xlsxwriter.Workbook(
            output_path,
            {"nan_inf_to_errors": True, "strings_to_formulas": False},
        )
    return _FallbackWorkbook(output_path)


def _make_formats(workbook: Any) -> dict[str, Any]:
    return {
        "bold": workbook.add_format({"bold": True}),
        "italic": workbook.add_format({"italic": True}),
        "bold_italic": workbook.add_format({"bold": True, "italic": True}),
        "header": workbook.add_format(
            {
                "bold": True,
                "bg_color": "#F2F2F2",
                "border": 1,
            }
        ),
        "pct": workbook.add_format({"num_format": "0.0%"}),
        "count": workbook.add_format({"num_format": "#,##0"}),
        "stat": workbook.add_format({"num_format": "0.00"}),
        "link": workbook.add_format({"font_color": "blue", "underline": 1}),
        "filter_header": workbook.add_format({"bold": True, "bg_color": "#F2F2F2"}),
        "filter_info": workbook.add_format({"bg_color": "#F2F2F2"}),
        "section": workbook.add_format({"bold": True, "bg_color": "#E8E8E8"}),
        "red_header": workbook.add_format(
            {"bold": True, "font_color": "#FFFFFF", "bg_color": "#CC0000"}
        ),
        "theme_header": workbook.add_format(
            {
                "bold": True,
                "font_color": "#FFFFFF",
                "bg_color": "#CC0000",
                "font_size": 12,
            }
        ),
        "theme_subtle": workbook.add_format(
            {"italic": True, "font_color": "#666666", "font_size": 9}
        ),
        "question_title": workbook.add_format({"bold": True, "font_size": 11}),
        "separator": workbook.add_format({"bottom": 1, "bottom_color": "#BFBFBF"}),
        "formula_pct": workbook.add_format({"num_format": "0.0%"}),
        "formula_count": workbook.add_format({"num_format": "#,##0"}),
    }


def _write_run_summary(
    workbook: Any,
    schema: SurveySchema,
    quality_report: DataQualityReport,
    results: list[SingleCutResult],
    skips: list[SkipRecord],
    log: CalculationLog,
    formats: dict[str, Any],
    cross_cut_results: list[CrossCutResult],
    cross_cut_skips: list[SkipRecord],
) -> None:
    ws = workbook.add_worksheet("Run_Summary")
    filter_log_entries = _filter_log_entry_count(cross_cut_results)
    rows = [
        ("Source datamap:", schema.source_datamap_path),
        ("Source raw data:", schema.source_rawdata_path),
        ("Run timestamp:", schema.parsed_at.isoformat()),
        ("Total respondents:", schema.total_respondents),
        ("Total questions:", len(schema.questions)),
        ("Results produced:", len(results)),
        ("Questions skipped:", len(skips)),
        ("Cross cuts produced:", len(cross_cut_results)),
        ("Cross cut skips:", len(cross_cut_skips)),
        ("Filter Log entries:", filter_log_entries),
        ("Audit log records:", len(log)),
        (
            "Calculation errors:",
            sum(1 for skip in skips if skip.skip_reason == "calculation_error"),
        ),
        ("Quality warnings:", len(quality_report.warnings)),
    ]
    for row_index, (label, value) in enumerate(rows):
        _write(ws, row_index, 0, label, formats["bold"])
        _write(ws, row_index, 1, value)
    ws.set_column(0, 0, 25)
    _autofit(ws)


def _write_question_metadata(
    workbook: Any, schema: SurveySchema, formats: dict[str, Any]
) -> None:
    ws = workbook.add_worksheet("Question_Metadata")
    headers = [
        "Question ID",
        "Canonical ID",
        "Question Text",
        "Type",
        "Raw Columns",
        "Options Count",
        "Analysis Eligible",
        "Exclusion Reason",
        "Parent Question",
    ]
    _write_header_row(ws, 0, headers, formats)
    for row_index, spec in enumerate(schema.questions, start=1):
        values = [
            spec.question_id,
            spec.canonical_id,
            spec.question_text,
            spec.question_type.value,
            ", ".join(spec.raw_columns),
            len(spec.option_map),
            "Yes" if spec.analysis_eligible else "No",
            spec.exclusion_reason or "",
            spec.parent_question_id or "",
        ]
        _write_row(ws, row_index, values)
    ws.freeze_panes(1, 0)
    _autofit(ws)


def _write_single_cut_index(
    workbook: Any,
    results: list[SingleCutResult],
    result_sheet_names: dict[str, str],
    result_theme_names: dict[str, str],
    formats: dict[str, Any],
) -> None:
    ws = workbook.add_worksheet("Single_Cut_Index")
    headers = [
        "Question ID",
        "Canonical ID",
        "Type",
        "Theme",
        "Sheet Name",
        "Valid N",
        "Missing N",
        "Missing %",
        "Warnings",
    ]
    _write_header_row(ws, 0, headers, formats)
    for row_index, result in enumerate(results, start=1):
        sheet_name = result_sheet_names.get(result.question_id, "")
        denominator = result.valid_n + result.missing_n
        missing_pct = result.missing_n / denominator if denominator > 0 else 0.0
        _write(ws, row_index, 0, result.question_id)
        _write(ws, row_index, 1, result.question_id)
        _write(ws, row_index, 2, result.question_type.value)
        _write(ws, row_index, 3, result_theme_names.get(result.question_id, ""))
        if sheet_name:
            _write_url(
                ws,
                row_index,
                4,
                f"internal:'{_quote_sheet_name(sheet_name)}'!A1",
                formats["link"],
                sheet_name,
            )
        _write(ws, row_index, 5, result.valid_n, formats["count"])
        _write(ws, row_index, 6, result.missing_n, formats["count"])
        _write(ws, row_index, 7, missing_pct, formats["pct"])
        _write(ws, row_index, 8, " | ".join(result.warnings))
    ws.freeze_panes(1, 0)
    _autofit(ws)


def _write_skip_log(
    workbook: Any,
    skips: list[SkipRecord],
    cross_cut_skips: list[SkipRecord],
    formats: dict[str, Any],
) -> None:
    ws = workbook.add_worksheet("Skip_Log")
    headers = ["Source", "Question/Cut ID", "Type", "Skip Reason", "Details"]
    _write_header_row(ws, 0, headers, formats)
    row_index = 1
    for skip in skips:
        _write_row(
            ws,
            row_index,
            [
                "single_cut",
                skip.canonical_id,
                skip.question_type.value,
                skip.skip_reason,
                skip.details or "",
            ],
        )
        row_index += 1
    for skip in cross_cut_skips:
        _write_row(
            ws,
            row_index,
            [
                "cross_cut",
                skip.canonical_id,
                skip.question_type.value,
                skip.skip_reason,
                skip.details or "",
            ],
        )
        row_index += 1
    ws.freeze_panes(1, 0)
    _autofit(ws)


def _reserved_sheet_names() -> set[str]:
    return {
        "Run_Summary",
        "Question_Metadata",
        "Single_Cut_Index",
        "Filters",
        "Skip_Log",
        "Calculation_Log",
        "Filter_Log",
        "Data_Quality",
        "Warnings",
        "Cross_Cut_Index",
        "_CohortDefinition",
        INPUT_RAW_SHEET_NAME,
        INPUT_DATAMAP_SHEET_NAME,
    }


def _theme_groups_for_results(
    schema: SurveySchema,
    results: list[SingleCutResult],
    themes: dict | None,
) -> list[tuple[str, list[SingleCutResult]]]:
    result_by_id = {result.question_id: result for result in results}
    assigned: set[str] = set()
    groups: list[tuple[str, list[SingleCutResult]]] = []

    if themes and isinstance(themes.get("themes"), list):
        for theme in themes["themes"]:
            theme_name = str(theme.get("name") or "Theme")
            theme_results = []
            for question_id in theme.get("question_ids", []):
                if question_id in result_by_id and question_id not in assigned:
                    theme_results.append(result_by_id[question_id])
                    assigned.add(question_id)
            if theme_results:
                groups.append((theme_name, theme_results))

    missing_results = [
        result
        for result in results
        if result.question_id not in assigned
    ]
    if missing_results:
        fallback = _fallback_theme_groups(schema, missing_results)
        for theme_name, theme_results in fallback:
            existing = next(
                (items for name, items in groups if name == theme_name),
                None,
            )
            if existing is not None:
                existing.extend(theme_results)
            else:
                groups.append((theme_name, theme_results))

    return groups or [("All Questions", list(results))]


def _fallback_theme_groups(
    schema: SurveySchema,
    results: list[SingleCutResult],
) -> list[tuple[str, list[SingleCutResult]]]:
    demographics: list[SingleCutResult] = []
    other: list[SingleCutResult] = []
    for result in results:
        question = schema.get_question(result.question_id)
        if question is not None and question.is_demographic:
            demographics.append(result)
        else:
            other.append(result)

    groups: list[tuple[str, list[SingleCutResult]]] = []
    if demographics:
        groups.append(("Demographics", demographics))
    if other:
        groups.append(("All Questions", other))
    return groups


def _write_theme_sheet(
    workbook: Any,
    theme_name: str,
    sheet_name: str,
    results: list[SingleCutResult],
    schema: SurveySchema,
    decoded_df: Any | None,
    formats: dict[str, Any],
    workbook_custom_filter_count: int = DEFAULT_WORKBOOK_CUSTOM_FILTER_COUNT,
) -> None:
    ws = workbook.add_worksheet(sheet_name)
    row_index = _write_sheet_header(ws, theme_name, formats)
    row_index = _write_demographic_filter_row(
        ws,
        row_index,
        list(schema.demographic_questions()),
        decoded_df,
        formats,
    )
    row_index = _write_custom_filter_slots(
        ws,
        row_index,
        [question.canonical_id for question in schema.analysis_eligible_questions()],
        formats,
        workbook_custom_filter_count,
    )

    for result in results:
        question = schema.get_question(result.question_id)
        if question is None:
            continue
        row_index = _write_question_block(ws, row_index, result, question, schema, formats)

    ws.freeze_panes(10, 0)
    _set_theme_column_widths(ws)
    _autofit(ws)


def _write_sheet_header(ws: Any, theme_name: str, formats: dict[str, Any]) -> int:
    _merge_range(ws, 0, 0, 0, 6, f"THEME: {theme_name}", formats["theme_header"])
    return 2


def _write_demographic_filter_row(
    ws: Any,
    start_row: int,
    demographic_questions: list[Any],
    decoded_df: Any | None,
    formats: dict[str, Any],
) -> int:
    if not demographic_questions:
        return start_row

    max_demographics = min(7, len(demographic_questions))
    _apply_light_gray_fill(ws, start_row, formats)
    _write(ws, start_row, 0, "FILTERS", formats["filter_header"])
    header_row = start_row + 1
    value_row = start_row + 2

    for index, question in enumerate(demographic_questions[:max_demographics]):
        label = f"{question.canonical_id} - {question.question_text[:30]}"
        _write(ws, header_row, index, label, formats["bold"])
        _write(ws, value_row, index, "(All)")
        _add_dropdown_validation(
            ws,
            value_row,
            index,
            ["(All)", *_demographic_values(question, decoded_df)],
        )

    _apply_autofilter(ws, header_row, value_row, max_demographics)
    return value_row + 2


def _write_custom_filter_slots(
    ws: Any,
    start_row: int,
    all_question_ids: list[str],
    formats: dict[str, Any],
    slot_count: int = DEFAULT_WORKBOOK_CUSTOM_FILTER_COUNT,
) -> int:
    _write(
        ws,
        start_row,
        0,
        "Optional custom filters (pick any question):",
        formats["theme_subtle"],
    )
    row_index = start_row + 1
    question_options = ["(Pick question)", *all_question_ids[:50]]
    for slot in range(1, _normalise_slot_count(slot_count, DEFAULT_WORKBOOK_CUSTOM_FILTER_COUNT, 5) + 1):
        _write(ws, row_index, 0, f"Filter {slot}:")
        _write(ws, row_index, 1, "(Pick question)")
        _write(ws, row_index, 2, "=")
        _write(ws, row_index, 3, "(All)")
        _add_dropdown_validation(ws, row_index, 1, question_options)
        row_index += 1
    return row_index + 1


def _write_question_block(
    ws: Any,
    start_row: int,
    result: SingleCutResult,
    question: Any,
    schema: SurveySchema,
    formats: dict[str, Any],
) -> int:
    for col in range(7):
        _write_blank(ws, start_row, col, formats["separator"])
    start_row += 1

    _merge_range(
        ws,
        start_row,
        0,
        start_row,
        6,
        f"{question.canonical_id} - {question.question_text}",
        formats["question_title"],
    )
    start_row += 1
    _write(
        ws,
        start_row,
        0,
        (
            f"Type: {question.question_type.value} | "
            f"Valid N: {result.valid_n:,} | Missing: {result.missing_n:,}"
        ),
        formats["theme_subtle"],
    )
    start_row += 2

    headers = ["Option", "Count", "%", "Denominator"]
    table_header_row = start_row
    _write_header_row(ws, table_header_row, headers, formats)
    data_start_row = table_header_row + 1
    rows_written = _write_distribution_rows(ws, data_start_row, result, question, schema, formats)
    if rows_written == 0:
        _write(ws, data_start_row, 0, "(no data)")
        _write(ws, data_start_row, 1, 0, formats["count"])
        rows_written = 1

    data_end_row = data_start_row + rows_written - 1
    first_excel_row = data_start_row + 1
    last_excel_row = data_end_row + 1
    for row_index in range(data_start_row, data_end_row + 1):
        excel_row = row_index + 1
        pct_formula = (
            f"=IFERROR(B{excel_row}/SUBTOTAL(9,B{first_excel_row}:"
            f"B{last_excel_row}),0)"
        )
        denom_formula = f"=SUBTOTAL(9,B{first_excel_row}:B{last_excel_row})"
        _write_formula(ws, row_index, 2, pct_formula, formats["formula_pct"], 0)
        _write_formula(ws, row_index, 3, denom_formula, formats["formula_count"], 0)

    _add_excel_table(ws, table_header_row, 0, data_end_row, 3, headers)
    return data_end_row + 2


def _write_distribution_rows(
    ws: Any,
    start_row: int,
    result: SingleCutResult,
    question: Any,
    schema: SurveySchema,
    formats: dict[str, Any],
) -> int:
    rows = _distribution_rows_for_result(result, question, schema)
    for offset, (label, count) in enumerate(rows):
        row_index = start_row + offset
        _write(ws, row_index, 0, label)
        _write(ws, row_index, 1, count, formats["count"] if isinstance(count, int) else formats["stat"])
    return len(rows)


def _distribution_rows_for_result(
    result: SingleCutResult,
    question: Any,
    schema: SurveySchema,
) -> list[tuple[str, int | float]]:
    if isinstance(result, SingleSelectResult):
        rows = [
            (str(payload["label"]), int(payload["count"]))
            for _code, payload in result.distribution.items()
        ]
        rows.sort(key=lambda item: item[1], reverse=True)
        return rows

    if isinstance(result, MultiSelectResult):
        rows = [
            (str(payload["label"]), int(payload["count"]))
            for _column_id, payload in result.selections.items()
            if int(payload["count"]) > 0
        ]
        rows.sort(key=lambda item: item[1], reverse=True)
        return rows

    if isinstance(result, NPSResult):
        return [
            (f"{entity.entity_label} NPS", float(entity.nps_score))
            for entity in result.entities
        ]

    if isinstance(result, GridSingleSelectResult):
        return [
            (str(row["option"]), int(row["count"]))
            for row in _grid_single_select_display_rows(result, schema)
            if int(row["count"]) > 0
        ]

    if isinstance(result, NumericResult):
        if result.question_type is QuestionType.NUMERIC_ALLOCATION:
            rows = []
            for option_id, payload in (result.per_option_stats or {}).items():
                label = question.option_map.get(option_id, option_id)
                rows.append((f"{label} mean", float(payload.get("mean", 0.0))))
                rows.append((f"{label} median", float(payload.get("median", 0.0))))
            return rows
        return [
            ("Valid N", result.valid_n),
            ("Missing N", result.missing_n),
            ("Mean", result.mean),
            ("Median", result.median),
            ("Std Dev", result.std),
            ("Min", result.min_val),
            ("Max", result.max_val),
            ("25th percentile", result.percentiles[25]),
            ("50th percentile", result.percentiles[50]),
            ("75th percentile", result.percentiles[75]),
        ]

    return []


def _demographic_values(question: Any, decoded_df: Any | None) -> list[str]:
    values: list[str] = []
    candidate_columns = [question.canonical_id, *question.raw_columns]
    if decoded_df is not None:
        columns = getattr(decoded_df, "columns", [])
        for column in candidate_columns:
            if column in columns:
                series = decoded_df[column].dropna()
                values = [str(value) for value in series.unique().tolist()]
                break
    if not values and question.option_map:
        values = [str(value) for value in question.option_map.values()]
    return _validation_values(values)


def _validation_values(values: list[str]) -> list[str]:
    cleaned: list[str] = []
    for value in values:
        text = str(value).replace(",", " ").strip()
        if text and text not in cleaned:
            cleaned.append(text[:60])
        if len(cleaned) >= 50:
            break
    return cleaned


def _add_dropdown_validation(
    ws: Any,
    row: int,
    col: int,
    values: list[str],
) -> None:
    values = _validation_values(values)
    if not values:
        values = ["(All)"]
    if hasattr(ws, "data_validation"):
        ws.data_validation(row, col, row, col, {"validate": "list", "source": values})


def _add_excel_table(
    ws: Any,
    first_row: int,
    first_col: int,
    last_row: int,
    last_col: int,
    headers: list[str],
) -> None:
    if hasattr(ws, "add_table"):
        ws.add_table(
            first_row,
            first_col,
            last_row,
            last_col,
            {
                "style": "Table Style Light 9",
                "columns": [{"header": header} for header in headers],
            },
        )
        return
    _apply_autofilter(ws, first_row, last_row, len(headers))


def _set_theme_column_widths(ws: Any) -> None:
    widths = [45, 12, 12, 14, 20, 20, 20]
    for col_index, width in enumerate(widths):
        ws.set_column(col_index, col_index, width)


def _write_filter_header_block(
    ws: Any,
    filters: list[FilterSpec] | tuple[FilterSpec, ...] | None,
    schema: SurveySchema,
    formats: dict[str, Any],
    start_row: int,
    filtered_n: int | None = None,
    total_n: int | None = None,
) -> int:
    """Write sheet-level filter documentation and return the next row."""

    if not filters:
        _apply_light_gray_fill(ws, start_row, formats)
        _write(ws, start_row, 0, "Filter", formats["filter_header"])
        _write(ws, start_row, 1, "None (full sample)", formats["filter_info"])
        return start_row + 2

    _apply_light_gray_fill(ws, start_row, formats)
    _write(ws, start_row, 0, "Filters applied", formats["filter_header"])
    row_index = start_row + 1
    excluded_n = (
        total_n - filtered_n
        if total_n is not None and filtered_n is not None and total_n >= filtered_n
        else None
    )
    for filter_spec in filters:
        _write(ws, row_index, 0, filter_spec.filter_question_id)
        _write(ws, row_index, 1, _filter_condition_text(filter_spec, schema))
        if filtered_n is not None:
            _write(ws, row_index, 2, f"n={filtered_n}")
        if excluded_n is not None:
            _write(ws, row_index, 3, f"({excluded_n} excluded)")
        row_index += 1
    return row_index + 1


def _write_filtered_distribution_section(
    ws: Any,
    filtered_result: SingleCutResult,
    schema: SurveySchema,
    formats: dict[str, Any],
    filter_description: str,
    start_row: int,
    filtered_n: int | None = None,
) -> int:
    section_n = filtered_n if filtered_n is not None else filtered_result.valid_n
    _apply_red_header_fill(ws, start_row, formats)
    _write(
        ws,
        start_row,
        0,
        f"FILTERED VIEW: {filter_description} (n={section_n:,})",
        formats["red_header"],
    )
    _write(
        ws,
        start_row + 1,
        0,
        f"n={section_n:,} respondents match filter",
        formats["italic"],
    )
    table_header_row = start_row + 3
    return _write_single_cut_result_body(
        ws,
        filtered_result,
        schema,
        formats,
        table_header_row,
    )


def _apply_light_gray_fill(ws: Any, row_num: int, formats: dict[str, Any]) -> None:
    for col in range(5):
        _write_blank(ws, row_num, col, formats["filter_info"])


def _apply_red_header_fill(ws: Any, row_num: int, formats: dict[str, Any]) -> None:
    for col in range(5):
        _write_blank(ws, row_num, col, formats["red_header"])


def _apply_autofilter(
    ws: Any,
    header_row: int,
    last_data_row: int,
    num_cols: int,
) -> None:
    if last_data_row < header_row or num_cols <= 0:
        return
    last_col = num_cols - 1
    if hasattr(ws, "autofilter"):
        ws.autofilter(header_row, 0, last_data_row, last_col)
    elif hasattr(ws, "set_autofilter"):
        ws.set_autofilter(header_row, 0, last_data_row, last_col)


def _set_single_cut_column_widths(ws: Any) -> None:
    widths = [45, 10, 10, 14, 20]
    for col_index, width in enumerate(widths):
        ws.set_column(col_index, col_index, width)


def _filter_condition_text(filter_spec: FilterSpec, schema: SurveySchema) -> str:
    if filter_spec.filter_value is None:
        return "(breakdown)"
    return f"= {_filter_option_label(filter_spec, schema)}"


def _filters_description(
    filters: tuple[FilterSpec, ...] | list[FilterSpec],
    schema: SurveySchema,
) -> str:
    if not filters:
        return "None (full sample)"
    return " | ".join(
        _filter_description(filter_spec, schema, compact_breakdown=False)
        for filter_spec in filters
    )


def _write_sc_sheet(
    workbook: Any,
    result: SingleCutResult,
    schema: SurveySchema,
    formats: dict[str, Any],
    sheet_name: str,
) -> None:
    spec = schema.get_question(result.question_id)
    question_text = spec.question_text if spec is not None else result.question_id
    ws = workbook.add_worksheet(sheet_name)

    row_index = _write_filter_header_block(ws, None, schema, formats, start_row=0)
    _write(ws, row_index, 0, "Question:", formats["bold"])
    _write(ws, row_index, 1, question_text)
    row_index += 1
    _write(ws, row_index, 0, "Type:", formats["bold"])
    _write(ws, row_index, 1, result.question_type.value)
    row_index += 1
    _write(ws, row_index, 0, "Denominator:", formats["bold"])
    _write(ws, row_index, 1, _denominator_description(result))
    row_index += 1
    _write(ws, row_index, 0, "Valid N:", formats["bold"])
    _write(ws, row_index, 1, result.valid_n, formats["count"])
    row_index += 1
    _write(ws, row_index, 0, "Missing N:", formats["bold"])
    _write(ws, row_index, 1, result.missing_n, formats["count"])
    row_index += 2

    _write(
        ws,
        row_index,
        0,
        f"FULL DISTRIBUTION (n={result.valid_n:,})",
        formats["section"],
    )
    table_header_row = row_index + 1
    _write_single_cut_result_body(ws, result, schema, formats, table_header_row)

    if result.warnings:
        last_row = _find_last_used_row(ws) + 2
        _write(ws, last_row, 0, "Warnings:", formats["bold"])
        for offset, warning in enumerate(result.warnings, start=1):
            _write(ws, last_row + offset, 0, warning)

    ws.freeze_panes(table_header_row + 1, 0)
    _autofit(ws)
    _set_single_cut_column_widths(ws)


def _write_single_select_body(
    ws: Any,
    result: SingleSelectResult,
    formats: dict[str, Any],
    start_row: int,
) -> int:
    _write_header_row(ws, start_row, ["Code", "Label", "Count", "%"], formats)
    row_index = start_row + 1
    total_count = 0
    total_rate = 0.0
    for code, payload in sorted(result.distribution.items(), key=lambda item: _sort_key(item[0])):
        total_count += int(payload["count"])
        total_rate += float(payload["rate"])
        _write(ws, row_index, 0, code)
        _write(ws, row_index, 1, payload["label"])
        _write(ws, row_index, 2, payload["count"], formats["count"])
        _write(ws, row_index, 3, payload["rate"], formats["pct"])
        row_index += 1
    total_rate = 1.0 if math.isclose(total_rate, 1.0) else total_rate
    _write(ws, row_index, 1, "Total", formats["bold"])
    _write(ws, row_index, 2, total_count, formats["count"])
    _write(ws, row_index, 3, total_rate, formats["pct"])
    _apply_autofilter(ws, start_row, row_index, 4)
    return row_index


def _is_grid_display_option(code: object, count: int, rate: float) -> bool:
    return code not in (0, "0") and (count > 0 or rate > 0.0)


def _grid_single_select_display_rows(
    result: GridSingleSelectResult,
    schema: SurveySchema,
) -> list[dict[str, Any]]:
    spec = schema.get_question(result.question_id)
    row_labels = spec.grid_row_labels if spec and spec.grid_row_labels else {}
    denominator = int(result.valid_n)
    rows: list[dict[str, Any]] = []

    for sub_column_id, row_result in result.rows.items():
        count = 0
        for code, payload in row_result.distribution.items():
            option_count = int(payload["count"])
            option_rate = float(payload["rate"])
            if _is_grid_display_option(code, option_count, option_rate):
                count += option_count
        if count == 0:
            continue

        rows.append(
            {
                "option": row_labels.get(sub_column_id, sub_column_id),
                "count": count,
                "rate": (float(count) / denominator) if denominator else 0.0,
                "denominator": denominator,
            }
        )

    rows.sort(key=lambda row: int(row["count"]), reverse=True)
    return rows


def _write_multi_select_body(
    ws: Any,
    result: MultiSelectResult,
    formats: dict[str, Any],
    start_row: int,
) -> int:
    _write_header_row(
        ws,
        start_row,
        ["Sub-Column ID", "Label", "Count selected", "Selection %"],
        formats,
    )
    row_index = start_row + 1
    sorted_items = sorted(
        result.selections.items(),
        key=lambda item: int(item[1]["count"]),
        reverse=True,
    )
    for sub_column_id, payload in sorted_items:
        if int(payload["count"]) == 0:
            continue
        _write(ws, row_index, 0, sub_column_id)
        _write(ws, row_index, 1, payload["label"])
        _write(ws, row_index, 2, payload["count"], formats["count"])
        _write(ws, row_index, 3, payload["selection_rate"], formats["pct"])
        row_index += 1
    _apply_autofilter(ws, start_row, max(start_row, row_index - 1), 4)
    row_index += 1
    _write(
        ws,
        row_index,
        0,
        f"Respondents who answered any: {result.respondents_who_answered_any}",
        formats["bold"],
    )
    return row_index


def _write_numeric_body(
    ws: Any,
    result: NumericResult,
    formats: dict[str, Any],
    start_row: int,
    schema: SurveySchema | None = None,
) -> int:
    rows = [
        ("Mean", result.mean),
        ("Median", result.median),
        ("Std Dev", result.std),
        ("Min", result.min_val),
        ("Max", result.max_val),
        ("25th percentile", result.percentiles[25]),
        ("50th percentile", result.percentiles[50]),
        ("75th percentile", result.percentiles[75]),
    ]
    row_index = start_row
    _write_header_row(ws, row_index, ["Statistic", "Value"], formats)
    row_index += 1
    for label, value in rows:
        _write(ws, row_index, 0, label)
        _write(ws, row_index, 1, value, formats["stat"])
        row_index += 1
    _apply_autofilter(ws, start_row, row_index - 1, 2)

    if result.question_type is QuestionType.NUMERIC_ALLOCATION:
        row_index += 1
        _write(ws, row_index, 0, "Per-option allocation statistics", formats["bold"])
        row_index += 1
        _write_header_row(
            ws,
            row_index,
            ["Option", "Mean", "Median", "Valid N", "Missing N"],
            formats,
        )
        row_index += 1
        spec = schema.get_question(result.question_id) if schema is not None else None
        for option_id, payload in (result.per_option_stats or {}).items():
            option_label = (
                spec.option_map.get(option_id, option_id)
                if spec is not None
                else option_id
            )
            _write(ws, row_index, 0, option_label)
            _write(ws, row_index, 1, payload["mean"], formats["stat"])
            _write(ws, row_index, 2, payload["median"], formats["stat"])
            _write(ws, row_index, 3, payload.get("valid_n", 0), formats["count"])
            _write(ws, row_index, 4, payload.get("missing_n", 0), formats["count"])
            row_index += 1
        row_index += 1
        allocation_rows = [
            ("Allocation target:", result.allocation_target),
            ("Tolerance:", result.allocation_tolerance),
            ("Excluded (out of tolerance):", result.allocation_excluded_n),
        ]
        for label, value in allocation_rows:
            _write(ws, row_index, 0, label)
            _write(ws, row_index, 1, value, formats["stat"])
            row_index += 1
    return row_index


def _write_grid_body(
    ws: Any,
    result: GridSingleSelectResult,
    schema: SurveySchema,
    formats: dict[str, Any],
    start_row: int,
) -> int:
    rows = _grid_single_select_display_rows(result, schema)
    row_index = start_row
    _write_header_row(ws, row_index, ["Option", "Count", "%", "Denominator"], formats)
    row_index += 1

    for row in rows:
        _write(ws, row_index, 0, row["option"])
        _write(ws, row_index, 1, row["count"], formats["count"])
        _write(ws, row_index, 2, row["rate"], formats["pct"])
        _write(ws, row_index, 3, row["denominator"], formats["count"])
        row_index += 1

    total_count = sum(int(row["count"]) for row in rows)
    denominator = int(result.valid_n)
    _write(ws, row_index, 0, "TOTAL (selected at least once)", formats["bold"])
    _write(ws, row_index, 1, total_count, formats["count"])
    _write(ws, row_index, 2, 1.0, formats["pct"])
    _write(ws, row_index, 3, denominator, formats["count"])
    _apply_autofilter(ws, start_row, row_index, 4)
    return row_index


def _write_cross_cut_index(
    workbook: Any,
    cross_cut_results: list[CrossCutResult],
    cc_sheet_names: dict[str, str],
    schema: SurveySchema,
    formats: dict[str, Any],
) -> None:
    ws = workbook.add_worksheet("Cross_Cut_Index")
    headers = [
        "Cross Cut ID",
        "Title",
        "Analysis Type",
        "Sheet Name",
        "Source Questions",
        "question_labels",
        "Filter",
        "Audit Records",
        "Warnings",
    ]
    _write_header_row(ws, 0, headers, formats)
    for row_index, result in enumerate(cross_cut_results, start=1):
        sheet_name = cc_sheet_names[result.cross_cut_id]
        filter_expr = (
            result.result_table.get("filter_expr", "")
            if result.analysis_type is AnalysisType.SEGMENT_PROFILE
            else ""
        )
        _write(ws, row_index, 0, result.cross_cut_id)
        _write(ws, row_index, 1, result.synthetic_question_title)
        _write(ws, row_index, 2, result.analysis_type.value)
        _write_url(
            ws,
            row_index,
            3,
            f"internal:'{_quote_sheet_name(sheet_name)}'!A1",
            formats["link"],
            sheet_name,
        )
        _write(ws, row_index, 4, ", ".join(result.source_question_ids))
        _write(ws, row_index, 5, _question_labels_text(schema, result.source_question_ids))
        _write(ws, row_index, 6, filter_expr)
        _write(ws, row_index, 7, len(result.audit_records), formats["count"])
        _write(ws, row_index, 8, " | ".join(result.warnings))
    ws.freeze_panes(1, 0)
    _autofit(ws)


def _write_filtered_run_summary(
    workbook: Any,
    filtered_results: list[FilteredSingleCutResult],
    schema: SurveySchema,
    log: CalculationLog,
    formats: dict[str, Any],
) -> None:
    ws = workbook.add_worksheet("Run_Summary")
    rows = [
        ("Workbook type:", "Filtered single cuts"),
        ("Source datamap:", schema.source_datamap_path),
        ("Source raw data:", schema.source_rawdata_path),
        ("Run timestamp:", schema.parsed_at.isoformat()),
        ("Filtered analyses included:", len(filtered_results)),
        ("Audit log records:", len(log)),
    ]
    for row_index, (label, value) in enumerate(rows):
        _write(ws, row_index, 0, label, formats["bold"])
        _write(ws, row_index, 1, value)
    ws.set_column(0, 0, 30)
    _autofit(ws)


def _write_filtered_cut_index(
    workbook: Any,
    filtered_results: list[FilteredSingleCutResult],
    fsc_sheet_names: dict[int, str],
    schema: SurveySchema,
    formats: dict[str, Any],
) -> None:
    ws = workbook.add_worksheet("Filtered_Cut_Index")
    headers = [
        "Sheet Name",
        "Target Question",
        "Filters Applied",
        "Dispatch Mode",
        "Filtered N",
        "Warnings",
    ]
    _write_header_row(ws, 0, headers, formats)
    for row_index, result in enumerate(filtered_results, start=1):
        sheet_name = fsc_sheet_names[row_index - 1]
        _write_url(
            ws,
            row_index,
            0,
            f"internal:'{_quote_sheet_name(sheet_name)}'!A1",
            formats["link"],
            sheet_name,
        )
        _write(ws, row_index, 1, result.target_question_id)
        _write(
            ws,
            row_index,
            2,
            " | ".join(
                _filter_description(filter_spec, schema, compact_breakdown=True)
                for filter_spec in result.filters_applied
            ),
        )
        _write(ws, row_index, 3, result.dispatch_mode)
        _write(ws, row_index, 4, result.filtered_n, formats["count"])
        _write(ws, row_index, 5, " | ".join(result.warnings))
    ws.freeze_panes(1, 0)
    _autofit(ws)


def _write_fsc_sheet(
    workbook: Any,
    result: FilteredSingleCutResult,
    schema: SurveySchema,
    formats: dict[str, Any],
    sheet_name: str,
) -> None:
    ws = workbook.add_worksheet(sheet_name)
    target_spec = schema.get_question(result.target_question_id)
    target_text = target_spec.question_text if target_spec else result.target_question_id

    row_index = _write_filter_header_block(
        ws,
        list(result.filters_applied),
        schema,
        formats,
        start_row=0,
        filtered_n=result.filtered_n,
        total_n=schema.total_respondents,
    )
    _write(ws, row_index, 0, "Target question:", formats["bold"])
    _write(ws, row_index, 1, target_text)
    row_index += 1
    _write(ws, row_index, 0, "Target ID:", formats["bold"])
    _write(ws, row_index, 1, result.target_question_id)
    row_index += 1
    _write(ws, row_index, 0, "Filtered N:", formats["bold"])
    _write(ws, row_index, 1, result.filtered_n, formats["count"])
    row_index += 1
    _write(ws, row_index, 0, "Dispatch mode:", formats["bold"])
    _write(ws, row_index, 1, result.dispatch_mode)
    row_index += 2

    if result.dispatch_mode == "single_cut_filtered":
        if result.single_cut_result is None:
            _write(ws, row_index, 0, "Missing filtered single-cut result", formats["bold"])
            last_row = row_index
        else:
            filter_description = _filters_description(result.filters_applied, schema)
            last_row = _write_filtered_distribution_section(
                ws,
                result.single_cut_result,
                schema,
                formats,
                filter_description,
                row_index,
                filtered_n=result.filtered_n,
            )
    elif result.dispatch_mode == "cross_cut_breakdown":
        if result.cross_cut_result is None:
            _write(ws, row_index, 0, "Missing cross-cut breakdown result", formats["bold"])
            last_row = row_index
        else:
            last_row = _write_cross_cut_result_body(
                workbook,
                ws,
                result.cross_cut_result,
                schema,
                formats,
                row_index,
                sheet_name,
            )
    else:
        _write(ws, row_index, 0, "Unsupported filtered result mode", formats["bold"])
        last_row = row_index

    if result.warnings:
        last_row += 2
        _write(ws, last_row, 0, "Warnings:", formats["bold"])
        for offset, warning in enumerate(result.warnings, start=1):
            _write(ws, last_row + offset, 0, warning)

    ws.freeze_panes(row_index + 4, 0)
    _autofit(ws)
    _set_single_cut_column_widths(ws)


def _write_single_cut_result_body(
    ws: Any,
    result: SingleCutResult,
    schema: SurveySchema,
    formats: dict[str, Any],
    start_row: int,
) -> int:
    if isinstance(result, GridSingleSelectResult):
        return _write_grid_body(ws, result, schema, formats, start_row)
    if isinstance(result, SingleSelectResult):
        return _write_single_select_body(ws, result, formats, start_row)
    if isinstance(result, MultiSelectResult):
        return _write_multi_select_body(ws, result, formats, start_row)
    if isinstance(result, NumericResult):
        return _write_numeric_body(ws, result, formats, start_row, schema)
    _write(ws, start_row, 0, "Unsupported single-cut result type", formats["bold"])
    return start_row


def _write_cross_cut_result_body(
    workbook: Any,
    ws: Any,
    result: CrossCutResult,
    schema: SurveySchema,
    formats: dict[str, Any],
    start_row: int,
    sheet_name: str,
) -> int:
    if result.analysis_type is AnalysisType.CROSS_TAB:
        return _write_cross_tab_body(
            workbook, ws, result, schema, formats, start_row, sheet_name
        )
    if result.analysis_type is AnalysisType.SEGMENT_PROFILE:
        return _write_segment_profile_body(
            workbook, ws, result, schema, formats, start_row, sheet_name
        )
    if result.analysis_type is AnalysisType.GROUP_COMPARISON:
        return _write_group_comparison_body(
            workbook, ws, result, schema, formats, start_row, sheet_name
        )
    if result.analysis_type is AnalysisType.EXPECTED_VS_REALIZED:
        return _write_expected_vs_realized_body(
            workbook, ws, result, schema, formats, start_row, sheet_name
        )
    _write(ws, start_row, 0, "Unsupported cross-cut result type", formats["bold"])
    return start_row


def _write_cc_sheet(
    workbook: Any,
    result: CrossCutResult,
    schema: SurveySchema,
    formats: dict[str, Any],
    sheet_name: str,
) -> None:
    ws = workbook.add_worksheet(sheet_name)
    filter_expr = result.result_table.get("filter_expr") or "<no filter>"
    header_rows = [
        ("Title:", result.synthetic_question_title),
        ("Analysis type:", result.analysis_type.value),
        ("Source questions:", ", ".join(result.source_question_ids)),
        ("Filter:", filter_expr),
        ("AI insight:", result.ai_insight or "<none>"),
        ("Audit records:", len(result.audit_records)),
    ]
    for row_index, (label, value) in enumerate(header_rows):
        _write(ws, row_index, 0, label, formats["bold"])
        _write(ws, row_index, 1, value)
    _write(ws, 2, 2, "question_labels", formats["bold"])
    _write(ws, 2, 3, _question_labels_text(schema, result.source_question_ids))

    body_start_row = len(header_rows) + 1
    if result.analysis_type is AnalysisType.CROSS_TAB:
        last_row = _write_cross_tab_body(
            workbook, ws, result, schema, formats, body_start_row, sheet_name
        )
    elif result.analysis_type is AnalysisType.SEGMENT_PROFILE:
        last_row = _write_segment_profile_body(
            workbook, ws, result, schema, formats, body_start_row, sheet_name
        )
    elif result.analysis_type is AnalysisType.GROUP_COMPARISON:
        last_row = _write_group_comparison_body(
            workbook, ws, result, schema, formats, body_start_row, sheet_name
        )
    elif result.analysis_type is AnalysisType.EXPECTED_VS_REALIZED:
        last_row = _write_expected_vs_realized_body(
            workbook, ws, result, schema, formats, body_start_row, sheet_name
        )
    else:
        _write(ws, body_start_row, 0, "Unsupported cross-cut result type", formats["bold"])
        last_row = body_start_row

    if result.warnings:
        last_row += 2
        _write(ws, last_row, 0, "Warnings:", formats["bold"])
        for offset, warning in enumerate(result.warnings, start=1):
            _write(ws, last_row + offset, 0, warning)

    ws.freeze_panes(body_start_row, 0)
    _autofit(ws)


def _write_cross_tab_body(
    workbook: Any,
    ws: Any,
    result: CrossCutResult,
    schema: SurveySchema,
    formats: dict[str, Any],
    start_row: int,
    sheet_name: str,
) -> int:
    result_table = result.result_table
    row_question_id = result_table.get("row_question_id", result.source_question_ids[0])
    column_question_id = result_table.get(
        "column_question_id", result.source_question_ids[1]
    )
    row_question_text = _question_text(schema, row_question_id)
    column_question_text = _question_text(schema, column_question_id)
    _merge_range(
        ws,
        start_row,
        0,
        start_row,
        3,
        f"Rows (vertical): {row_question_id}",
        formats["bold_italic"],
    )
    _merge_range(ws, start_row + 1, 0, start_row + 1, 3, row_question_text)
    _merge_range(
        ws,
        start_row + 2,
        0,
        start_row + 2,
        3,
        f"Columns (horizontal): {column_question_id}",
        formats["bold_italic"],
    )
    _merge_range(ws, start_row + 3, 0, start_row + 3, 3, column_question_text)
    orientation_label = f"↓ {row_question_id}  →  {column_question_id}"
    blocks_to_render = []
    if result.display_mode in ("counts", "both", "all"):
        blocks_to_render.append(("Counts", "counts", formats["count"]))
    if result.display_mode in ("row_pct", "both", "all"):
        blocks_to_render.append(("Row %", "row_pct", formats["pct"]))
    if result.display_mode in ("col_pct", "all"):
        blocks_to_render.append(("Column %", "column_pct", formats["pct"]))

    row_index = start_row + 5
    for block_index, (header, key, cell_format) in enumerate(blocks_to_render):
        if block_index > 0:
            row_index += 2
        block_start_row = row_index
        row_index = _write_cross_tab_matrix(
            ws,
            header,
            result_table.get(key, {}),
            result_table.get("row_label_map", {}),
            result_table.get("column_label_map", {}),
            formats,
            row_index,
            cell_format,
            orientation_label,
        )
        if header == "Counts":
            counts_matrix = result_table.get("counts", {})
            column_label_map = result_table.get("column_label_map", {})
            _add_cross_tab_chart(
                workbook,
                ws,
                sheet_name,
                block_start_row,
                len(counts_matrix),
                len(_cross_tab_column_codes(counts_matrix, column_label_map)),
                result.synthetic_question_title,
                row_question_id,
                column_question_id,
            )
    row_index += 2
    row_index = _write_cross_tab_copy_friendly(
        ws,
        result_table.get("counts", {}),
        result_table.get("row_label_map", {}),
        result_table.get("column_label_map", {}),
        formats,
        row_index,
    )
    row_index += 2
    _write(ws, row_index, 0, "Grand total:", formats["bold"])
    _write(ws, row_index, 1, result_table.get("grand_total", ""), formats["count"])
    return row_index


def _write_cross_tab_matrix(
    ws: Any,
    title: str,
    matrix: dict,
    row_label_map: dict,
    column_label_map: dict,
    formats: dict[str, Any],
    start_row: int,
    value_format: Any,
    orientation_label: str,
) -> int:
    row_codes = sorted(matrix.keys(), key=_sort_key)
    column_codes = _cross_tab_column_codes(matrix, column_label_map)
    _write(ws, start_row, 0, title, formats["bold"])
    _write(ws, start_row + 1, 0, orientation_label, formats["italic"])
    _write(ws, start_row + 1, 1, "")
    for offset, column_code in enumerate(column_codes, start=2):
        _write(ws, start_row + 1, offset, column_code, formats["header"])
        _write(
            ws,
            start_row + 2,
            offset,
            _label_for_code(column_label_map, column_code),
            formats["header"],
        )
    row_index = start_row + 3
    for row_code in row_codes:
        _write(ws, row_index, 0, row_code)
        _write(ws, row_index, 1, _label_for_code(row_label_map, row_code))
        for offset, column_code in enumerate(column_codes, start=2):
            _write(
                ws,
                row_index,
                offset,
                _nested_lookup(matrix, row_code, column_code),
                value_format,
            )
        row_index += 1
    return row_index - 1


def _write_cross_tab_copy_friendly(
    ws: Any,
    counts: dict,
    row_label_map: dict,
    column_label_map: dict,
    formats: dict[str, Any],
    start_row: int,
) -> int:
    row_codes = sorted(counts.keys(), key=_sort_key)
    column_codes = _cross_tab_column_codes(counts, column_label_map)
    _write(ws, start_row, 0, "Copy-friendly", formats["bold"])
    _write_header_row(
        ws,
        start_row + 1,
        ["Row Code", "Row Label", "Column Code", "Column Label", "Count"],
        formats,
    )
    row_index = start_row + 2
    for row_code in row_codes:
        for column_code in column_codes:
            _write(ws, row_index, 0, row_code)
            _write(ws, row_index, 1, _label_for_code(row_label_map, row_code))
            _write(ws, row_index, 2, column_code)
            _write(
                ws,
                row_index,
                3,
                _label_for_code(column_label_map, column_code),
            )
            _write(
                ws,
                row_index,
                4,
                _nested_lookup(counts, row_code, column_code),
                formats["count"],
            )
            row_index += 1
    return row_index - 1


def _add_cross_tab_chart(
    workbook: Any,
    worksheet: Any,
    sheet_name: str,
    counts_start_row: int,
    counts_n_rows: int,
    counts_n_cols: int,
    title: str,
    x_axis_label: str,
    y_axis_label: str,
) -> None:
    """Insert a clustered column chart referencing a cross-tab counts block."""

    if not _can_add_chart(workbook, worksheet) or counts_n_rows <= 0 or counts_n_cols <= 0:
        return

    chart = workbook.add_chart({"type": "column", "subtype": "clustered"})
    header_row = counts_start_row + 2
    data_first_row = counts_start_row + 3
    data_last_row = data_first_row + counts_n_rows - 1

    for col_idx in range(2, counts_n_cols + 2):
        chart.add_series(
            {
                "name": [sheet_name, header_row, col_idx],
                "categories": [sheet_name, data_first_row, 1, data_last_row, 1],
                "values": [sheet_name, data_first_row, col_idx, data_last_row, col_idx],
                "fill": {
                    "color": _CHART_COLORS[(col_idx - 2) % len(_CHART_COLORS)]
                },
            }
        )

    chart.set_title(
        {
            "name": title,
            "name_font": {
                "name": "Arial",
                "size": 12,
                "bold": True,
                "color": "#0A0A0A",
            },
        }
    )
    chart.set_x_axis(
        {
            "name": x_axis_label,
            "name_font": {"name": "Arial", "size": 10, "color": "#666666"},
            "num_font": {"name": "Arial", "size": 9},
        }
    )
    chart.set_y_axis(
        {
            "name": "Count",
            "name_font": {"name": "Arial", "size": 10, "color": "#666666"},
            "num_font": {"name": "Arial", "size": 9},
        }
    )
    chart.set_legend({"position": "right", "font": {"name": "Arial", "size": 9}})
    chart.set_size({"width": 640, "height": 400})
    chart.set_chartarea(
        {"border": {"color": "#E0E0E0"}, "fill": {"color": "#FFFFFF"}}
    )

    insert_col = counts_n_cols + 3
    worksheet.insert_chart(
        counts_start_row,
        insert_col,
        chart,
        {"x_offset": 10, "y_offset": 0},
    )


def _add_segment_chart(
    workbook: Any,
    worksheet: Any,
    sheet_name: str,
    data_start_row: int,
    n_segments: int,
    label_col: int,
    value_col: int,
    title: str,
    value_label: str,
) -> None:
    """Insert a simple column chart for segment-style tables."""

    if not _can_add_chart(workbook, worksheet) or n_segments <= 0:
        return

    chart = workbook.add_chart({"type": "column"})
    data_last_row = data_start_row + n_segments - 1
    chart.add_series(
        {
            "name": value_label,
            "categories": [
                sheet_name,
                data_start_row,
                label_col,
                data_last_row,
                label_col,
            ],
            "values": [
                sheet_name,
                data_start_row,
                value_col,
                data_last_row,
                value_col,
            ],
            "fill": {"color": "#CC0000"},
        }
    )
    chart.set_title(
        {
            "name": title,
            "name_font": {"name": "Arial", "size": 12, "bold": True},
        }
    )
    chart.set_legend({"none": True})
    chart.set_size({"width": 580, "height": 360})
    worksheet.insert_chart(
        data_start_row,
        value_col + 3,
        chart,
        {"x_offset": 10},
    )


def _add_segment_profile_chart(
    workbook: Any,
    worksheet: Any,
    sheet_name: str,
    payload: dict,
    start_row: int,
    title: str,
) -> None:
    """Add a chart for the serialized single-cut result in a segment profile."""

    if "distribution" in payload:
        n_rows = len(payload.get("distribution", {}))
        _add_segment_chart(
            workbook, worksheet, sheet_name, start_row + 1, n_rows, 1, 2, title, "Count"
        )
        return
    if "selections" in payload:
        n_rows = sum(
            1
            for selection in payload.get("selections", {}).values()
            if int(selection.get("count", 0)) != 0
        )
        _add_segment_chart(
            workbook, worksheet, sheet_name, start_row + 1, n_rows, 1, 2, title, "Count"
        )
        return
    if "mean" in payload:
        _add_segment_chart(
            workbook, worksheet, sheet_name, start_row, 3, 0, 1, title, "Statistic"
        )


def _add_expected_vs_realized_chart(
    workbook: Any,
    worksheet: Any,
    sheet_name: str,
    data_start_row: int,
    n_metrics: int,
    title: str,
) -> None:
    """Insert expected and realized metric columns as a clustered chart."""

    if not _can_add_chart(workbook, worksheet) or n_metrics <= 0:
        return

    chart = workbook.add_chart({"type": "column", "subtype": "clustered"})
    data_last_row = data_start_row + n_metrics - 1
    header_row = data_start_row - 1
    for offset, color in ((1, "#CC0000"), (2, "#666666")):
        chart.add_series(
            {
                "name": [sheet_name, header_row, offset],
                "categories": [sheet_name, data_start_row, 0, data_last_row, 0],
                "values": [sheet_name, data_start_row, offset, data_last_row, offset],
                "fill": {"color": color},
            }
        )
    chart.set_title(
        {
            "name": title,
            "name_font": {"name": "Arial", "size": 12, "bold": True},
        }
    )
    chart.set_legend({"position": "right", "font": {"name": "Arial", "size": 9}})
    chart.set_size({"width": 580, "height": 360})
    worksheet.insert_chart(data_start_row, 6, chart, {"x_offset": 10})


def _can_add_chart(workbook: Any, worksheet: Any) -> bool:
    return hasattr(workbook, "add_chart") and hasattr(worksheet, "insert_chart")


def _write_segment_profile_body(
    workbook: Any,
    ws: Any,
    result: CrossCutResult,
    schema: SurveySchema,
    formats: dict[str, Any],
    start_row: int,
    sheet_name: str,
) -> int:
    result_table = result.result_table
    filter_question_id = result.source_question_ids[0]
    target_question_id = result_table.get(
        "target_question_id", result.source_question_ids[1]
    )
    filter_spec = schema.get_question(filter_question_id)
    _write(ws, start_row, 0, "Filter applied:", formats["bold"])
    _write(ws, start_row + 1, 0, _question_label(schema, filter_question_id))
    _write(
        ws,
        start_row + 2,
        0,
        f"= {_filter_value_label(result_table.get('filter_expr'), filter_spec)}",
    )
    _write(ws, start_row + 3, 0, "Target question:", formats["bold"])
    _write(ws, start_row + 4, 0, _question_label(schema, target_question_id))
    _write(ws, start_row + 6, 0, "Target distribution", formats["bold"])
    target_payload = result_table.get("target_result", {})
    target_start_row = start_row + 7
    last_row = _write_serialized_single_cut_body(
        ws,
        target_payload,
        formats,
        target_start_row,
        _segment_profile_display_mode(result.display_mode),
    )
    _add_segment_profile_chart(
        workbook,
        ws,
        sheet_name,
        target_payload,
        target_start_row,
        result.synthetic_question_title,
    )
    return last_row


def _write_serialized_single_cut_body(
    ws: Any,
    payload: dict,
    formats: dict[str, Any],
    start_row: int,
    display_mode: str = "all",
) -> int:
    if "distribution" in payload:
        return _write_distribution_dict(
            ws, payload.get("distribution", {}), formats, start_row, display_mode
        )
    if "selections" in payload:
        return _write_selections_dict(
            ws, payload.get("selections", {}), formats, start_row
        )
    if "rows" in payload:
        return _write_grid_dict(ws, payload.get("rows", {}), formats, start_row)
    if "mean" in payload:
        return _write_numeric_dict(ws, payload, formats, start_row)
    _write(ws, start_row, 0, "Unsupported target result")
    return start_row


def _write_distribution_dict(
    ws: Any,
    distribution: dict,
    formats: dict[str, Any],
    start_row: int,
    display_mode: str = "all",
) -> int:
    headers = ["Code", "Label", "Count"]
    include_rate = display_mode != "counts"
    if include_rate:
        headers.append("%")
    _write_header_row(ws, start_row, headers, formats)
    row_index = start_row + 1
    for code, payload in sorted(distribution.items(), key=lambda item: _sort_key(item[0])):
        _write(ws, row_index, 0, code)
        _write(ws, row_index, 1, payload.get("label", ""))
        _write(ws, row_index, 2, payload.get("count", 0), formats["count"])
        if include_rate:
            _write(ws, row_index, 3, payload.get("rate", 0.0), formats["pct"])
        row_index += 1
    return row_index - 1


def _write_selections_dict(
    ws: Any,
    selections: dict,
    formats: dict[str, Any],
    start_row: int,
) -> int:
    _write_header_row(
        ws,
        start_row,
        ["Sub-Column ID", "Label", "Count selected", "Selection %"],
        formats,
    )
    row_index = start_row + 1
    sorted_items = sorted(
        selections.items(),
        key=lambda item: int(item[1].get("count", 0)),
        reverse=True,
    )
    for sub_column_id, payload in sorted_items:
        if int(payload.get("count", 0)) == 0:
            continue
        _write(ws, row_index, 0, sub_column_id)
        _write(ws, row_index, 1, payload.get("label", ""))
        _write(ws, row_index, 2, payload.get("count", 0), formats["count"])
        _write(ws, row_index, 3, payload.get("selection_rate", 0.0), formats["pct"])
        row_index += 1
    return row_index - 1


def _write_numeric_dict(
    ws: Any,
    payload: dict,
    formats: dict[str, Any],
    start_row: int,
) -> int:
    rows = [
        ("Mean", payload.get("mean")),
        ("Median", payload.get("median")),
        ("Std Dev", payload.get("std")),
        ("Min", payload.get("min_val")),
        ("Max", payload.get("max_val")),
        ("25th percentile", payload.get("percentiles", {}).get(25)),
        ("50th percentile", payload.get("percentiles", {}).get(50)),
        ("75th percentile", payload.get("percentiles", {}).get(75)),
    ]
    row_index = start_row
    for label, value in rows:
        _write(ws, row_index, 0, label)
        _write(ws, row_index, 1, value, formats["stat"])
        row_index += 1
    return row_index - 1


def _write_grid_dict(
    ws: Any,
    rows: dict,
    formats: dict[str, Any],
    start_row: int,
) -> int:
    row_index = start_row
    for row_id, row_payload in rows.items():
        _write(
            ws,
            row_index,
            0,
            f"{row_id} (n={row_payload.get('valid_n', 0)})",
            formats["bold"],
        )
        row_index += 1
        row_index = _write_distribution_dict(
            ws,
            row_payload.get("distribution", {}),
            formats,
            row_index,
        ) + 2
    return row_index - 1


def _write_group_comparison_body(
    workbook: Any,
    ws: Any,
    result: CrossCutResult,
    schema: SurveySchema,
    formats: dict[str, Any],
    start_row: int,
    sheet_name: str,
) -> int:
    result_table = result.result_table
    if result_table.get("selection_rate_rows"):
        return _write_selection_rate_group_comparison_body(
            workbook,
            ws,
            result,
            schema,
            formats,
            start_row,
            sheet_name,
        )
    if result_table.get("rank_rows"):
        return _write_rank_group_comparison_body(
            workbook,
            ws,
            result,
            schema,
            formats,
            start_row,
            sheet_name,
        )
    if result_table.get("allocation_rows"):
        return _write_rank_group_comparison_body(
            workbook,
            ws,
            result,
            schema,
            formats,
            start_row,
            sheet_name,
            rows_key="allocation_rows",
            metric_label="Numeric allocation metric:",
            mean_title="Mean allocation by segment",
            median_title="Median allocation by segment",
            detail_title="Per-option segment allocation detail",
        )
    if result_table.get("nps_entities"):
        return _write_nps_group_comparison_body(
            workbook,
            ws,
            result,
            schema,
            formats,
            start_row,
            sheet_name,
        )
    if result_table.get("grid_rows"):
        return _write_grid_rated_group_comparison_body(
            workbook,
            ws,
            result,
            schema,
            formats,
            start_row,
            sheet_name,
        )
    segment_question_id = result_table.get(
        "segment_question_id", result.source_question_ids[0]
    )
    metric_question_id = result_table.get(
        "metric_question_id", result.source_question_ids[1]
    )
    _write(ws, start_row, 0, "Segments (rows):", formats["bold"])
    _write(ws, start_row + 1, 0, _question_label(schema, segment_question_id))
    _write(ws, start_row + 2, 0, "Metric (columns):", formats["bold"])
    _write(ws, start_row + 3, 0, _question_label(schema, metric_question_id))
    table_start = start_row + 5
    _write(ws, table_start, 0, "Per-segment comparison", formats["bold"])
    _write_header_row(ws, table_start + 1, ["Segment", "Label", "N", "Mean", "Median", "Std"], formats)
    row_index = table_start + 2
    per_segment = result_table.get("per_segment", {})
    for segment, payload in per_segment.items():
        _write(ws, row_index, 0, segment)
        _write(ws, row_index, 1, payload.get("label", ""))
        _write(ws, row_index, 2, payload.get("n", 0), formats["count"])
        _write(ws, row_index, 3, payload.get("mean"), formats["stat"])
        _write(ws, row_index, 4, payload.get("median"), formats["stat"])
        _write(ws, row_index, 5, payload.get("std"), formats["stat"])
        row_index += 1
    overall = result_table.get("overall", {})
    _write(ws, row_index, 0, "Overall", formats["bold"])
    _write(ws, row_index, 2, overall.get("n", 0), formats["count"])
    _write(ws, row_index, 3, overall.get("mean"), formats["stat"])
    _write(ws, row_index, 4, overall.get("median"), formats["stat"])
    _write(ws, row_index, 5, overall.get("std"), formats["stat"])
    if per_segment:
        _add_segment_chart(
            workbook,
            ws,
            sheet_name,
            table_start + 2,
            len(per_segment),
            1,
            3,
            result.synthetic_question_title,
            "Mean",
        )
    return row_index


def _write_selection_rate_group_comparison_body(
    workbook: Any,
    ws: Any,
    result: CrossCutResult,
    schema: SurveySchema,
    formats: dict[str, Any],
    start_row: int,
    sheet_name: str,
) -> int:
    del workbook, sheet_name
    result_table = result.result_table
    segment_question_id = result_table.get(
        "segment_question_id", result.source_question_ids[0]
    )
    metric_question_id = result_table.get(
        "metric_question_id", result.source_question_ids[1]
    )
    _write(ws, start_row, 0, "Segments (rows):", formats["bold"])
    _write(ws, start_row + 1, 0, _question_label(schema, segment_question_id))
    _write(ws, start_row + 2, 0, "Selection-rate metric:", formats["bold"])
    _write(ws, start_row + 3, 0, _question_label(schema, metric_question_id))
    table_start = start_row + 5
    rows_payload = result_table.get("selection_rate_rows", {}) or {}
    segment_columns: list[tuple[Any, str]] = []
    for row_payload in rows_payload.values():
        if not isinstance(row_payload, dict):
            continue
        for segment, payload in (row_payload.get("per_segment", {}) or {}).items():
            label = payload.get("label", str(segment)) if isinstance(payload, dict) else str(segment)
            key = (segment, label)
            if key not in segment_columns:
                segment_columns.append(key)

    _write(ws, table_start, 0, "Selection rate by segment", formats["bold"])
    _write_header_row(
        ws,
        table_start + 1,
        ["Metric row", *(label for _segment, label in segment_columns), "Overall"],
        formats,
    )
    row_index = table_start + 2
    for row_id, row_payload in rows_payload.items():
        row_label = row_payload.get("label", row_id)
        _write(ws, row_index, 0, row_label)
        per_segment = row_payload.get("per_segment", {}) if isinstance(row_payload, dict) else {}
        for col_offset, (segment, _label) in enumerate(segment_columns, start=1):
            payload = per_segment.get(segment, {}) if isinstance(per_segment, dict) else {}
            _write(ws, row_index, col_offset, payload.get("selection_rate"), formats["pct"])
        overall = row_payload.get("overall", {}) if isinstance(row_payload, dict) else {}
        _write(ws, row_index, len(segment_columns) + 1, overall.get("selection_rate"), formats["pct"])
        row_index += 1

    detail_start = row_index + 2
    _write(ws, detail_start, 0, "Per-row segment selection detail", formats["bold"])
    _write_header_row(
        ws,
        detail_start + 1,
        ["Metric row", "Segment", "Label", "N", "Count", "% selected"],
        formats,
    )
    row_index = detail_start + 2
    for row_id, row_payload in rows_payload.items():
        row_label = row_payload.get("label", row_id)
        for segment, payload in (row_payload.get("per_segment", {}) or {}).items():
            _write(ws, row_index, 0, row_label)
            _write(ws, row_index, 1, segment)
            _write(ws, row_index, 2, payload.get("label", ""))
            _write(ws, row_index, 3, payload.get("n", 0), formats["count"])
            _write(ws, row_index, 4, payload.get("count", 0), formats["count"])
            _write(ws, row_index, 5, payload.get("selection_rate"), formats["pct"])
            row_index += 1
        overall = row_payload.get("overall", {}) if isinstance(row_payload, dict) else {}
        _write(ws, row_index, 0, row_label, formats["bold"])
        _write(ws, row_index, 1, "Overall", formats["bold"])
        _write(ws, row_index, 3, overall.get("n", 0), formats["count"])
        _write(ws, row_index, 4, overall.get("count", 0), formats["count"])
        _write(ws, row_index, 5, overall.get("selection_rate"), formats["pct"])
        row_index += 1
    if row_index > detail_start + 2:
        _apply_autofilter(ws, detail_start + 1, row_index - 1, 6)
    return row_index


def _write_rank_group_comparison_body(
    workbook: Any,
    ws: Any,
    result: CrossCutResult,
    schema: SurveySchema,
    formats: dict[str, Any],
    start_row: int,
    sheet_name: str,
    *,
    rows_key: str = "rank_rows",
    metric_label: str = "Rank-order metric:",
    mean_title: str = "Mean rank by segment",
    median_title: str = "Median rank by segment",
    detail_title: str = "Per-row segment rank detail",
    row_header: str = "Metric row",
) -> int:
    del workbook, sheet_name
    result_table = result.result_table
    segment_question_id = result_table.get(
        "segment_question_id", result.source_question_ids[0]
    )
    metric_question_id = result_table.get(
        "metric_question_id", result.source_question_ids[1]
    )
    _write(ws, start_row, 0, "Segments (rows):", formats["bold"])
    _write(ws, start_row + 1, 0, _question_label(schema, segment_question_id))
    _write(ws, start_row + 2, 0, metric_label, formats["bold"])
    _write(ws, start_row + 3, 0, _question_label(schema, metric_question_id))
    table_start = start_row + 5
    rows_payload = result_table.get(rows_key, {}) or {}
    segment_columns: list[tuple[Any, str]] = []
    for row_payload in rows_payload.values():
        if not isinstance(row_payload, dict):
            continue
        for segment, payload in (row_payload.get("per_segment", {}) or {}).items():
            label = payload.get("label", str(segment)) if isinstance(payload, dict) else str(segment)
            key = (segment, label)
            if key not in segment_columns:
                segment_columns.append(key)

    _write(ws, table_start, 0, mean_title, formats["bold"])
    _write_header_row(
        ws,
        table_start + 1,
        [row_header, *(label for _segment, label in segment_columns), "Overall"],
        formats,
    )
    row_index = table_start + 2
    for row_id, row_payload in rows_payload.items():
        row_label = row_payload.get("label", row_id)
        _write(ws, row_index, 0, row_label)
        per_segment = row_payload.get("per_segment", {}) if isinstance(row_payload, dict) else {}
        for col_offset, (segment, _label) in enumerate(segment_columns, start=1):
            payload = per_segment.get(segment, {}) if isinstance(per_segment, dict) else {}
            _write(ws, row_index, col_offset, payload.get("mean"), formats["stat"])
        overall = row_payload.get("overall", {}) if isinstance(row_payload, dict) else {}
        _write(ws, row_index, len(segment_columns) + 1, overall.get("mean"), formats["stat"])
        row_index += 1

    median_start = row_index + 2
    _write(ws, median_start, 0, median_title, formats["bold"])
    _write_header_row(
        ws,
        median_start + 1,
        [row_header, *(label for _segment, label in segment_columns), "Overall"],
        formats,
    )
    row_index = median_start + 2
    for row_id, row_payload in rows_payload.items():
        row_label = row_payload.get("label", row_id)
        _write(ws, row_index, 0, row_label)
        per_segment = row_payload.get("per_segment", {}) if isinstance(row_payload, dict) else {}
        for col_offset, (segment, _label) in enumerate(segment_columns, start=1):
            payload = per_segment.get(segment, {}) if isinstance(per_segment, dict) else {}
            _write(ws, row_index, col_offset, payload.get("median"), formats["stat"])
        overall = row_payload.get("overall", {}) if isinstance(row_payload, dict) else {}
        _write(ws, row_index, len(segment_columns) + 1, overall.get("median"), formats["stat"])
        row_index += 1

    detail_start = row_index + 2
    _write(ws, detail_start, 0, detail_title, formats["bold"])
    _write_header_row(
        ws,
        detail_start + 1,
        [row_header, "Segment", "Label", "N", "Mean", "Median", "Std"],
        formats,
    )
    row_index = detail_start + 2
    for row_id, row_payload in rows_payload.items():
        row_label = row_payload.get("label", row_id)
        for segment, payload in (row_payload.get("per_segment", {}) or {}).items():
            _write(ws, row_index, 0, row_label)
            _write(ws, row_index, 1, segment)
            _write(ws, row_index, 2, payload.get("label", ""))
            _write(ws, row_index, 3, payload.get("n", 0), formats["count"])
            _write(ws, row_index, 4, payload.get("mean"), formats["stat"])
            _write(ws, row_index, 5, payload.get("median"), formats["stat"])
            _write(ws, row_index, 6, payload.get("std"), formats["stat"])
            row_index += 1
        overall = row_payload.get("overall", {}) if isinstance(row_payload, dict) else {}
        _write(ws, row_index, 0, row_label, formats["bold"])
        _write(ws, row_index, 1, "Overall", formats["bold"])
        _write(ws, row_index, 3, overall.get("n", 0), formats["count"])
        _write(ws, row_index, 4, overall.get("mean"), formats["stat"])
        _write(ws, row_index, 5, overall.get("median"), formats["stat"])
        _write(ws, row_index, 6, overall.get("std"), formats["stat"])
        row_index += 1
    if row_index > detail_start + 2:
        _apply_autofilter(ws, detail_start + 1, row_index - 1, 7)
    return row_index


def _write_nps_group_comparison_body(
    workbook: Any,
    ws: Any,
    result: CrossCutResult,
    schema: SurveySchema,
    formats: dict[str, Any],
    start_row: int,
    sheet_name: str,
) -> int:
    del workbook, sheet_name
    result_table = result.result_table
    segment_question_id = result_table.get(
        "segment_question_id", result.source_question_ids[0]
    )
    metric_question_id = result_table.get(
        "metric_question_id", result.source_question_ids[1]
    )
    _write(ws, start_row, 0, "Segments (rows):", formats["bold"])
    _write(ws, start_row + 1, 0, _question_label(schema, segment_question_id))
    _write(ws, start_row + 2, 0, "NPS metric:", formats["bold"])
    _write(ws, start_row + 3, 0, _question_label(schema, metric_question_id))
    table_start = start_row + 5
    segment_columns: list[tuple[Any, str]] = []
    for entity_payload in (result_table.get("nps_entities", {}) or {}).values():
        if not isinstance(entity_payload, dict):
            continue
        for segment, payload in (entity_payload.get("per_segment", {}) or {}).items():
            label = payload.get("label", str(segment)) if isinstance(payload, dict) else str(segment)
            key = (segment, label)
            if key not in segment_columns:
                segment_columns.append(key)

    _write(ws, table_start, 0, "NPS by segment", formats["bold"])
    _write_header_row(
        ws,
        table_start + 1,
        ["Entity", *(label for _segment, label in segment_columns), "Overall"],
        formats,
    )
    row_index = table_start + 2
    for entity_id, entity_payload in (result_table.get("nps_entities", {}) or {}).items():
        entity_label = (
            entity_payload.get("label", entity_id)
            if isinstance(entity_payload, dict)
            else entity_id
        )
        _write(ws, row_index, 0, entity_label)
        per_segment = (
            entity_payload.get("per_segment", {})
            if isinstance(entity_payload, dict)
            else {}
        )
        for col_offset, (segment, _label) in enumerate(segment_columns, start=1):
            payload = per_segment.get(segment, {}) if isinstance(per_segment, dict) else {}
            _write(ws, row_index, col_offset, payload.get("nps_score"), formats["stat"])
        overall = (
            entity_payload.get("overall", {})
            if isinstance(entity_payload, dict)
            else {}
        )
        _write(ws, row_index, len(segment_columns) + 1, overall.get("nps_score"), formats["stat"])
        row_index += 1

    detail_start = row_index + 2
    _write(ws, detail_start, 0, "Per-entity segment NPS detail", formats["bold"])
    _write_header_row(
        ws,
        detail_start + 1,
        [
            "Entity",
            "Segment",
            "Label",
            "Valid N",
            "Promoters %",
            "Passives %",
            "Detractors %",
            "NPS",
        ],
        formats,
    )
    row_index = detail_start + 2
    for entity_id, entity_payload in (result_table.get("nps_entities", {}) or {}).items():
        entity_label = (
            entity_payload.get("label", entity_id)
            if isinstance(entity_payload, dict)
            else entity_id
        )
        per_segment = (
            entity_payload.get("per_segment", {})
            if isinstance(entity_payload, dict)
            else {}
        )
        for segment, payload in per_segment.items():
            _write(ws, row_index, 0, entity_label)
            _write(ws, row_index, 1, segment)
            _write(ws, row_index, 2, payload.get("label", ""))
            _write(ws, row_index, 3, payload.get("valid_n", payload.get("n", 0)), formats["count"])
            _write(ws, row_index, 4, payload.get("pct_promoters"), formats["pct"])
            _write(ws, row_index, 5, payload.get("pct_passives"), formats["pct"])
            _write(ws, row_index, 6, payload.get("pct_detractors"), formats["pct"])
            _write(ws, row_index, 7, payload.get("nps_score"), formats["stat"])
            row_index += 1
        overall = (
            entity_payload.get("overall", {})
            if isinstance(entity_payload, dict)
            else {}
        )
        _write(ws, row_index, 0, entity_label, formats["bold"])
        _write(ws, row_index, 1, "Overall", formats["bold"])
        _write(ws, row_index, 3, overall.get("valid_n", overall.get("n", 0)), formats["count"])
        _write(ws, row_index, 4, overall.get("pct_promoters"), formats["pct"])
        _write(ws, row_index, 5, overall.get("pct_passives"), formats["pct"])
        _write(ws, row_index, 6, overall.get("pct_detractors"), formats["pct"])
        _write(ws, row_index, 7, overall.get("nps_score"), formats["stat"])
        row_index += 1
    if row_index > detail_start + 2:
        _apply_autofilter(ws, detail_start + 1, row_index - 1, 8)
    return row_index


def _write_grid_rated_group_comparison_body(
    workbook: Any,
    ws: Any,
    result: CrossCutResult,
    schema: SurveySchema,
    formats: dict[str, Any],
    start_row: int,
    sheet_name: str,
) -> int:
    result_table = result.result_table
    segment_question_id = result_table.get(
        "segment_question_id", result.source_question_ids[0]
    )
    metric_question_id = result_table.get(
        "metric_question_id", result.source_question_ids[1]
    )
    _write(ws, start_row, 0, "Segments (rows):", formats["bold"])
    _write(ws, start_row + 1, 0, _question_label(schema, segment_question_id))
    _write(ws, start_row + 2, 0, "Grid-rated metric:", formats["bold"])
    _write(ws, start_row + 3, 0, _question_label(schema, metric_question_id))
    table_start = start_row + 5
    segment_columns: list[tuple[Any, str]] = []
    for row_payload in (result_table.get("grid_rows", {}) or {}).values():
        if not isinstance(row_payload, dict):
            continue
        for segment, payload in (row_payload.get("per_segment", {}) or {}).items():
            label = payload.get("label", str(segment)) if isinstance(payload, dict) else str(segment)
            key = (segment, label)
            if key not in segment_columns:
                segment_columns.append(key)

    _write(ws, table_start, 0, "Mean rating by segment", formats["bold"])
    _write_header_row(
        ws,
        table_start + 1,
        ["Metric row", *(label for _segment, label in segment_columns), "Overall"],
        formats,
    )
    row_index = table_start + 2
    for row_id, row_payload in (result_table.get("grid_rows", {}) or {}).items():
        row_label = row_payload.get("label", row_id)
        _write(ws, row_index, 0, row_label)
        per_segment = row_payload.get("per_segment", {}) if isinstance(row_payload, dict) else {}
        for col_offset, (segment, _label) in enumerate(segment_columns, start=1):
            payload = per_segment.get(segment, {}) if isinstance(per_segment, dict) else {}
            _write(ws, row_index, col_offset, payload.get("mean"), formats["stat"])
        overall = row_payload.get("overall", {}) if isinstance(row_payload, dict) else {}
        _write(ws, row_index, len(segment_columns) + 1, overall.get("mean"), formats["stat"])
        row_index += 1

    median_start = row_index + 2
    _write(ws, median_start, 0, "Median rating by segment", formats["bold"])
    _write_header_row(
        ws,
        median_start + 1,
        ["Metric row", *(label for _segment, label in segment_columns), "Overall"],
        formats,
    )
    row_index = median_start + 2
    for row_id, row_payload in (result_table.get("grid_rows", {}) or {}).items():
        row_label = row_payload.get("label", row_id)
        _write(ws, row_index, 0, row_label)
        per_segment = row_payload.get("per_segment", {}) if isinstance(row_payload, dict) else {}
        for col_offset, (segment, _label) in enumerate(segment_columns, start=1):
            payload = per_segment.get(segment, {}) if isinstance(per_segment, dict) else {}
            _write(ws, row_index, col_offset, payload.get("median"), formats["stat"])
        overall = row_payload.get("overall", {}) if isinstance(row_payload, dict) else {}
        _write(ws, row_index, len(segment_columns) + 1, overall.get("median"), formats["stat"])
        row_index += 1

    detail_start = row_index + 2
    _write(ws, detail_start, 0, "Per-row segment comparison", formats["bold"])
    _write_header_row(
        ws,
        detail_start + 1,
        ["Metric row", "Segment", "Label", "N", "Mean", "Median", "Std"],
        formats,
    )
    row_index = detail_start + 2
    for row_id, row_payload in (result_table.get("grid_rows", {}) or {}).items():
        row_label = row_payload.get("label", row_id)
        for segment, payload in (row_payload.get("per_segment", {}) or {}).items():
            _write(ws, row_index, 0, row_label)
            _write(ws, row_index, 1, segment)
            _write(ws, row_index, 2, payload.get("label", ""))
            _write(ws, row_index, 3, payload.get("n", 0), formats["count"])
            _write(ws, row_index, 4, payload.get("mean"), formats["stat"])
            _write(ws, row_index, 5, payload.get("median"), formats["stat"])
            _write(ws, row_index, 6, payload.get("std"), formats["stat"])
            row_index += 1
        overall = row_payload.get("overall", {})
        _write(ws, row_index, 0, row_label, formats["bold"])
        _write(ws, row_index, 1, "Overall", formats["bold"])
        _write(ws, row_index, 3, overall.get("n", 0), formats["count"])
        _write(ws, row_index, 4, overall.get("mean"), formats["stat"])
        _write(ws, row_index, 5, overall.get("median"), formats["stat"])
        _write(ws, row_index, 6, overall.get("std"), formats["stat"])
        row_index += 1
    if row_index > detail_start + 2:
        _apply_autofilter(ws, detail_start + 1, row_index - 1, 7)
    return row_index


def _write_expected_vs_realized_body(
    workbook: Any,
    ws: Any,
    result: CrossCutResult,
    schema: SurveySchema,
    formats: dict[str, Any],
    start_row: int,
    sheet_name: str,
) -> int:
    result_table = result.result_table
    expected = result_table.get("expected", {})
    realized = result_table.get("realized", {})
    gap = result_table.get("gap", {})
    expected_question_id = result_table.get(
        "expected_question_id", result.source_question_ids[0]
    )
    realized_question_id = result_table.get(
        "realized_question_id", result.source_question_ids[1]
    )
    _write(ws, start_row, 0, "Expected:", formats["bold"])
    _write(ws, start_row + 1, 0, _question_label(schema, expected_question_id))
    _write(ws, start_row + 2, 0, "Realized:", formats["bold"])
    _write(ws, start_row + 3, 0, _question_label(schema, realized_question_id))
    table_start = start_row + 5
    _write(ws, table_start, 0, "Expected vs Realized", formats["bold"])
    _write_header_row(ws, table_start + 1, ["Metric", "Expected", "Realized", "Gap"], formats)
    rows = [
        ("Mean", "mean", formats["stat"]),
        ("Median", "median", formats["stat"]),
        ("Std", "std", formats["stat"]),
        ("N", "valid_n", formats["count"]),
    ]
    row_index = table_start + 2
    for label, key, cell_format in rows:
        _write(ws, row_index, 0, label)
        _write(ws, row_index, 1, expected.get(key), cell_format)
        _write(ws, row_index, 2, realized.get(key), cell_format)
        _write(ws, row_index, 3, gap.get(key), cell_format)
        row_index += 1
    row_index += 1
    _write(ws, row_index, 0, f"Paired N: {result_table.get('paired_n', 0)}", formats["bold"])
    _add_expected_vs_realized_chart(
        workbook,
        ws,
        sheet_name,
        table_start + 2,
        3,
        result.synthetic_question_title,
    )
    return row_index


def _write_calculation_log(
    workbook: Any, log: CalculationLog, formats: dict[str, Any]
) -> None:
    ws = workbook.add_worksheet("Calculation_Log")
    headers = [
        "Output Sheet",
        "Metric Name",
        "Source Question ID",
        "Source Columns",
        "Filter",
        "Numerator",
        "Denominator",
        "Formula",
        "Value Raw",
        "Valid N",
        "Missing N",
        "Timestamp",
    ]
    _write_header_row(ws, 0, headers, formats)
    for row_index, record in enumerate(log.all_records(), start=1):
        _write(ws, row_index, 0, record.output_sheet)
        _write(ws, row_index, 1, record.metric_name)
        _write(ws, row_index, 2, record.source_question_id)
        _write(ws, row_index, 3, ", ".join(record.source_columns))
        _write(ws, row_index, 4, record.filter_expr or "")
        _write(ws, row_index, 5, "" if record.numerator is None else record.numerator, formats["count"])
        _write(ws, row_index, 6, "" if record.denominator is None else record.denominator, formats["count"])
        _write(ws, row_index, 7, record.formula)
        _write(ws, row_index, 8, record.value_raw, formats["stat"])
        _write(ws, row_index, 9, record.valid_n, formats["count"])
        _write(ws, row_index, 10, record.missing_n, formats["count"])
        _write(ws, row_index, 11, record.timestamp.isoformat())
    ws.freeze_panes(1, 0)
    _autofit(ws)


def _write_filter_log(
    workbook: Any,
    cross_cut_results: list[CrossCutResult],
    cc_sheet_names: dict[str, str],
    formats: dict[str, Any],
) -> None:
    ws = workbook.add_worksheet("Filter_Log")
    _write_header_row(
        ws,
        0,
        [
            "Cross Cut ID",
            "Title",
            "Filter Expression",
            "Description",
            "Affects Sheet",
        ],
        formats,
    )
    row_index = 1
    for result in cross_cut_results:
        filter_expr = result.result_table.get("filter_expr")
        if not filter_expr:
            continue
        _write(ws, row_index, 0, result.cross_cut_id)
        _write(ws, row_index, 1, result.synthetic_question_title)
        _write(ws, row_index, 2, filter_expr)
        _write(
            ws,
            row_index,
            3,
            result.result_table.get("filter_mask_description") or filter_expr,
        )
        _write(ws, row_index, 4, cc_sheet_names.get(result.cross_cut_id, ""))
        row_index += 1
    ws.freeze_panes(1, 0)
    _autofit(ws)


def _write_filtered_calculation_log(
    workbook: Any,
    filtered_results: list[FilteredSingleCutResult],
    log: CalculationLog,
    formats: dict[str, Any],
) -> None:
    filtered_log = CalculationLog()
    for record in _filtered_audit_records(filtered_results, log):
        filtered_log.record(record)
    _write_calculation_log(workbook, filtered_log, formats)


def _write_filtered_filter_log(
    workbook: Any,
    filtered_results: list[FilteredSingleCutResult],
    fsc_sheet_names: dict[int, str],
    schema: SurveySchema,
    formats: dict[str, Any],
) -> None:
    ws = workbook.add_worksheet("Filter_Log")
    headers = ["Sheet Name", "Filter Question", "Filter Value", "Filter Description"]
    _write_header_row(ws, 0, headers, formats)
    row_index = 1
    for result_index, result in enumerate(filtered_results):
        sheet_name = fsc_sheet_names[result_index]
        for filter_spec in result.filters_applied:
            _write(ws, row_index, 0, sheet_name)
            _write(ws, row_index, 1, filter_spec.filter_question_id)
            _write(
                ws,
                row_index,
                2,
                "" if filter_spec.filter_value is None else filter_spec.filter_value,
            )
            _write(
                ws,
                row_index,
                3,
                _filter_description(filter_spec, schema, compact_breakdown=False),
            )
            row_index += 1
    ws.freeze_panes(1, 0)
    _autofit(ws)


def _write_data_quality(
    workbook: Any,
    quality_report: DataQualityReport,
    formats: dict[str, Any],
) -> None:
    ws = workbook.add_worksheet("Data_Quality")
    totals = [
        ("Total rows:", quality_report.total_rows),
        ("Total columns:", quality_report.total_columns),
        ("Columns in datamap:", quality_report.columns_in_datamap),
        ("Columns NOT in datamap:", len(quality_report.columns_not_in_datamap)),
    ]
    row_index = 0
    for label, value in totals:
        _write(ws, row_index, 0, label, formats["bold"])
        _write(ws, row_index, 1, value)
        row_index += 1

    row_index += 1
    _write(ws, row_index, 0, "Per-column missing %", formats["bold"])
    row_index += 1
    _write_header_row(ws, row_index, ["Column", "Missing %"], formats)
    row_index += 1
    for column, pct in sorted(
        quality_report.per_column_missing_pct.items(),
        key=lambda item: item[1],
        reverse=True,
    ):
        if pct > 0:
            _write(ws, row_index, 0, column)
            _write(ws, row_index, 1, pct, formats["pct"])
            row_index += 1

    row_index += 1
    _write(ws, row_index, 0, "Per-column out-of-range %", formats["bold"])
    row_index += 1
    _write_header_row(ws, row_index, ["Column", "Out-of-range %"], formats)
    row_index += 1
    for column, pct in quality_report.per_column_out_of_range_pct.items():
        if pct > 0:
            _write(ws, row_index, 0, column)
            _write(ws, row_index, 1, pct, formats["pct"])
            row_index += 1

    row_index += 1
    _write(ws, row_index, 0, "Coercion log", formats["bold"])
    row_index += 1
    _write_header_row(
        ws,
        row_index,
        ["Column", "From Type", "To Type", "Values Coerced", "Rows Affected"],
        formats,
    )
    row_index += 1
    for entry in quality_report.coercion_log:
        coerced_values = ", ".join(str(value) for value in entry.get("values_coerced", ()))
        _write(ws, row_index, 0, entry.get("column", ""))
        _write(ws, row_index, 1, entry.get("from_type", ""))
        _write(ws, row_index, 2, entry.get("to_type", ""))
        _write(ws, row_index, 3, coerced_values[:100])
        _write(ws, row_index, 4, entry.get("rows_affected", ""))
        row_index += 1

    ws.freeze_panes(6, 0)
    _autofit(ws)


def _write_warnings(
    workbook: Any,
    quality_report: DataQualityReport,
    results: list[SingleCutResult],
    skips: list[SkipRecord],
    formats: dict[str, Any],
    cross_cut_results: list[CrossCutResult],
    cross_cut_skips: list[SkipRecord],
) -> None:
    ws = workbook.add_worksheet("Warnings")
    _write_header_row(ws, 0, ["Source", "Warning"], formats)
    row_index = 1
    for warning in quality_report.warnings:
        _write_row(ws, row_index, ["data_quality", warning])
        row_index += 1
    for result in results:
        for warning in result.warnings:
            _write_row(ws, row_index, [f"result:{result.question_id}", warning])
            row_index += 1
    for result in cross_cut_results:
        for warning in result.warnings:
            _write_row(ws, row_index, [f"cross_result:{result.cross_cut_id}", warning])
            row_index += 1
    for skip in skips:
        if skip.details:
            _write_row(ws, row_index, [f"skip:{skip.question_id}", skip.details])
            row_index += 1
    for skip in cross_cut_skips:
        if skip.details:
            _write_row(ws, row_index, [f"cross_skip:{skip.question_id}", skip.details])
            row_index += 1
    decoder_warnings = tuple(getattr(quality_report, "decoder_warnings", ()) or ())
    if decoder_warnings:
        row_index = _write_decoder_warnings_section(
            ws,
            row_index + 1,
            decoder_warnings,
            formats,
        )
    ws.freeze_panes(1, 0)
    _autofit(ws)


def _write_decoder_warnings_section(
    worksheet: Any,
    row_index: int,
    decoder_warnings: tuple[dict, ...],
    formats: dict[str, Any],
) -> int:
    worksheet.write(row_index, 0, "Data Quality — Decoder Warnings", formats["bold"])
    row_index += 1

    summary_counts: dict[tuple[str, str], int] = defaultdict(int)
    for warning in decoder_warnings:
        question_id = str(warning.get("question_id", ""))
        action = str(warning.get("action", ""))
        summary_counts[(question_id, action)] += 1

    _write_header_row(worksheet, row_index, ["question_id", "action", "count"], formats)
    row_index += 1
    for (question_id, action), count in sorted(summary_counts.items()):
        _write_row(worksheet, row_index, [question_id, action, count])
        row_index += 1

    row_index += 1
    _write_header_row(
        worksheet,
        row_index,
        ["question_id", "column", "row", "raw_value", "action"],
        formats,
    )
    row_index += 1
    for warning in decoder_warnings:
        _write_row(
            worksheet,
            row_index,
            [
                warning.get("question_id", ""),
                warning.get("column", ""),
                warning.get("row", ""),
                warning.get("raw_value", ""),
                warning.get("action", ""),
            ],
        )
        row_index += 1
    return row_index


def _safe_sheet_name(name: str) -> str:
    """Excel sheet names: 31 chars max, no special chars."""
    cleaned = re.sub(r"[\[\]:*?/\\']+", "", name)
    cleaned = cleaned.strip() or "Theme"
    return cleaned[:31]


def _question_text(schema: SurveySchema, question_id: str) -> str:
    spec = schema.get_question(question_id)
    return spec.question_text if spec is not None else question_id


def _question_label(schema: SurveySchema, question_id: str) -> str:
    return f"{question_id}: {_question_text(schema, question_id)}"


def _question_labels_text(
    schema: SurveySchema,
    question_ids: tuple[str, ...],
) -> str:
    return " × ".join(_question_label(schema, question_id) for question_id in question_ids)


def _filter_value_label(filter_expr: str | None, filter_spec: Any) -> str:
    if not filter_expr:
        return "<unknown>"
    value_text = filter_expr.split("==", 1)[-1].strip().strip("\"'")
    try:
        value: int | str = int(value_text)
    except ValueError:
        value = value_text
    if filter_spec is not None and value in filter_spec.option_map:
        return str(filter_spec.option_map[value])
    return str(value)


def _segment_profile_display_mode(display_mode: str) -> str:
    return "counts" if display_mode == "counts" else "both"


def _filter_log_entry_count(cross_cut_results: list[CrossCutResult]) -> int:
    return sum(1 for result in cross_cut_results if result.result_table.get("filter_expr"))


def _filtered_sheet_names(
    filtered_results: list[FilteredSingleCutResult],
    used_sheet_names: set[str],
) -> dict[int, str]:
    target_counts: dict[str, int] = defaultdict(int)
    for result in filtered_results:
        target_counts[result.target_question_id] += 1

    target_seen: dict[str, int] = defaultdict(int)
    sheet_names: dict[int, str] = {}
    for index, result in enumerate(filtered_results):
        target_seen[result.target_question_id] += 1
        base_name = f"FSC_{result.target_question_id}"
        if target_counts[result.target_question_id] > 1:
            base_name = f"{base_name}_{target_seen[result.target_question_id]:02d}"
        sheet_names[index] = _unique_sheet_name(base_name, used_sheet_names)
    return sheet_names


def _filter_description(
    filter_spec: FilterSpec,
    schema: SurveySchema,
    compact_breakdown: bool,
) -> str:
    if filter_spec.filter_value is None:
        if compact_breakdown:
            return f"{filter_spec.filter_question_id} (breakdown)"
        return f"{filter_spec.filter_question_id} (breakdown - no specific value)"

    label = _filter_option_label(filter_spec, schema)
    if label:
        return (
            f"{filter_spec.filter_question_id} == {filter_spec.filter_value} "
            f"({label})"
        )
    return f"{filter_spec.filter_question_id} == {filter_spec.filter_value}"


def _filter_option_label(filter_spec: FilterSpec, schema: SurveySchema) -> str | None:
    spec = schema.get_question(filter_spec.filter_question_id)
    if spec is None:
        return None
    value = filter_spec.filter_value
    if value in spec.option_map:
        return str(spec.option_map[value])
    value_as_string = str(value)
    for option_code, label in spec.option_map.items():
        if str(option_code) == value_as_string:
            return str(label)
    return None


def _filtered_audit_records(
    filtered_results: list[FilteredSingleCutResult],
    log: CalculationLog,
) -> tuple[AuditRecord, ...]:
    source_ids = {
        result.target_question_id
        for result in filtered_results
    }
    for result in filtered_results:
        for filter_spec in result.filters_applied:
            source_ids.add(filter_spec.filter_question_id)

    records = []
    for record in log.all_records():
        if record.source_question_id in source_ids:
            records.append(record)
            continue
        if record.output_sheet.startswith(("FSC_", "CC_")):
            records.append(record)
    return tuple(records)


def _cross_tab_column_codes(matrix: dict, column_label_map: dict) -> list[Any]:
    codes = set(column_label_map)
    for row_payload in matrix.values():
        if isinstance(row_payload, dict):
            codes.update(row_payload)
    return sorted(codes, key=_sort_key)


def _label_for_code(label_map: dict, code: Any) -> str:
    if code in label_map:
        return str(label_map[code])
    code_as_string = str(code)
    if code_as_string in label_map:
        return str(label_map[code_as_string])
    return code_as_string


def _nested_lookup(matrix: dict, row_code: Any, column_code: Any) -> Any:
    row_payload = matrix.get(row_code, matrix.get(str(row_code), {}))
    if not isinstance(row_payload, dict):
        return ""
    return row_payload.get(column_code, row_payload.get(str(column_code), 0))


def _unique_sheet_name(name: str, used_sheet_names: set[str]) -> str:
    base = _safe_sheet_name(name)
    candidate = base
    suffix = 1
    while candidate in used_sheet_names:
        suffix_text = f"_{suffix}"
        candidate = f"{base[:31 - len(suffix_text)]}{suffix_text}"
        suffix += 1
    used_sheet_names.add(candidate)
    return candidate


def _autofit(ws: Any, max_width: int = 60) -> None:
    """Set widths from tracked cell content, capped at max_width."""
    widths = getattr(ws, "_sie_widths", {})
    for col_index, width in widths.items():
        ws.set_column(col_index, col_index, min(width + 2, max_width))


def _denominator_description(result: SingleCutResult) -> str:
    policy = result.denominator_policy.value
    if policy == "VALID_RESPONSES":
        return f"Valid responses (n={result.valid_n})"
    if policy == "ALL_RESPONDENTS":
        return f"All respondents (n={result.valid_n})"
    if policy == "EXPOSED_TO_QUESTION":
        return f"Exposed to question (n={result.valid_n})"
    return policy


def _write_header_row(
    ws: Any, row_index: int, headers: list[str], formats: dict[str, Any]
) -> None:
    for column_index, header in enumerate(headers):
        _write(ws, row_index, column_index, header, formats["header"])


def _write_row(
    ws: Any, row_index: int, values: list[Any], cell_format: Any | None = None
) -> None:
    for column_index, value in enumerate(values):
        _write(ws, row_index, column_index, value, cell_format)


def _write(ws: Any, row: int, col: int, value: Any, cell_format: Any | None = None) -> None:
    if isinstance(value, float) and math.isnan(value):
        value = ""
    ws.write(row, col, value, cell_format)
    _track_cell(ws, row, col, value)


def _write_blank(ws: Any, row: int, col: int, cell_format: Any | None = None) -> None:
    if hasattr(ws, "write_blank"):
        ws.write_blank(row, col, None, cell_format)
    else:
        ws.write(row, col, "", cell_format)
    _track_cell(ws, row, col, "")


def _write_formula(
    ws: Any,
    row: int,
    col: int,
    formula: str,
    cell_format: Any | None = None,
    value: Any | None = None,
) -> None:
    if hasattr(ws, "write_formula"):
        ws.write_formula(row, col, formula, cell_format, value)
    else:
        ws.write(row, col, formula, cell_format)
    _track_cell(ws, row, col, formula)


def _write_url(
    ws: Any,
    row: int,
    col: int,
    url: str,
    cell_format: Any,
    string: str,
) -> None:
    ws.write_url(row, col, url, cell_format, string=string)
    _track_cell(ws, row, col, string)


def _merge_range(
    ws: Any,
    first_row: int,
    first_col: int,
    last_row: int,
    last_col: int,
    value: Any,
    cell_format: Any | None = None,
) -> None:
    if hasattr(ws, "merge_range"):
        ws.merge_range(first_row, first_col, last_row, last_col, value, cell_format)
        _track_cell(ws, first_row, first_col, value)
        return
    _write(ws, first_row, first_col, value, cell_format)


def _track_cell(ws: Any, row: int, col: int, value: Any) -> None:
    if not hasattr(ws, "_sie_widths"):
        setattr(ws, "_sie_widths", {})
    if not hasattr(ws, "_sie_last_row"):
        setattr(ws, "_sie_last_row", 0)
    widths = getattr(ws, "_sie_widths")
    widths[col] = max(widths.get(col, 0), len(_cell_display(value)))
    setattr(ws, "_sie_last_row", max(getattr(ws, "_sie_last_row"), row))


def _find_last_used_row(ws: Any) -> int:
    return int(getattr(ws, "_sie_last_row", 0))


def _cell_display(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.isoformat()
    return str(value)


def _sort_key(value: Any) -> tuple[int, Any]:
    if isinstance(value, (int, float)) and not isinstance(value, bool):
        return (0, value)
    return (1, str(value))


def _rate_sort_value(value: Any) -> float:
    try:
        rate = float(value)
    except (TypeError, ValueError):
        return float("-inf")
    return rate if not math.isnan(rate) else float("-inf")


def _quote_sheet_name(sheet_name: str) -> str:
    return sheet_name.replace("'", "''")


class _FallbackFormat:
    def __init__(self, properties: dict[str, Any], style_id: int) -> None:
        self.properties = properties
        self.style_id = style_id


class _FallbackChart:
    def __init__(self, properties: dict[str, Any]) -> None:
        self.properties = properties
        self.series: list[dict[str, Any]] = []
        self.title: str = ""
        self.legend_none = False
        self.chart_index = 0

    def add_series(self, series: dict[str, Any]) -> None:
        self.series.append(series)

    def set_title(self, properties: dict[str, Any]) -> None:
        self.title = str(properties.get("name", ""))

    def set_x_axis(self, _properties: dict[str, Any]) -> None:
        return

    def set_y_axis(self, _properties: dict[str, Any]) -> None:
        return

    def set_legend(self, properties: dict[str, Any]) -> None:
        self.legend_none = bool(properties.get("none"))

    def set_size(self, _properties: dict[str, Any]) -> None:
        return

    def set_chartarea(self, _properties: dict[str, Any]) -> None:
        return

    def to_xml(self) -> str:
        title_xml = ""
        if self.title:
            title_xml = (
                "<c:title><c:tx><c:rich><a:bodyPr/><a:lstStyle/>"
                f"<a:p><a:r><a:t>{escape(self.title)}</a:t></a:r></a:p>"
                "</c:rich></c:tx><c:layout/></c:title>"
            )
        series_xml = "".join(
            self._series_xml(series, index) for index, series in enumerate(self.series)
        )
        legend_xml = "" if self.legend_none else (
            '<c:legend><c:legendPos val="r"/><c:layout/></c:legend>'
        )
        return (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<c:chartSpace xmlns:c="http://schemas.openxmlformats.org/drawingml/2006/chart" '
            'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
            'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
            "<c:chart>"
            f"{title_xml}<c:plotArea><c:layout/><c:barChart>"
            '<c:barDir val="col"/><c:grouping val="clustered"/>'
            f"{series_xml}<c:axId val=\"123456\"/><c:axId val=\"123457\"/>"
            "</c:barChart>"
            '<c:catAx><c:axId val="123456"/><c:scaling><c:orientation val="minMax"/>'
            '</c:scaling><c:axPos val="b"/><c:tickLblPos val="nextTo"/>'
            '<c:crossAx val="123457"/><c:crosses val="autoZero"/></c:catAx>'
            '<c:valAx><c:axId val="123457"/><c:scaling><c:orientation val="minMax"/>'
            '</c:scaling><c:axPos val="l"/><c:majorGridlines/>'
            '<c:numFmt formatCode="General" sourceLinked="1"/><c:tickLblPos val="nextTo"/>'
            '<c:crossAx val="123456"/><c:crosses val="autoZero"/></c:valAx>'
            f"</c:plotArea>{legend_xml}<c:plotVisOnly val=\"1\"/>"
            "</c:chart></c:chartSpace>"
        )

    def _series_xml(self, series: dict[str, Any], index: int) -> str:
        name_xml = _chart_text_ref_xml(series.get("name"), "tx")
        categories_xml = _chart_text_ref_xml(series.get("categories"), "cat")
        values_xml = _chart_num_ref_xml(series.get("values"), "val")
        return (
            f'<c:ser><c:idx val="{index}"/><c:order val="{index}"/>'
            f"{name_xml}{categories_xml}{values_xml}</c:ser>"
        )


class _FallbackWorkbook:
    def __init__(self, output_path: str) -> None:
        self.output_path = output_path
        self._worksheets: list[_FallbackWorksheet] = []

    def add_format(self, properties: dict[str, Any]) -> _FallbackFormat:
        return _FallbackFormat(properties, _fallback_style_id(properties))

    def add_chart(self, properties: dict[str, Any]) -> "_FallbackChart":
        return _FallbackChart(properties)

    def add_worksheet(self, name: str) -> "_FallbackWorksheet":
        worksheet = _FallbackWorksheet(name)
        self._worksheets.append(worksheet)
        return worksheet

    def close(self) -> None:
        Path(self.output_path).parent.mkdir(parents=True, exist_ok=True)
        self._assign_chart_parts()
        with ZipFile(self.output_path, "w", ZIP_DEFLATED) as archive:
            archive.writestr("[Content_Types].xml", self._content_types_xml())
            archive.writestr("_rels/.rels", self._root_rels_xml())
            archive.writestr("docProps/core.xml", self._core_xml())
            archive.writestr("docProps/app.xml", self._app_xml())
            archive.writestr("xl/workbook.xml", self._workbook_xml())
            archive.writestr("xl/_rels/workbook.xml.rels", self._workbook_rels_xml())
            archive.writestr("xl/styles.xml", self._styles_xml())
            for index, worksheet in enumerate(self._worksheets, start=1):
                archive.writestr(
                    f"xl/worksheets/sheet{index}.xml", worksheet.to_xml()
                )
                if worksheet.charts or worksheet.tables:
                    archive.writestr(
                        f"xl/worksheets/_rels/sheet{index}.xml.rels",
                        worksheet.sheet_rels_xml(),
                    )
                if worksheet.charts:
                    archive.writestr(
                        f"xl/drawings/drawing{worksheet.drawing_index}.xml",
                        worksheet.drawing_xml(),
                    )
                    archive.writestr(
                        f"xl/drawings/_rels/drawing{worksheet.drawing_index}.xml.rels",
                        worksheet.drawing_rels_xml(),
                    )
                    for chart, _row, _col in worksheet.charts:
                        archive.writestr(
                            f"xl/charts/chart{chart.chart_index}.xml",
                            chart.to_xml(),
                        )
                for table in worksheet.tables:
                    archive.writestr(
                        f"xl/tables/table{table['index']}.xml",
                        worksheet.table_xml(table),
                    )

    def _assign_chart_parts(self) -> None:
        drawing_index = 1
        chart_index = 1
        table_index = 1
        for worksheet in self._worksheets:
            if not worksheet.charts:
                worksheet.drawing_index = 0
            else:
                worksheet.drawing_index = drawing_index
                worksheet.drawing_rid = "rId1"
                drawing_index += 1
                for chart, _row, _col in worksheet.charts:
                    chart.chart_index = chart_index
                    chart_index += 1
            next_rid = 2 if worksheet.charts else 1
            for table in worksheet.tables:
                table["index"] = table_index
                table["name"] = f"Table{table_index}"
                table["rid"] = f"rId{next_rid}"
                table_index += 1
                next_rid += 1

    def _content_types_xml(self) -> str:
        sheet_overrides = "".join(
            f'<Override PartName="/xl/worksheets/sheet{index}.xml" '
            'ContentType="application/vnd.openxmlformats-officedocument.'
            'spreadsheetml.worksheet+xml"/>'
            for index in range(1, len(self._worksheets) + 1)
        )
        drawing_overrides = "".join(
            f'<Override PartName="/xl/drawings/drawing{worksheet.drawing_index}.xml" '
            'ContentType="application/vnd.openxmlformats-officedocument.'
            'drawing+xml"/>'
            for worksheet in self._worksheets
            if worksheet.charts
        )
        chart_overrides = "".join(
            f'<Override PartName="/xl/charts/chart{chart.chart_index}.xml" '
            'ContentType="application/vnd.openxmlformats-officedocument.'
            'drawingml.chart+xml"/>'
            for worksheet in self._worksheets
            for chart, _row, _col in worksheet.charts
        )
        table_overrides = "".join(
            f'<Override PartName="/xl/tables/table{table["index"]}.xml" '
            'ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.table+xml"/>'
            for worksheet in self._worksheets
            for table in worksheet.tables
        )
        return (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
            '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
            '<Default Extension="xml" ContentType="application/xml"/>'
            '<Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>'
            '<Override PartName="/xl/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.styles+xml"/>'
            '<Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>'
            '<Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/>'
            f"{sheet_overrides}{drawing_overrides}{chart_overrides}{table_overrides}</Types>"
        )

    def _root_rels_xml(self) -> str:
        return (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="xl/workbook.xml"/>'
            '<Relationship Id="rId2" Type="http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties" Target="docProps/core.xml"/>'
            '<Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties" Target="docProps/app.xml"/>'
            "</Relationships>"
        )

    def _workbook_xml(self) -> str:
        sheets = "".join(
            f'<sheet name="{escape(worksheet.name)}" sheetId="{index}" r:id="rId{index}"/>'
            for index, worksheet in enumerate(self._worksheets, start=1)
        )
        return (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" '
            'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
            f"<sheets>{sheets}</sheets></workbook>"
        )

    def _workbook_rels_xml(self) -> str:
        sheet_rels = "".join(
            f'<Relationship Id="rId{index}" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" '
            f'Target="worksheets/sheet{index}.xml"/>'
            for index in range(1, len(self._worksheets) + 1)
        )
        style_id = len(self._worksheets) + 1
        return (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            f"{sheet_rels}"
            f'<Relationship Id="rId{style_id}" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" '
            'Target="styles.xml"/>'
            "</Relationships>"
        )

    def _styles_xml(self) -> str:
        return (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<styleSheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">'
            '<numFmts count="3">'
            '<numFmt numFmtId="164" formatCode="0.0%"/>'
            '<numFmt numFmtId="165" formatCode="#,##0"/>'
            '<numFmt numFmtId="166" formatCode="0.00"/>'
            "</numFmts>"
            '<fonts count="3"><font><sz val="11"/><name val="Calibri"/></font>'
            '<font><b/><sz val="11"/><name val="Calibri"/></font>'
            '<font><u/><color rgb="000000FF"/><sz val="11"/><name val="Calibri"/></font></fonts>'
            '<fills count="3"><fill><patternFill patternType="none"/></fill>'
            '<fill><patternFill patternType="gray125"/></fill>'
            '<fill><patternFill patternType="solid"><fgColor rgb="FFF2F2F2"/><bgColor indexed="64"/></patternFill></fill></fills>'
            '<borders count="2"><border><left/><right/><top/><bottom/><diagonal/></border>'
            '<border><left style="thin"/><right style="thin"/><top style="thin"/><bottom style="thin"/><diagonal/></border></borders>'
            '<cellStyleXfs count="1"><xf numFmtId="0" fontId="0" fillId="0" borderId="0"/></cellStyleXfs>'
            '<cellXfs count="7">'
            '<xf numFmtId="0" fontId="0" fillId="0" borderId="0" xfId="0"/>'
            '<xf numFmtId="0" fontId="1" fillId="0" borderId="0" xfId="0" applyFont="1"/>'
            '<xf numFmtId="0" fontId="1" fillId="2" borderId="1" xfId="0" applyFont="1" applyFill="1" applyBorder="1"/>'
            '<xf numFmtId="164" fontId="0" fillId="0" borderId="0" xfId="0" applyNumberFormat="1"/>'
            '<xf numFmtId="165" fontId="0" fillId="0" borderId="0" xfId="0" applyNumberFormat="1"/>'
            '<xf numFmtId="166" fontId="0" fillId="0" borderId="0" xfId="0" applyNumberFormat="1"/>'
            '<xf numFmtId="0" fontId="2" fillId="0" borderId="0" xfId="0" applyFont="1"/>'
            "</cellXfs>"
            '<cellStyles count="1"><cellStyle name="Normal" xfId="0" builtinId="0"/></cellStyles>'
            "</styleSheet>"
        )

    def _core_xml(self) -> str:
        timestamp = datetime.now(timezone.utc).isoformat()
        return (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" '
            'xmlns:dc="http://purl.org/dc/elements/1.1/" '
            'xmlns:dcterms="http://purl.org/dc/terms/" '
            'xmlns:dcmitype="http://purl.org/dc/dcmitype/" '
            'xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">'
            "<dc:creator>Survey Insight Engine</dc:creator>"
            f'<dcterms:created xsi:type="dcterms:W3CDTF">{timestamp}</dcterms:created>'
            f'<dcterms:modified xsi:type="dcterms:W3CDTF">{timestamp}</dcterms:modified>'
            "</cp:coreProperties>"
        )

    def _app_xml(self) -> str:
        return (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties" '
            'xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">'
            "<Application>Survey Insight Engine</Application>"
            "</Properties>"
        )


class _FallbackWorksheet:
    def __init__(self, name: str) -> None:
        self.name = name
        self.cells: dict[tuple[int, int], tuple[Any, int]] = {}
        self.hyperlinks: dict[tuple[int, int], str] = {}
        self.widths: dict[int, float] = {}
        self.freeze: tuple[int, int] | None = None
        self.autofilter_range: tuple[int, int, int, int] | None = None
        self.validations: list[tuple[str, list[str]]] = []
        self.tables: list[dict[str, Any]] = []
        self.charts: list[tuple[_FallbackChart, int, int]] = []
        self.drawing_index = 0
        self.drawing_rid = "rId1"

    def write(self, row: int, col: int, value: Any, cell_format: _FallbackFormat | None = None) -> None:
        style_id = cell_format.style_id if cell_format is not None else 0
        self.cells[(row, col)] = (value, style_id)

    def write_url(
        self,
        row: int,
        col: int,
        url: str,
        cell_format: _FallbackFormat | None = None,
        string: str | None = None,
    ) -> None:
        display = string if string is not None else url
        self.write(row, col, display, cell_format)
        if url.startswith("internal:"):
            self.hyperlinks[(row, col)] = url.removeprefix("internal:")

    def freeze_panes(self, row: int, col: int) -> None:
        self.freeze = (row, col)

    def set_column(self, first_col: int, last_col: int, width: float) -> None:
        for col in range(first_col, last_col + 1):
            self.widths[col] = width

    def autofilter(
        self, first_row: int, first_col: int, last_row: int, last_col: int
    ) -> None:
        self.autofilter_range = (first_row, first_col, last_row, last_col)

    def data_validation(
        self,
        first_row: int,
        first_col: int,
        last_row: int,
        last_col: int,
        options: dict[str, Any],
    ) -> None:
        if options.get("validate") != "list":
            return
        source = options.get("source", [])
        if isinstance(source, str):
            values = [source]
        else:
            values = [str(value) for value in source]
        sqref = f"{_cell_ref(first_row, first_col)}:{_cell_ref(last_row, last_col)}"
        if first_row == last_row and first_col == last_col:
            sqref = _cell_ref(first_row, first_col)
        self.validations.append((sqref, values))

    def add_table(
        self,
        first_row: int,
        first_col: int,
        last_row: int,
        last_col: int,
        options: dict[str, Any] | None = None,
    ) -> None:
        options = options or {}
        for offset, column in enumerate(options.get("columns", [])):
            header = column.get("header", "")
            if header:
                self.write(first_row, first_col + offset, header)
        headers = [
            str(column.get("header", f"Column{index + 1}"))
            for index, column in enumerate(options.get("columns", []))
        ]
        self.tables.append(
            {
                "first_row": first_row,
                "first_col": first_col,
                "last_row": last_row,
                "last_col": last_col,
                "headers": headers,
            }
        )

    def insert_chart(
        self,
        row: int,
        col: int,
        chart: _FallbackChart,
        _options: dict[str, Any] | None = None,
    ) -> None:
        self.charts.append((chart, row, col))

    def to_xml(self) -> str:
        rows: dict[int, list[tuple[int, Any, int]]] = defaultdict(list)
        for (row, col), (value, style_id) in self.cells.items():
            rows[row].append((col, value, style_id))
        dimension_xml = self._dimension_xml()
        row_xml = []
        for row in sorted(rows):
            cells = "".join(
                self._cell_xml(row, col, value, style_id)
                for col, value, style_id in sorted(rows[row])
            )
            row_xml.append(f'<row r="{row + 1}">{cells}</row>')
        cols_xml = ""
        if self.widths:
            cols_xml = "<cols>" + "".join(
                f'<col min="{col + 1}" max="{col + 1}" width="{width}" customWidth="1"/>'
                for col, width in sorted(self.widths.items())
            ) + "</cols>"
        hyperlinks_xml = ""
        if self.hyperlinks:
            links = "".join(
                f'<hyperlink ref="{_cell_ref(row, col)}" location="{escape(location)}" display="{escape(str(self.cells[(row, col)][0]))}"/>'
                for (row, col), location in sorted(self.hyperlinks.items())
            )
            hyperlinks_xml = f"<hyperlinks>{links}</hyperlinks>"
        validations_xml = ""
        if self.validations:
            validations = []
            for sqref, values in self.validations:
                list_formula = '"' + ",".join(values)[:250] + '"'
                validations.append(
                    f'<dataValidation type="list" allowBlank="1" sqref="{sqref}">'
                    f"<formula1>{escape(list_formula)}</formula1></dataValidation>"
                )
            validations_xml = (
                f'<dataValidations count="{len(validations)}">'
                f"{''.join(validations)}</dataValidations>"
            )
        autofilter_xml = ""
        if self.autofilter_range is not None:
            first_row, first_col, last_row, last_col = self.autofilter_range
            autofilter_xml = (
                f'<autoFilter ref="{_cell_ref(first_row, first_col)}:'
                f'{_cell_ref(last_row, last_col)}"/>'
            )
        sheet_views = self._sheet_views_xml()
        drawing_xml = f'<drawing r:id="{self.drawing_rid}"/>' if self.charts else ""
        table_parts_xml = ""
        if self.tables:
            parts = "".join(
                f'<tablePart r:id="{table["rid"]}"/>' for table in self.tables
            )
            table_parts_xml = f'<tableParts count="{len(self.tables)}">{parts}</tableParts>'
        return (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" '
            'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
            f"{dimension_xml}{sheet_views}{cols_xml}<sheetData>{''.join(row_xml)}</sheetData>{autofilter_xml}{validations_xml}{hyperlinks_xml}{drawing_xml}{table_parts_xml}"
            "</worksheet>"
        )

    def sheet_rels_xml(self) -> str:
        relationships = []
        if self.charts:
            relationships.append(
                f'<Relationship Id="{self.drawing_rid}" '
                'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/drawing" '
                f'Target="../drawings/drawing{self.drawing_index}.xml"/>'
            )
        for table in self.tables:
            relationships.append(
                f'<Relationship Id="{table["rid"]}" '
                'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/table" '
                f'Target="../tables/table{table["index"]}.xml"/>'
            )
        return (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            f"{''.join(relationships)}</Relationships>"
        )

    def table_xml(self, table: dict[str, Any]) -> str:
        ref = (
            f"{_cell_ref(table['first_row'], table['first_col'])}:"
            f"{_cell_ref(table['last_row'], table['last_col'])}"
        )
        columns = "".join(
            f'<tableColumn id="{index}" name="{escape(header)}"/>'
            for index, header in enumerate(table["headers"], start=1)
        )
        return (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<table xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" '
            f'id="{table["index"]}" name="{table["name"]}" displayName="{table["name"]}" '
            f'ref="{ref}" totalsRowShown="0">'
            f'<autoFilter ref="{ref}"/>'
            f'<tableColumns count="{len(table["headers"])}">{columns}</tableColumns>'
            '<tableStyleInfo name="TableStyleLight9" showFirstColumn="0" '
            'showLastColumn="0" showRowStripes="1" showColumnStripes="0"/>'
            '</table>'
        )

    def drawing_rels_xml(self) -> str:
        relationships = "".join(
            f'<Relationship Id="rId{index}" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/chart" '
            f'Target="../charts/chart{chart.chart_index}.xml"/>'
            for index, (chart, _row, _col) in enumerate(self.charts, start=1)
        )
        return (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            f"{relationships}</Relationships>"
        )

    def drawing_xml(self) -> str:
        anchors = "".join(
            self._chart_anchor_xml(chart, row, col, index)
            for index, (chart, row, col) in enumerate(self.charts, start=1)
        )
        return (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<xdr:wsDr xmlns:xdr="http://schemas.openxmlformats.org/drawingml/2006/spreadsheetDrawing" '
            'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
            f"{anchors}</xdr:wsDr>"
        )

    def _chart_anchor_xml(
        self,
        _chart: _FallbackChart,
        row: int,
        col: int,
        rel_index: int,
    ) -> str:
        return (
            "<xdr:twoCellAnchor>"
            f"<xdr:from><xdr:col>{col}</xdr:col><xdr:colOff>0</xdr:colOff>"
            f"<xdr:row>{row}</xdr:row><xdr:rowOff>0</xdr:rowOff></xdr:from>"
            f"<xdr:to><xdr:col>{col + 8}</xdr:col><xdr:colOff>0</xdr:colOff>"
            f"<xdr:row>{row + 20}</xdr:row><xdr:rowOff>0</xdr:rowOff></xdr:to>"
            '<xdr:graphicFrame macro="">'
            "<xdr:nvGraphicFramePr>"
            f'<xdr:cNvPr id="{rel_index + 1}" name="Chart {rel_index}"/>'
            "<xdr:cNvGraphicFramePr/>"
            "</xdr:nvGraphicFramePr>"
            "<xdr:xfrm><a:off x=\"0\" y=\"0\"/><a:ext cx=\"0\" cy=\"0\"/></xdr:xfrm>"
            '<a:graphic><a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/chart">'
            f'<c:chart xmlns:c="http://schemas.openxmlformats.org/drawingml/2006/chart" '
            f'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
            f'r:id="rId{rel_index}"/>'
            "</a:graphicData></a:graphic>"
            "</xdr:graphicFrame><xdr:clientData/></xdr:twoCellAnchor>"
        )

    def _dimension_xml(self) -> str:
        if not self.cells:
            return '<dimension ref="A1"/>'
        max_row = max(row for row, _ in self.cells)
        max_col = max(col for _, col in self.cells)
        return f'<dimension ref="A1:{_cell_ref(max_row, max_col)}"/>'

    def _sheet_views_xml(self) -> str:
        if self.freeze is None:
            return '<sheetViews><sheetView workbookViewId="0"/></sheetViews>'
        row, col = self.freeze
        top_left = _cell_ref(row, col)
        return (
            '<sheetViews><sheetView workbookViewId="0">'
            f'<pane ySplit="{row}" xSplit="{col}" topLeftCell="{top_left}" '
            'activePane="bottomRight" state="frozen"/>'
            "</sheetView></sheetViews>"
        )

    def _cell_xml(self, row: int, col: int, value: Any, style_id: int) -> str:
        cell_ref = _cell_ref(row, col)
        style_attr = f' s="{style_id}"' if style_id else ""
        if value is None or value == "":
            return f'<c r="{cell_ref}"{style_attr}/>'
        if isinstance(value, bool):
            return f'<c r="{cell_ref}" t="b"{style_attr}><v>{1 if value else 0}</v></c>'
        if isinstance(value, (int, float)) and not isinstance(value, bool):
            return f'<c r="{cell_ref}"{style_attr}><v>{value}</v></c>'
        text = escape(str(value))
        preserve = ' xml:space="preserve"' if str(value).strip() != str(value) else ""
        return f'<c r="{cell_ref}" t="inlineStr"{style_attr}><is><t{preserve}>{text}</t></is></c>'


def _chart_text_ref_xml(reference: Any, tag_name: str) -> str:
    if isinstance(reference, str):
        return f"<c:{tag_name}><c:v>{escape(reference)}</c:v></c:{tag_name}>"
    formula = _chart_reference_formula(reference)
    if not formula:
        return ""
    return (
        f"<c:{tag_name}><c:strRef><c:f>{escape(formula)}</c:f></c:strRef>"
        f"</c:{tag_name}>"
    )


def _chart_num_ref_xml(reference: Any, tag_name: str) -> str:
    formula = _chart_reference_formula(reference)
    if not formula:
        return ""
    return (
        f"<c:{tag_name}><c:numRef><c:f>{escape(formula)}</c:f></c:numRef>"
        f"</c:{tag_name}>"
    )


def _chart_reference_formula(reference: Any) -> str:
    if not isinstance(reference, list) or len(reference) not in (3, 5):
        return ""
    sheet_name = str(reference[0]).replace("'", "''")
    if len(reference) == 3:
        row = int(reference[1])
        col = int(reference[2])
        return f"'{sheet_name}'!{_absolute_cell_ref(row, col)}"
    first_row = int(reference[1])
    first_col = int(reference[2])
    last_row = int(reference[3])
    last_col = int(reference[4])
    return (
        f"'{sheet_name}'!{_absolute_cell_ref(first_row, first_col)}:"
        f"{_absolute_cell_ref(last_row, last_col)}"
    )


def _absolute_cell_ref(row: int, col: int) -> str:
    return f"${_column_name(col)}${row + 1}"


def _fallback_style_id(properties: dict[str, Any]) -> int:
    if properties.get("num_format") == "0.0%":
        return 3
    if properties.get("num_format") == "#,##0":
        return 4
    if properties.get("num_format") == "0.00":
        return 5
    if properties.get("underline"):
        return 6
    if properties.get("bold") and properties.get("bg_color"):
        return 2
    if properties.get("bold"):
        return 1
    if properties.get("italic"):
        return 0
    return 0


def _cell_ref(row: int, col: int) -> str:
    return f"{_column_name(col)}{row + 1}"


def _column_name(col: int) -> str:
    name = ""
    col += 1
    while col:
        col, remainder = divmod(col - 1, 26)
        name = chr(65 + remainder) + name
    return name
